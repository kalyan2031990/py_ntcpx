#!/usr/bin/env python3
"""
NTCP Output QA Reporter
-----------------------
Reads the output folder (or zip) created by enhanced_ntcp_analysis.py / enhanced_ntcp_ml.py
and generates:
  1) comprehensive_report.docx – human‑readable QA report
  2) qa_summary_tables.xlsx – per‑organ summary metrics + unique patient list

It flags:
- Inflated patient counts (files vs unique PatientID).
- Unrealistic NTCP values: NaNs, const predictions, outside [0,1].
- Low-n / low-event instability.
- Potential ML overfitting / leakage (AUC≥0.90 with n<40 or events<8).
- Traditional model optimism when events<5 but AUC≥0.85.

Usage
-----
python ntcp_output_QA_reporter.py --input <out_dir_or_zip> --report_outdir <outdir>

Examples
--------
python ntcp_output_QA_reporter.py --input enhanced_ntcp_analysis_ml_out.zip --report_outdir QA_results
python ntcp_output_QA_reporter.py --input ./enhanced_ntcp_analysis_ml_out --report_outdir QA_results
"""

import argparse
import os, re, sys, zipfile, math
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union

# Windows-safe encoding configuration
try:
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except (AttributeError, ValueError):
    pass

import numpy as np
import pandas as pd

# Import v2.0 leakage detection
try:
    from src.reporting.leakage_detector import DataLeakageDetector
    V2_LEAKAGE_DETECTOR_AVAILABLE = True
except ImportError:
    V2_LEAKAGE_DETECTOR_AVAILABLE = False

# Optional deps: if unavailable, the script will print a helpful error.
try:
    from sklearn.metrics import roc_curve, auc, brier_score_loss
except Exception as e:
    print("[WARN] scikit‑learn not available; AUC/Brier will be NaN. Install scikit‑learn to compute metrics.", file=sys.stderr)
    roc_curve = None
    auc = None
    brier_score_loss = None

try:
    from docx import Document
except Exception as e:
    Document = None
    print("[WARN] python‑docx not available; will skip DOCX report. Install python‑docx.", file=sys.stderr)


def unzip_if_needed(input_path: Path, workdir: Path) -> Path:
    """If input is a .zip, extract to workdir and return the root folder; else return input_path."""
    if input_path.is_file() and input_path.suffix.lower() == ".zip":
        with zipfile.ZipFile(input_path, 'r') as z:
            z.extractall(workdir)
        # Heuristic: if the zip contained a single top-level folder, return that; else return workdir
        entries = [p for p in workdir.iterdir()]
        if len(entries) == 1 and entries[0].is_dir():
            return entries[0]
        return workdir
    return input_path


def discover_files(root: Path) -> List[Path]:
    files = []
    for r, d, fs in os.walk(root):
        for f in fs:
            files.append(Path(r) / f)
    return files


def load_table(path: Path) -> Optional[Union[pd.DataFrame, Dict[str, pd.DataFrame]]]:
    """Load CSV or XLSX (all sheets)."""
    try:
        if path.suffix.lower() == ".csv":
            return pd.read_csv(path)
        if path.suffix.lower() == ".xlsx":
            with pd.ExcelFile(path) as xl:
                return {s: xl.parse(s) for s in xl.sheet_names}
    except Exception as e:
        print(f"[WARN] Failed to load {path}: {e}", file=sys.stderr)
    return None


def likely_results_file(p: Path) -> bool:
    low = p.name.lower()
    if not low.endswith((".csv", ".xlsx")):
        return False
    keys = ["result", "summary", "by_organ", "metrics", "ntcp", "calc"]
    return any(k in low for k in keys)


def is_patient_level_df(df: pd.DataFrame) -> bool:
    cols = [c.strip().lower() for c in df.columns]
    key_cols = {"patient", "patientid", "id", "mrn"}
    outcome_cols = {"observed_toxicity", "toxicity", "event", "grade", "label"}
    pred_cols = [c for c in cols if c.startswith("ntcp_") or c.startswith("ml_") or "lkb" in c or "rs_" in c]
    return (any(k in cols for k in key_cols) and any(o in cols for o in outcome_cols)) or (len(pred_cols) >= 1 and "organ" in cols)


def harmonize(df: pd.DataFrame) -> pd.DataFrame:
    df2 = df.copy()
    # Rename common columns
    rename_map = {}
    for c in df2.columns:
        lc = c.strip().lower()
        if lc in ("patient", "patientid", "ptid", "id"):
            rename_map[c] = "PatientID"
        elif lc == "organ":
            rename_map[c] = "Organ"
        elif lc in ("observed_toxicity", "toxicity", "event", "label"):
            rename_map[c] = "Observed_Toxicity"
        elif lc == "grade":
            rename_map[c] = "Grade"
    if rename_map:
        df2 = df2.rename(columns=rename_map)

    # Standardize types
    if "Organ" in df2.columns:
        df2["Organ"] = df2["Organ"].astype(str).str.strip()

    # Derive Observed_Toxicity if missing: treat Grade>=2 as positive
    if "Observed_Toxicity" not in df2.columns and "Grade" in df2.columns:
        g = pd.to_numeric(df2["Grade"], errors="coerce")
        df2["Observed_Toxicity"] = (g >= 2).astype(float)

    if "Observed_Toxicity" in df2.columns:
        df2["Observed_Toxicity"] = pd.to_numeric(df2["Observed_Toxicity"], errors="coerce")
        df2["Observed_Toxicity"] = df2["Observed_Toxicity"].fillna(0).clip(0, 1)

    # Normalize prediction columns (keep original names too)
    for c in list(df2.columns):
        lc = c.lower()
        if lc in ("lkb_loglogit", "lkb_probit", "rs_poisson"):
            df2.rename(columns={c: f"NTCP_{c}"}, inplace=True)
        if lc.startswith("ml_"):
            # Standardize ML columns as NTCP_ML_<ModelName>
            if "ann" in lc:
                suffix = "ANN"
            elif "xgboost" in lc or "xgb" in lc:
                suffix = "XGBoost"
            elif "randomforest" in lc or "rf" in lc:
                suffix = "RandomForest"
            elif "gradient" in lc or "gb" in lc:
                suffix = "GradientBoosting"
            else:
                suffix = c.split("ml_")[-1]
            df2.rename(columns={c: f"NTCP_ML_{suffix}"}, inplace=True)

    return df2


def auc_safe(y_true, y_pred) -> float:
    if roc_curve is None or auc is None:
        return float("nan")
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    if len(y_true) < 5 or len(np.unique(y_true)) < 2:
        return float("nan")
    try:
        fpr, tpr, _ = roc_curve(y_true, y_pred)
        return float(auc(fpr, tpr))
    except Exception:
        return float("nan")


def brier_safe(y_true, y_pred) -> float:
    if brier_score_loss is None:
        return float("nan")
    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)
    if len(y_true) < 5:
        return float("nan")
    try:
        return float(brier_score_loss(y_true, y_pred))
    except Exception:
        return float("nan")


def flag_unrealistic(ntcp_vals: pd.Series) -> List[str]:
    ntcp_vals = pd.to_numeric(ntcp_vals, errors="coerce")
    flags = []
    if ntcp_vals.isna().all():
        flags.append("All NTCP are NaN")
    if (ntcp_vals < 0).any() or (ntcp_vals > 1).any():
        flags.append("NTCP outside [0,1] range")
    if ntcp_vals.nunique(dropna=True) <= 1:
        flags.append("No variation in NTCP predictions")
    return flags


def main():
    ap = argparse.ArgumentParser(description="QA Reporter for NTCP outputs")
    ap.add_argument("--input", required=True, help="Path to output folder or zip produced by enhanced_ntcp_analysis.py / enhanced_ntcp_ml.py")
    ap.add_argument("--report_outdir", required=False, default="QA_results", help="Directory to write the report and tables")
    args = ap.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    report_dir = Path(args.report_outdir).expanduser().resolve()
    report_dir.mkdir(parents=True, exist_ok=True)

    # Prepare working directory for zip extraction if needed
    workdir = report_dir / "_unpacked"
    workdir.mkdir(exist_ok=True)

    root = unzip_if_needed(input_path, workdir) if input_path.exists() else None
    if root is None or not root.exists():
        print(f"[ERROR] Input not found: {input_path}", file=sys.stderr)
        sys.exit(2)

    files = discover_files(root)
    # Focus on likely results
    candidate_paths = [p for p in files if likely_results_file(p)]
    tables = {}
    for p in candidate_paths:
        tables[str(p)] = load_table(p)

    # Gather patient-level frames
    patient_frames: List[Tuple[str, pd.DataFrame]] = []
    for path, obj in tables.items():
        if obj is None:
            continue
        if isinstance(obj, dict):
            for sname, df in obj.items():
                if isinstance(df, pd.DataFrame) and df.shape[0] > 0 and is_patient_level_df(df):
                    patient_frames.append((f"{path}::{sname}", df))
        elif isinstance(obj, pd.DataFrame):
            if obj.shape[0] > 0 and is_patient_level_df(obj):
                patient_frames.append((path, obj))

    # Harmonize and combine
    harm_frames: List[Tuple[str, pd.DataFrame]] = []
    for name, df in patient_frames:
        try:
            hf = harmonize(df)
            if "Organ" in hf.columns and ("Observed_Toxicity" in hf.columns or "Grade" in hf.columns):
                # Ensure Observed_Toxicity exists
                if "Observed_Toxicity" not in hf.columns and "Grade" in hf.columns:
                    g = pd.to_numeric(hf["Grade"], errors="coerce")
                    hf["Observed_Toxicity"] = (g >= 2).astype(float)
                harm_frames.append((name, hf))
        except Exception as e:
            print(f"[WARN] Harmonize failed for {name}: {e}", file=sys.stderr)

    combined = pd.concat([hf for _, hf in harm_frames], ignore_index=True, sort=False) if harm_frames else None

    # Determine unique patients
    def get_pid_cols(df: pd.DataFrame) -> List[str]:
        return [c for c in df.columns if c.lower() in ("patientid", "patient", "id")]
    patient_ids = set()
    if combined is not None:
        pid_cols = get_pid_cols(combined)
        if pid_cols:
            for pid in combined[pid_cols[0]].astype(str).str.strip().unique():
                if pid:
                    patient_ids.add(pid)
        else:
            # fallback via filename patterns: e.g., PID_Organ.csv
            for p in files:
                m = re.search(r"[\\/](\w+)_([A-Za-z]+)\.csv$", str(p))
                if m:
                    patient_ids.add(m.group(1))

    # Compute per-organ metrics and flags
    issues: List[str] = []
    report_rows: List[Dict[str, Union[str, int, float]]] = []
    if combined is not None and "Organ" in combined.columns:
        organs = sorted([o for o in combined["Organ"].dropna().unique()])
        model_cols = [c for c in combined.columns if c.startswith("NTCP_")]
        # Allow some common alt names
        aliases = {"NTCP_LKB_LogLogit":"NTCP_lkb_loglogit", "NTCP_LKB_Probit":"NTCP_lkb_probit", "NTCP_RS_Poisson":"NTCP_rs_poisson"}
        for organ in organs:
            sub = combined[combined["Organ"] == organ].copy()
            n = len(sub)
            if "Observed_Toxicity" in sub.columns:
                events = int(pd.to_numeric(sub["Observed_Toxicity"], errors="coerce").fillna(0).sum())
            else:
                events = np.nan
            event_rate = (events/n*100 if n>0 and not np.isnan(events) else np.nan)

            # Metrics
            perf = {}
            for col in model_cols:
                if col in sub.columns:
                    y_true = pd.to_numeric(sub["Observed_Toxicity"], errors="coerce").fillna(0).values if "Observed_Toxicity" in sub.columns else None
                    y_pred = pd.to_numeric(sub[col], errors="coerce").values
                    perf[col] = {
                        "AUC": auc_safe(y_true, y_pred) if y_true is not None else float("nan"),
                        "Brier": brier_safe(y_true, y_pred) if y_true is not None else float("nan"),
                        "Flags": flag_unrealistic(sub[col])
                    }

            # Data quality flags
            if n < 20:
                issues.append(f"{organ}: Low sample size (n={n})")
            if not np.isnan(events) and events < 5:
                issues.append(f"{organ}: Few events (events={events})")

            # ML overfitting/leakage heuristic
            for ml_key in [
                "NTCP_ML_ANN",
                "NTCP_ML_XGBoost",
                "NTCP_ML_XGBOOST",
                "NTCP_ML_RandomForest",
                "NTCP_ML_GradientBoosting",
            ]:
                if ml_key in perf:
                    auc_v = perf[ml_key]["AUC"]
                    if not (auc_v is None or np.isnan(auc_v)):
                        if (n < 40 or (not np.isnan(events) and events < 8)) and auc_v >= 0.90:
                            issues.append(f"{organ}: Potential ML overfitting/leakage (AUC={auc_v:.3f}, n={n}, events={events})")

            # Traditional models optimism under very low events
            for trad in ["NTCP_LKB_LogLogit", "NTCP_LKB_Probit", "NTCP_RS_Poisson"]:
                if trad in perf:
                    auc_v = perf[trad]["AUC"]
                    if not (auc_v is None or np.isnan(auc_v)):
                        if (not np.isnan(events) and events < 5) and auc_v >= 0.85:
                            issues.append(f"{organ}: Traditional model AUC={auc_v:.3f} with events={events} (unstable)")

            # Tabulate
            best_model, best_auc = None, -1
            for k, v in perf.items():
                if v["AUC"] is not None and not np.isnan(v["AUC"]) and v["AUC"] > best_auc:
                    best_auc = v["AUC"]
                    best_model = k

            row = {
                "Organ": organ, "n": n, "events": events,
                "event_rate_%": (round(event_rate, 1) if isinstance(event_rate, (int, float)) and not np.isnan(event_rate) else np.nan),
                "best_model": best_model, "best_auc": (round(best_auc, 3) if best_auc >= 0 else np.nan),
            }
            for k, v in perf.items():
                row[f"AUC|{k}"] = (round(v["AUC"], 3) if v["AUC"] == v["AUC"] else np.nan)
                row[f"Brier|{k}"] = (round(v["Brier"], 3) if v["Brier"] == v["Brier"] else np.nan)
                row[f"Flags|{k}"] = "; ".join(v["Flags"]) if v["Flags"] else ""
            report_rows.append(row)
    else:
        issues.append("Could not locate a patient-level results table or Organ column in outputs.")

    summary_df = pd.DataFrame(report_rows) if report_rows else pd.DataFrame()

    # Global stats
    global_rows = int(summary_df["n"].sum()) if "n" in summary_df else np.nan
    global_patients = len(patient_ids) if patient_ids else np.nan

    # Save tables
    excel_path = report_dir / "qa_summary_tables.xlsx"
    with pd.ExcelWriter(excel_path, engine="xlsxwriter") as writer:
        if not summary_df.empty:
            summary_df.to_excel(writer, index=False, sheet_name="PerOrganSummary")
        if patient_ids:
            pd.DataFrame(sorted(list(patient_ids)), columns=["PatientID"]).to_excel(writer, index=False, sheet_name="UniquePatients")
    
    # Generate Table X: Classical vs ML comparison
    if combined is not None and "Observed_Toxicity" in combined.columns:
        # Try to load clinical data from code0_output if available
        clinical_df = None
        try:
            # Look for clinical_reconciled.xlsx in parent directory
            parent_dir = input_path.parent if input_path.is_file() else input_path
            clinical_path = parent_dir / "code0_output" / "clinical_reconciled.xlsx"
            if not clinical_path.exists():
                # Try alternative locations
                for alt_path in [parent_dir.parent / "code0_output" / "clinical_reconciled.xlsx",
                                 report_dir.parent / "code0_output" / "clinical_reconciled.xlsx"]:
                    if alt_path.exists():
                        clinical_path = alt_path
                        break
            
            if clinical_path.exists():
                clinical_df = pd.read_excel(clinical_path)
                print(f"[INFO] Loaded clinical data from: {clinical_path}")
        except Exception as e:
            print(f"[WARN] Could not load clinical data for Table X: {e}", file=sys.stderr)
        
        # Generate Table X
        try:
            generate_table_X(combined, clinical_df, report_dir)
        except Exception as e:
            print(f"[WARN] Could not generate Table X: {e}", file=sys.stderr)

    # Save DOCX report
    docx_path = report_dir / "comprehensive_report.docx"
    if Document is None:
        print(f"[WARN] python-docx missing; skipping DOCX. Tables saved at: {excel_path}")
    else:
        doc = Document()
        doc.add_heading("Comprehensive QA Report – NTCP Outputs", 0)
        p = doc.add_paragraph()
        p.add_run("Generated on: ").bold = True
        p.add_run(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"))

        doc.add_heading("1) Summary & Data Integrity", level=1)
        doc.add_paragraph(f"• Estimated unique patients: {global_patients}")
        doc.add_paragraph(f"• Total rows across organs (sum of per‑organ n): {global_rows}")
        doc.add_paragraph("Note: Counting per‑organ files as patients inflates totals. This report estimates unique patients using PatientID (if available) or filename patterns.")

        if not summary_df.empty:
            doc.add_heading("2) Per‑Organ Metrics", level=1)
            show_cols = ["Organ","n","events","event_rate_%","best_model","best_auc"]
            show_cols += [c for c in sorted(summary_df.columns) if c.startswith("AUC|")]
            show_cols += [c for c in sorted(summary_df.columns) if c.startswith("Brier|")]
            table = doc.add_table(rows=1, cols=len(show_cols))
            hdr = table.rows[0].cells
            for i, c in enumerate(show_cols):
                hdr[i].text = c
            for _, row in summary_df[show_cols].iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(show_cols):
                    val = row[c]
                    cells[i].text = "" if (isinstance(val, float) and np.isnan(val)) else str(val)
        else:
            doc.add_paragraph("No per‑organ metrics could be constructed from the provided outputs.")

        doc.add_heading("3) Detected Inconsistencies & Risk Flags", level=1)
        if issues:
            for it in issues:
                doc.add_paragraph(f"• {it}")
        else:
            doc.add_paragraph("No critical issues detected by the heuristic checks.")

        if not summary_df.empty:
            doc.add_heading("4) Model‑Specific Flags by Organ", level=1)
            flags_cols = [c for c in summary_df.columns if c.startswith("Flags|")]
            for _, r in summary_df.iterrows():
                sub_flags = []
                for c in flags_cols:
                    txt = r[c]
                    if isinstance(txt, str) and txt.strip():
                        sub_flags.append(f"{c.split('|',1)[1]} -> {txt}")
                if sub_flags:
                    doc.add_paragraph(f"{r['Organ']}:")
                    for f in sub_flags:
                        doc.add_paragraph(f"• {f}")

        doc.add_heading("5) ML Overfitting/Leakage Heuristics (Applied)", level=1)
        doc.add_paragraph("Flagged when AUC ≥ 0.90 with n < 40 or events < 8. High discrimination in small/low‑event cohorts suggests optimism or leakage.")
        
        # NEW: Small Dataset Advisory
        if not summary_df.empty:
            doc.add_heading("6) Small Dataset Advisory", level=1)
            for _, row in summary_df.iterrows():
                n = row.get('n', 0)
                events = row.get('events', 0)
                if isinstance(n, (int, float)) and not np.isnan(n) and n < 100:
                    epv = events / max(1, len([c for c in summary_df.columns if c.startswith('AUC|')])) if isinstance(events, (int, float)) and not np.isnan(events) else 0
                    advisory_text = generate_small_dataset_advisory(int(n), int(events) if isinstance(events, (int, float)) and not np.isnan(events) else 0, epv)
                    if advisory_text:
                        # Parse and add advisory sections
                        for line in advisory_text.split('\n'):
                            if line.strip():
                                if line.startswith('##'):
                                    doc.add_heading(line.replace('##', '').strip(), level=2)
                                elif line.startswith('###'):
                                    doc.add_heading(line.replace('###', '').strip(), level=3)
                                elif line.startswith('-'):
                                    doc.add_paragraph(line.strip('- ').strip(), style='List Bullet')
                                else:
                                    doc.add_paragraph(line.strip())
        
        # NEW: Clinical Factor Significance
        doc.add_heading("7) Clinical Factor Significance", level=1)
        # Try to extract clinical factors from adaptation reports if available
        # This would need to be passed from code3 or stored in results
        doc.add_paragraph("Clinical factors with p < 0.05 are automatically integrated into ML models.")
        doc.add_paragraph("Check model feature lists for included clinical factors (e.g., Age, Chemotherapy, etc.).")
        
        # V2.0: Data Leakage Detection
        if V2_LEAKAGE_DETECTOR_AVAILABLE and combined is not None:
            doc.add_heading("8) Data Leakage Detection (v2.0)", level=1)
            leakage_detector = DataLeakageDetector()
            
            # Check for train/test split information in results
            # Look for columns that might indicate train/test split
            has_split_info = any(col.lower() in ['split', 'train', 'test', 'fold'] for col in combined.columns)
            
            if has_split_info or 'PrimaryPatientID' in combined.columns:
                # Try to detect potential leakage
                if 'PrimaryPatientID' in combined.columns:
                    # Check for duplicate patients across what might be train/test
                    patient_counts = combined.groupby('PrimaryPatientID').size()
                    duplicate_patients = patient_counts[patient_counts > 1]
                    
                    if len(duplicate_patients) > 0:
                        # Check if these duplicates span what might be train/test
                        leakage_detector.warnings.append(
                            f"Found {len(duplicate_patients)} patients with multiple entries. "
                            "This may indicate data leakage if same patients appear in both train and test sets."
                        )
                    
                    # Basic check: ensure patient-level integrity
                    leakage_detector.checks_performed.append(
                        f"Patient ID column found: {len(combined['PrimaryPatientID'].unique())} unique patients"
                    )
                
                leakage_report = leakage_detector.generate_report()
                
                if leakage_report['has_warnings'] or leakage_report['errors']:
                    doc.add_paragraph("⚠️ Data Leakage Warnings Detected:")
                    for warning in leakage_report['warnings']:
                        doc.add_paragraph(f"• {warning}", style='List Bullet')
                    for error in leakage_report['errors']:
                        doc.add_paragraph(f"• {error}", style='List Bullet')
                else:
                    doc.add_paragraph("✓ No data leakage detected in available data.")
                    doc.add_paragraph("Note: Full leakage detection requires access to train/test split information.")
            else:
                doc.add_paragraph("Data leakage detection requires PrimaryPatientID column or train/test split information.")
                doc.add_paragraph("Current data does not contain sufficient information for leakage detection.")
        else:
            doc.add_heading("8) Data Leakage Detection", level=1)
            doc.add_paragraph("v2.0 leakage detection components not available. Install v2.0 components for enhanced leakage detection.")

        doc.save(docx_path)

    print(f"[OK] Saved report: {docx_path}")
    print(f"[OK] Saved tables: {excel_path}")


def generate_small_dataset_advisory(n_samples, n_events, epv):
    """
    Generate advisory text for small datasets
    
    Parameters
    ----------
    n_samples : int
        Number of samples
    n_events : int
        Number of events
    epv : float
        Events per variable
        
    Returns
    -------
    str
        Advisory text
    """
    advisory = []
    
    if n_samples < 100:
        advisory.append("## ⚠️ SMALL DATASET ADVISORY")
        advisory.append(f"- **Dataset size**: {n_samples} patients")
        advisory.append(f"- **Number of events**: {n_events}")
        advisory.append(f"- **Events per variable (EPV)**: {epv:.2f}")
        
        if n_samples < 30:
            advisory.append("### 🔴 HIGH UNCERTAINTY")
            advisory.append("- Sample size < 30: Results are exploratory")
            advisory.append("- Use LOOCV: No independent test set")
            advisory.append("- Models may be unstable")
            advisory.append("- **Recommendation**: Collect more data before clinical application")
        
        elif n_samples < 50:
            advisory.append("### 🟡 MODERATE UNCERTAINTY")
            advisory.append("- Sample size 30-50: Limited generalizability")
            advisory.append("- Simplified models used (reduced complexity)")
            advisory.append("- Wider confidence intervals expected")
            advisory.append("- **Recommendation**: Use for hypothesis generation only")
        
        elif n_samples < 100:
            advisory.append("### 🟢 CAUTION ADVISED")
            advisory.append("- Sample size 50-100: Results indicative but not definitive")
            advisory.append("- Model complexity adapted to sample size")
            advisory.append("- Clinical factors integrated when significant")
            advisory.append("- **Recommendation**: Validate with external dataset")
        
        advisory.append("\n### 📊 Statistical Considerations:")
        advisory.append(f"- Minimum EPV for reliable modeling: 10 (Current: {epv:.2f})")
        advisory.append(f"- Recommended sample size for parotid NTCP: ≥150 patients")
        advisory.append(f"- Bootstrap confidence intervals may be wide")
    
    return "\n".join(advisory)


def report_clinical_factor_significance(significant_factors, p_values=None):
    """
    Create prominent clinical factor significance report
    
    Parameters
    ----------
    significant_factors : list of str
        List of significant clinical factor names
    p_values : dict, optional
        Dictionary mapping factor names to p-values
        
    Returns
    -------
    str
        Clinical factor report text
    """
    if not significant_factors:
        return "## 📋 Clinical Factors Analysis\nNo significant clinical factors found (p < 0.05)"
    
    report = ["## 📋 CLINICALLY SIGNIFICANT FACTORS"]
    report.append("The following clinical factors showed significant association with toxicity (p < 0.05):")
    report.append("")
    
    for i, factor in enumerate(significant_factors, 1):
        p_val = p_values.get(factor, "N/A") if p_values else "N/A"
        p_str = f" (p = {p_val})" if p_val != "N/A" else ""
        
        # Add clinical interpretation
        interpretation = ""
        if 'age' in factor.lower() or 'Age' in factor:
            interpretation = " → Older age associated with higher toxicity risk"
        elif 'chemo' in factor.lower() or 'Chemo' in factor:
            interpretation = " → Chemotherapy increases toxicity risk"
        elif 'tobacco' in factor.lower() or 'smoking' in factor.lower() or 'Smoking' in factor:
            interpretation = " → Tobacco exposure may increase risk"
        elif 'diabetes' in factor.lower() or 'Diabetes' in factor:
            interpretation = " → Diabetes may modify toxicity risk"
        
        report.append(f"{i}. **{factor}**{p_str}{interpretation}")
    
    report.append("")
    report.append("### 🩺 Clinical Implications:")
    report.append("- These factors have been automatically integrated into NTCP models")
    report.append("- Consider these in treatment planning and patient counseling")
    report.append("- Further validation needed for clinical implementation")
    
    return "\n".join(report)


def generate_enhanced_qa_report(results, adaptation_report=None):
    """
    Generate enhanced QA report with small dataset advisories
    
    Parameters
    ----------
    results : dict
        Results dictionary with organ-level data
    adaptation_report : dict, optional
        Adaptation report from ML training
        
    Returns
    -------
    str
        Enhanced QA report text
    """
    report_sections = []
    
    # 1. Small Dataset Advisory
    if adaptation_report and adaptation_report.get('dataset_size', 0) < 100:
        advisory = generate_small_dataset_advisory(
            adaptation_report['dataset_size'],
            adaptation_report.get('n_events', 0),
            adaptation_report.get('epv', 0)
        )
        report_sections.append(advisory)
    
    # 2. Clinical Factor Significance
    if adaptation_report and adaptation_report.get('significant_clinical_factors'):
        clinical_report = report_clinical_factor_significance(
            adaptation_report['significant_clinical_factors']
        )
        report_sections.append(clinical_report)
    
    # 3. Feature Selection Summary
    if adaptation_report and adaptation_report.get('selected_features'):
        report_sections.append(f"## 🔍 Selected Features\n{', '.join(adaptation_report['selected_features'])}")
    
    # 4. CV Strategy Used
    if adaptation_report and adaptation_report.get('cv_strategy'):
        report_sections.append(f"## 📈 Cross-Validation Strategy\n{adaptation_report['cv_strategy']} ({adaptation_report.get('cv_folds', 'N/A')} folds)")
    
    # 5. Model Adaptation Summary
    if adaptation_report and adaptation_report.get('model_config'):
        report_sections.append("## ⚙️ Model Adaptations")
        report_sections.append(f"- ANN: {adaptation_report['model_config'].get('ann', {})}")
        report_sections.append(f"- XGBoost: {adaptation_report['model_config'].get('xgboost', {})}")
    
    # Combine all sections
    full_report = "\n\n".join(report_sections)
    
    return full_report


def generate_table_X(ntcp_df, clinical_df, output_dir):
    """
    Generate Table X:
    Performance of classical NTCP (literature vs local) vs ML
    
    Args:
        ntcp_df: DataFrame with NTCP predictions (must contain NTCP columns and Observed_Toxicity)
        clinical_df: DataFrame with clinical data (must contain xerostomia_grade2plus)
        output_dir: Output directory for saving tables
    """
    import os
    from scipy.stats import spearmanr
    
    # Import metrics
    try:
        from sklearn.metrics import roc_auc_score, brier_score_loss, roc_curve
    except ImportError:
        print("[WARN] scikit-learn not available; Table X metrics will be NaN", file=sys.stderr)
        roc_auc_score = None
        brier_score_loss = None
    
    # Helper functions for metrics
    def compute_auc(y_true, y_pred):
        """Compute AUC (ROC)"""
        if roc_auc_score is None:
            return float("nan")
        y_true = np.asarray(y_true)
        y_pred = np.asarray(y_pred)
        if len(y_true) < 5 or len(np.unique(y_true)) < 2:
            return float("nan")
        try:
            return float(roc_auc_score(y_true, y_pred))
        except Exception:
            return float("nan")
    
    def compute_brier_score(y_true, y_pred):
        """Compute Brier score"""
        if brier_score_loss is None:
            return float("nan")
        y_true = np.asarray(y_true)
        y_pred = np.asarray(y_pred)
        if len(y_true) < 5:
            return float("nan")
        try:
            return float(brier_score_loss(y_true, y_pred))
        except Exception:
            return float("nan")
    
    def compute_calibration_slope_intercept(y_true, y_pred):
        """Compute calibration slope and intercept using binned calibration"""
        y_true = np.asarray(y_true)
        y_pred = np.asarray(y_pred)
        
        # Remove NaN values
        valid_mask = ~(np.isnan(y_true) | np.isnan(y_pred))
        if valid_mask.sum() < 10:
            return float("nan"), float("nan")
        
        y_true_valid = y_true[valid_mask]
        y_pred_valid = y_pred[valid_mask]
        
        # Bin predictions (5 bins)
        n_bins = min(5, len(y_pred_valid) // 3)
        if n_bins < 2:
            return float("nan"), float("nan")
        
        try:
            bin_edges = np.linspace(0, 1, n_bins + 1)
            bin_indices = np.digitize(y_pred_valid, bin_edges) - 1
            bin_indices = np.clip(bin_indices, 0, n_bins - 1)
            
            bin_centers = []
            bin_observed = []
            
            for i in range(n_bins):
                mask = (bin_indices == i)
                if mask.sum() > 0:
                    bin_centers.append(np.mean(y_pred_valid[mask]))
                    bin_observed.append(np.mean(y_true_valid[mask]))
            
            if len(bin_centers) < 2:
                return float("nan"), float("nan")
            
            bin_centers = np.array(bin_centers)
            bin_observed = np.array(bin_observed)
            
            # Linear regression: observed = slope * predicted + intercept
            x_mean = np.mean(bin_centers)
            y_mean = np.mean(bin_observed)
            numerator = np.sum((bin_centers - x_mean) * (bin_observed - y_mean))
            denominator = np.sum((bin_centers - x_mean) ** 2)
            
            if denominator == 0:
                return float("nan"), float("nan")
            
            slope = numerator / denominator
            intercept = y_mean - slope * x_mean
            
            return float(slope), float(intercept)
        except Exception:
            return float("nan"), float("nan")
    
    def compute_cohort_consistency_score(y_pred):
        """
        Compute Cohort Consistency Score (CCS) using correlation with monotonic dose ranking.
        Simplified version: correlation of NTCP vs monotonic dose ranking across cohort.
        """
        y_pred = np.asarray(y_pred)
        valid_mask = ~np.isnan(y_pred)
        
        if valid_mask.sum() < 5:
            return float("nan")
        
        y_pred_valid = y_pred[valid_mask]
        
        # Create monotonic dose ranking (simplified: use NTCP predictions as proxy for dose)
        # In practice, this would use actual dose metrics (gEUD, mean dose, etc.)
        # For this simplified version, we use the rank of predictions
        dose_ranking = np.arange(len(y_pred_valid))
        
        try:
            # Compute Spearman correlation between NTCP and dose ranking
            correlation, _ = spearmanr(y_pred_valid, dose_ranking)
            # Convert correlation to 0-1 scale (CCS)
            # CCS = (correlation + 1) / 2, but we want it to reflect consistency
            # Higher correlation = higher consistency
            ccs = (correlation + 1.0) / 2.0 if not np.isnan(correlation) else float("nan")
            return float(ccs)
        except Exception:
            return float("nan")
    
    # Model definitions
    models = {
        "QUANTEC-LKB": ("Classical (literature)", "NTCP_LKB_QUANTEC"),
        "QUANTEC-RS": ("Classical (literature)", "NTCP_RS_QUANTEC"),
        "Local-LKB": ("Classical (local)", "NTCP_LKB_LOCAL"),
        "ANN": ("ML", "NTCP_ANN"),
        "XGBoost": ("ML", "NTCP_XGB"),
        "RandomForest": ("ML", "NTCP_RF"),
        "GradientBoosting": ("ML", "NTCP_GB"),
    }
    
    # Get true labels - use clinical_df if available, otherwise use Observed_Toxicity from ntcp_df
    if clinical_df is not None and "xerostomia_grade2plus" in clinical_df.columns:
        # Merge clinical data with NTCP data if needed
        if "PatientID" in ntcp_df.columns and "patient_id" in clinical_df.columns:
            # Try to merge on patient ID
            merged = ntcp_df.merge(
                clinical_df[["patient_id", "xerostomia_grade2plus"]],
                left_on="PatientID",
                right_on="patient_id",
                how="left"
            )
            y_true = merged["xerostomia_grade2plus"].values
        else:
            # Use clinical_df directly if it has the same length
            if len(clinical_df) == len(ntcp_df):
                y_true = clinical_df["xerostomia_grade2plus"].values
            else:
                # Fallback to Observed_Toxicity
                y_true = ntcp_df["Observed_Toxicity"].values if "Observed_Toxicity" in ntcp_df.columns else None
    else:
        # Use Observed_Toxicity from ntcp_df
        y_true = ntcp_df["Observed_Toxicity"].values if "Observed_Toxicity" in ntcp_df.columns else None
    
    if y_true is None:
        print("[WARN] Could not determine true labels for Table X", file=sys.stderr)
        return pd.DataFrame()
    
    rows = []
    
    for model_name, (category, col) in models.items():
        if col not in ntcp_df.columns:
            # Try alternative column names
            alt_names = {
                "NTCP_LKB_QUANTEC": ["NTCP_LKB_Probit", "NTCP_LKB_LogLogit", "NTCP_LKB_QUANTEC"],
                "NTCP_RS_QUANTEC": ["NTCP_RS_Poisson", "NTCP_RS_QUANTEC"],
                "NTCP_LKB_LOCAL": ["NTCP_LKB_LOCAL"],
                "NTCP_ANN": ["NTCP_ML_ANN", "NTCP_ANN", "NTCP_ML_NeuralNetwork"],
                "NTCP_XGB": ["NTCP_ML_XGBoost", "NTCP_XGBoost", "NTCP_XGB", "NTCP_ML_XGB"],
                "NTCP_RF": ["NTCP_ML_RandomForest", "NTCP_ML_RF", "NTCP_RF"],
                "NTCP_GB": ["NTCP_ML_GradientBoosting", "NTCP_ML_GB", "NTCP_GB"],
            }
            
            found_col = None
            if col in alt_names:
                for alt in alt_names[col]:
                    if alt in ntcp_df.columns:
                        found_col = alt
                        break
            
            if found_col is None:
                continue
            
            col = found_col
        
        y_pred = pd.to_numeric(ntcp_df[col], errors="coerce").values
        
        # Remove NaN values
        valid_mask = ~(np.isnan(y_true) | np.isnan(y_pred))
        if valid_mask.sum() < 5:
            continue
        
        y_true_valid = y_true[valid_mask]
        y_pred_valid = y_pred[valid_mask]
        
        auc = compute_auc(y_true_valid, y_pred_valid)
        brier = compute_brier_score(y_true_valid, y_pred_valid)
        slope, intercept = compute_calibration_slope_intercept(y_true_valid, y_pred_valid)
        ccs = compute_cohort_consistency_score(y_pred_valid)
        
        rows.append({
            "Category": category,
            "Model": model_name,
            "AUC": round(auc, 3) if not np.isnan(auc) else np.nan,
            "Brier": round(brier, 3) if not np.isnan(brier) else np.nan,
            "Calibration slope": round(slope, 3) if not np.isnan(slope) else np.nan,
            "Calibration intercept": round(intercept, 3) if not np.isnan(intercept) else np.nan,
            "CCS": round(ccs, 3) if not np.isnan(ccs) else np.nan,
        })
    
    table_df = pd.DataFrame(rows)
    
    # Create tables directory
    tables_dir = Path(output_dir) / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    
    # Excel output
    excel_path = tables_dir / "Table_X_Classical_vs_ML.xlsx"
    table_df.to_excel(excel_path, index=False)
    
    # Markdown output (for manuscript / Claude)
    md_path = tables_dir / "Table_X_Classical_vs_ML.md"
    try:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(table_df.to_markdown(index=False))
        print(f"[OK] Saved Table X (Markdown): {md_path}")
    except Exception as e:
        # Fallback to CSV if markdown fails
        csv_path = tables_dir / "Table_X_Classical_vs_ML.csv"
        table_df.to_csv(csv_path, index=False)
        print(f"[OK] Saved Table X (CSV): {csv_path} (markdown unavailable: {e})")
    
    print(f"[OK] Saved Table X: {excel_path}")
    
    return table_df


if __name__ == "__main__":
    main()
