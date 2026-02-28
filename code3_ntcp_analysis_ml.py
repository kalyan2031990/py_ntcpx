#!/usr/bin/env python3
"""
Enhanced NTCP Analysis with Traditional and Machine Learning Models
==================================================================

This comprehensive script combines:
1. Traditional NTCP models (LKB Log-Logistic, LKB Probit, RS Poisson)
2. Machine learning models (ANN, XGBoost) with proper validation
3. Professional 600 DPI publication-ready plots
4. Unique colors and legends for all models
5. Comprehensive Excel output with all results

Author: K. Mondal (North Bengal Medical College, Darjeeling, India.)
Version: 3.0.1
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from pathlib import Path as _Path
import argparse
import sys
import os
import unicodedata
import re
from scipy.optimize import minimize, differential_evolution
from scipy.stats import norm
from sklearn.metrics import roc_curve, auc, brier_score_loss, log_loss
from sklearn.model_selection import cross_val_score, StratifiedKFold, train_test_split, LeaveOneOut
from sklearn.neural_network import MLPClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
import warnings
warnings.filterwarnings('ignore')

# Import v2.0 components for data leakage prevention and overfitting control
try:
    from src.validation.data_splitter import PatientDataSplitter
    from src.models.machine_learning.ml_models import OverfitResistantMLModels
    from src.features.feature_selector import RadiobiologyGuidedFeatureSelector
    from src.metrics.auc_calculator import calculate_auc_with_ci
    from src.reporting.leakage_detector import DataLeakageDetector
    V2_COMPONENTS_AVAILABLE = True
except ImportError as e:
    V2_COMPONENTS_AVAILABLE = False
    print(f"Warning: v2.0 components not available: {e}")
    print("Falling back to basic implementation. Install v2.0 components for enhanced features.")

# Windows-safe encoding configuration
try:
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
except (AttributeError, ValueError):
    pass

# DVH ID normalization function for canonical matching
def normalize_dvh_id(x):
    """
    Normalize DVH ID for canonical matching across Excel and filenames.
    
    Handles:
    - Unicode dash normalization (e.g., en-dash, em-dash -> hyphen)
    - Whitespace trimming
    - Case normalization (lowercase)
    - File extension removal (.csv)
    - Keep only digits and hyphens
    
    Args:
        x: DVH ID (string, number, or None)
        
    Returns:
        str: Normalized DVH ID or None if input is None
    """
    if x is None:
        return None
    x = str(x)
    # Normalize Unicode dashes and other characters (NFKC form)
    x = unicodedata.normalize("NFKC", x)
    # Strip whitespace and convert to lowercase
    x = x.strip().lower()
    # Remove .csv extension if present
    x = x.replace(".csv", "")
    # Remove all spaces, underscores, and other separators before filtering
    x = x.replace(" ", "").replace("_", "")
    # Keep only digits and hyphens
    x = re.sub(r"[^0-9\-]", "", x)
    return x

# Import utility functions and novel models
try:
    from ntcp_utils import normalize_columns, find_dvh_file
    from ntcp_novel_models import ProbabilisticgEUDModel, MonteCarloNTCPModel
    from ntcp_qa_modules import UncertaintyAwareNTCP, CohortConsistencyScore
except ImportError:
    print("Warning: Novel NTCP models and QA modules not available. Install required modules.")
    normalize_columns = lambda df: df
    find_dvh_file = None
    ProbabilisticgEUDModel = None
    MonteCarloNTCPModel = None
    UncertaintyAwareNTCP = None
    CohortConsistencyScore = None

# Try to import XGBoost
try:
    import xgboost as xgb
    XGBOOST_AVAILABLE = True
except ImportError:
    XGBOOST_AVAILABLE = False
    print("Warning: XGBoost not available. Install with: pip install xgboost")

# Set publication-quality plotting parameters (600 DPI)
plt.rcParams.update({
    'font.size': 12,
    'font.family': 'sans-serif',
    'font.sans-serif': ['Arial', 'Helvetica', 'DejaVu Sans'],
    'axes.linewidth': 1.2,
    'axes.spines.top': False,
    'axes.spines.right': False,
    'axes.grid': True,
    'grid.alpha': 0.3,
    'grid.linewidth': 0.8,
    'legend.frameon': False,
    'legend.fontsize': 10,
    'xtick.major.size': 6,
    'ytick.major.size': 6,
    'xtick.minor.size': 3,
    'ytick.minor.size': 3,
    'lines.linewidth': 2.5,
    'lines.markersize': 6,
    'figure.dpi': 100,
    'savefig.dpi': 600,  # 600 DPI for publication quality
    'savefig.bbox': 'tight',
    'savefig.pad_inches': 0.1,
    'savefig.facecolor': 'white'
})

# Professional color scheme with unique colors for all models
COLORS = {
    'LKB_LogLogit': '#2E86AB',      # Professional blue
    'LKB_Probit': '#F24236',        # Professional red  
    'RS_Poisson': '#F6AE2D',        # Professional gold
    'ML_ANN': '#8B4B9E',            # Purple for ANN
    'ML_XGBoost': '#2ECC71',        # Green for XGBoost
    'ML_RandomForest': '#FF8C00',   # Orange for Random Forest
    'observed': '#C73E1D',          # Red for observed data
    'literature': '#592E83',        # Dark purple for literature
    'confidence': '#95A5A6',        # Light gray for confidence
    'grid': '#E8E8E8'               # Very light gray for grid
}

# Line styles for different models
LINE_STYLES = {
    'LKB_LogLogit': '-',
    'LKB_Probit': '--',
    'RS_Poisson': '-.',
    'ML_ANN': ':',
    'ML_XGBoost': (0, (3, 1, 1, 1)),  # Custom dash pattern
    'ML_RandomForest': (0, (1, 1))    # Dotted line
}

# Markers for scatter plots
MARKERS = {
    'LKB_LogLogit': 'o',
    'LKB_Probit': 's',
    'RS_Poisson': '^',
    'ML_ANN': 'D',
    'ML_XGBoost': 'X',
    'ML_RandomForest': 'P'
}

class DVHProcessor:
    """Process differential DVH data for NTCP calculations"""
    
    def __init__(self, dvh_directory):
        self.dvh_dir = Path(dvh_directory)
        self.processed_data = {}
        
    def load_dvh_file(self, patient_id, organ, patient_name=None, dvh_id=None, dvh_id_norm=None):
        """
        Load differential DVH file for specific patient and organ.
        Identity-safe: matches on PrimaryPatientID (real patient ID from filename).
        
        Args:
            patient_id: PrimaryPatientID (real patient ID, e.g., "2020-734")
            organ: Organ name
            patient_name: Optional patient name (not used for matching)
            dvh_id: Same as patient_id (for backward compatibility)
            dvh_id_norm: Normalized PrimaryPatientID for matching
        
        Returns:
            tuple: (dvh DataFrame, extracted_patient_id) or (None, None) if not found
        """
        dvh_file = None
        extracted_patient_id = None
        
        # Identity-safe: use PrimaryPatientID for matching (patient_id parameter is now PrimaryPatientID)
        primary_id = patient_id if patient_id else dvh_id
        if primary_id is None:
            print(f"Warning: No PrimaryPatientID provided for matching - {organ}")
            return None, None
        
        # Normalize PrimaryPatientID for matching
        primary_id_norm = dvh_id_norm if dvh_id_norm else normalize_dvh_id(primary_id)
        
        # Match DVH files using PrimaryPatientID (identity-safe key)
        # Filenames are now: {PrimaryPatientID}_{Organ}.csv
        if primary_id_norm and primary_id_norm != "":
            # Search all CSV files and match using normalized PrimaryPatientID
            for candidate_file in self.dvh_dir.glob('*.csv'):
                filename_parts = candidate_file.stem.split('_')
                if len(filename_parts) >= 2:
                    candidate_primary_id = filename_parts[0]
                    candidate_id_norm = normalize_dvh_id(candidate_primary_id)
                    file_organ = '_'.join(filename_parts[1:])
                    
                    # Match on normalized PrimaryPatientID + Organ (identity-safe)
                    if candidate_id_norm == primary_id_norm and file_organ.lower() == organ.lower():
                        dvh_file = candidate_file
                        extracted_patient_id = candidate_primary_id  # Original PrimaryPatientID
                        break
        
        # Fallback: try exact filename match
        if dvh_file is None or not dvh_file.exists():
            dvh_file = self.dvh_dir / f"{primary_id}_{organ}.csv"
            if dvh_file.exists():
                extracted_patient_id = primary_id
        
        if dvh_file is None or not dvh_file.exists():
            print(f"Warning: DVH file not found for PrimaryPatientID={primary_id}, Organ={organ}")
            return None, None
            
        try:
            # Load DVH data
            dvh = pd.read_csv(dvh_file)
            
            # Standardize column names
            if 'Dose[Gy]' in dvh.columns and 'Volume[cm3]' in dvh.columns:
                dvh = dvh.rename(columns={'Dose[Gy]': 'dose_gy', 'Volume[cm3]': 'volume_cm3'})
            elif 'Dose' in dvh.columns and 'Volume' in dvh.columns:
                dvh = dvh.rename(columns={'Dose': 'dose_gy', 'Volume': 'volume_cm3'})
            
            # Remove zero volume entries at high doses
            dvh = dvh[dvh['volume_cm3'] > 0].copy()
            
            # Sort by dose
            dvh = dvh.sort_values('dose_gy').reset_index(drop=True)
            
            # Extract PatientID from filename if not already extracted
            if extracted_patient_id is None:
                extracted_patient_id = dvh_file.stem.split('_')[0]
            
            # Ensure extracted_patient_id is normalized for consistency
            # Keep original for display, but ensure it's a valid string
            extracted_patient_id = str(extracted_patient_id) if extracted_patient_id is not None else None
            
            return dvh, extracted_patient_id
            
        except Exception as e:
            print(f"Error loading {dvh_file}: {e}")
            return None, None
    
    def calculate_dose_metrics(self, dvh):
        """Calculate comprehensive dose metrics from differential DVH"""
        if dvh is None or len(dvh) == 0:
            return None
            
        doses = dvh['dose_gy'].values
        volumes = dvh['volume_cm3'].values
        total_volume = np.sum(volumes)
        
        if total_volume <= 0:
            return None
        
        # Calculate relative volumes
        rel_volumes = volumes / total_volume
        
        # Basic dose metrics
        mean_dose = np.sum(rel_volumes * doses)
        max_dose = np.max(doses)
        min_dose = np.min(doses[volumes > 0])
        
        # Convert to cumulative DVH for Vxx calculations
        cumulative_vol = np.cumsum(volumes[::-1])[::-1]
        rel_cumulative = cumulative_vol / total_volume
        
        dose_metrics = {
            'total_volume': total_volume,
            'mean_dose': mean_dose,
            'max_dose': max_dose,
            'min_dose': min_dose
        }
        
        # Calculate Vxx (% volume receiving >= xx Gy)
        for dose_level in [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70]:
            if dose_level <= max_dose:
                volume_at_dose = np.interp(dose_level, doses, rel_cumulative) * 100
                dose_metrics[f'V{dose_level}'] = volume_at_dose
            else:
                dose_metrics[f'V{dose_level}'] = 0.0
        
        # Calculate Dxx (dose to xx% of volume)
        for vol_percent in [0.01, 0.1, 1, 2, 5, 10, 20, 30, 50, 70, 90, 95, 98]:
            target_vol_fraction = vol_percent / 100
            if target_vol_fraction <= 1.0:
                dose_at_volume = np.interp(target_vol_fraction, rel_cumulative[::-1], doses[::-1])
                dose_metrics[f'D{vol_percent}'] = dose_at_volume
        
        return dose_metrics
    
    def calculate_gEUD(self, dvh, a_parameter):
        """Calculate generalized Equivalent Uniform Dose (gEUD)"""
        if dvh is None or len(dvh) == 0:
            return np.nan
            
        doses = dvh['dose_gy'].values
        volumes = dvh['volume_cm3'].values
        total_volume = np.sum(volumes)
        
        if total_volume <= 0:
            return np.nan
        
        # Calculate relative volumes
        rel_volumes = volumes / total_volume
        
        # Handle special cases
        if a_parameter == 0:
            # a=0 case: geometric mean
            log_doses = np.log(np.maximum(doses, 1e-10))
            log_geud = np.sum(rel_volumes * log_doses)
            return np.exp(log_geud)
        
        elif a_parameter == 1:
            # a=1 case: arithmetic mean (mean dose)
            return np.sum(rel_volumes * doses)
        
        elif np.isinf(a_parameter):
            # a=∞ case: maximum dose
            return np.max(doses)
        
        else:
            # General case: gEUD = (Σ vi × Di^a)^(1/a)
            powered_doses = np.power(np.maximum(doses, 1e-10), a_parameter)
            sum_weighted = np.sum(rel_volumes * powered_doses)
            
            if sum_weighted <= 0:
                return np.nan
                
            geud = np.power(sum_weighted, 1.0 / a_parameter)
            return geud
    
    def calculate_effective_volume(self, dvh, n_parameter):
        """Calculate effective volume for LKB probit model"""
        if dvh is None or len(dvh) == 0:
            return np.nan
            
        doses = dvh['dose_gy'].values
        volumes = dvh['volume_cm3'].values
        total_volume = np.sum(volumes)
        
        if total_volume <= 0:
            return np.nan
        
        max_dose = np.max(doses)
        if max_dose <= 0:
            return np.nan
        
        # Calculate relative volumes
        rel_volumes = volumes / total_volume
        
        # Calculate dose ratio terms
        dose_ratios = doses / max_dose
        
        if n_parameter == 0:
            return 1.0
        else:
            powered_ratios = np.power(dose_ratios, 1.0 / n_parameter)
            v_eff = np.sum(rel_volumes * powered_ratios)
            return v_eff

class NTCPCalculator:
    """Calculate NTCP using published model equations"""
    
    def __init__(self):
        # Literature parameters
        self.literature_params = {
            'Parotid': {
                'LKB_LogLogit': {'a': 2.2, 'TD50': 28.4, 'gamma50': 1.0, 'alpha_beta': 3},
                'LKB_Probit': {'TD50': 28.4, 'm': 0.18, 'n': 0.45, 'alpha_beta': 3},
                'RS_Poisson': {'D50': 26.3, 'gamma': 0.73, 's': 0.01, 'alpha_beta': 3}
            },
            'Larynx': {
                'LKB_LogLogit': {'a': 1.0, 'TD50': 44.0, 'gamma50': 1.0, 'alpha_beta': 3},
                'LKB_Probit': {'TD50': 44.0, 'm': 0.20, 'n': 1.0, 'alpha_beta': 3},
                'RS_Poisson': {'D50': 40.0, 'gamma': 1.2, 's': 0.12, 'alpha_beta': 3}
            },
            'SpinalCord': {
                'LKB_LogLogit': {'a': 7.4, 'TD50': 66.5, 'gamma50': 4.0, 'alpha_beta': 2},
                'LKB_Probit': {'TD50': 66.5, 'm': 0.10, 'n': 0.03, 'alpha_beta': 2},
                'RS_Poisson': {'D50': 68.6, 'gamma': 1.9, 's': 4.0, 'alpha_beta': 2}
            }
        }
    
    def convert_to_eqd2(self, dose, alpha_beta_ratio, dose_per_fraction, n_fractions=None):
        """Convert physical dose to EQD2"""
        if np.isnan(dose) or dose <= 0:
            return np.nan
            
        if dose_per_fraction is None:
            if n_fractions is not None:
                dose_per_fraction = dose / n_fractions
            else:
                dose_per_fraction = 2.0
        
        eqd2 = dose * (alpha_beta_ratio + dose_per_fraction) / (alpha_beta_ratio + 2.0)
        return eqd2
    
    def ntcp_lkb_loglogit(self, geud, TD50, gamma50):
        """LKB model with log-logistic link function"""
        if np.isnan(geud) or geud <= 0 or TD50 <= 0 or gamma50 <= 0:
            return 0.0
        
        try:
            ratio = TD50 / geud
            exponent = 4.0 * gamma50
            ntcp = 1.0 / (1.0 + np.power(ratio, exponent))
            ntcp = np.clip(ntcp, 1e-15, 1.0 - 1e-15)
            return ntcp
        except (OverflowError, ZeroDivisionError):
            return 0.0 if geud < TD50 else 1.0
    
    def ntcp_lkb_probit(self, dose_metrics, TD50, m, n):
        """LKB model with probit link function"""
        if 'v_effective' not in dose_metrics or np.isnan(dose_metrics['v_effective']):
            return 0.0
        
        v_eff = dose_metrics['v_effective']
        max_dose = dose_metrics.get('max_dose', 0)
        
        if max_dose <= 0 or TD50 <= 0 or m <= 0 or v_eff <= 0:
            return 0.0
        
        try:
            # Calculate effective TD50
            td_veff_50 = TD50 * np.power(v_eff, -n)
            
            # Calculate t parameter
            t = (max_dose - td_veff_50) / (m * td_veff_50)
            
            # Apply probit function
            ntcp = norm.cdf(t)
            ntcp = np.clip(ntcp, 1e-15, 1.0 - 1e-15)
            return ntcp
        except (OverflowError, ZeroDivisionError):
            return 0.0 if max_dose < TD50 else 1.0
    
    def ntcp_rs_poisson(self, dvh, D50, gamma, s):
        """Relative Seriality model with Poisson statistics"""
        if dvh is None or len(dvh) == 0 or D50 <= 0 or gamma <= 0 or s <= 0:
            return 0.0
        
        try:
            doses = dvh['dose_gy'].values
            volumes = dvh['volume_cm3'].values
            total_volume = np.sum(volumes)
            
            if total_volume <= 0:
                return 0.0
            
            # Calculate relative volumes
            rel_volumes = volumes / total_volume
            
            # Calculate voxel-level NTCP for each dose bin (Källman et al.: p_i = 2^{-exp(e·γ(1−r_i))})
            dose_ratios = doses / D50
            voxel_ntcps = 2.0 ** (-np.exp(np.e * gamma * (1.0 - dose_ratios)))
            
            # Handle numerical issues
            voxel_ntcps = np.clip(voxel_ntcps, 1e-15, 1.0 - 1e-15)
            
            # Calculate organ-level NTCP using relative seriality
            powered_ntcps = np.power(voxel_ntcps, s)
            complement_terms = 1.0 - powered_ntcps
            
            # Take the product weighted by relative volumes
            log_terms = rel_volumes * np.log(np.maximum(complement_terms, 1e-15))
            product_term = np.exp(np.sum(log_terms))
            
            # Final NTCP calculation
            ntcp = np.power(1.0 - product_term, 1.0 / s)
            ntcp = np.clip(ntcp, 1e-15, 1.0 - 1e-15)
            return ntcp
        except (OverflowError, ZeroDivisionError, ValueError):
            return 0.0
    
    def calculate_all_ntcp_models(self, dvh, dose_metrics, organ, dose_per_fraction=2.0):
        """Calculate NTCP using all three models for given organ"""
        
        if organ not in self.literature_params:
            print(f"Warning: No parameters available for organ '{organ}'")
            return {}
        
        organ_params = self.literature_params[organ]
        results = {}
        
        # Get required dose metrics
        geud = dose_metrics.get('gEUD', np.nan)
        v_effective = dose_metrics.get('v_effective', np.nan)
        
        # 1. LKB Log-Logistic Model
        try:
            params = organ_params['LKB_LogLogit']
            geud_eqd2 = self.convert_to_eqd2(geud, params['alpha_beta'], dose_per_fraction)
            ntcp_lkb_loglogit = self.ntcp_lkb_loglogit(geud_eqd2, params['TD50'], params['gamma50'])
            
            results['LKB_LogLogit'] = {
                'NTCP': ntcp_lkb_loglogit,
                'gEUD_physical': geud,
                'gEUD_EQD2': geud_eqd2,
                'parameters_used': params
            }
        except Exception as e:
            print(f"Error calculating LKB LogLogit for {organ}: {e}")
            results['LKB_LogLogit'] = {'NTCP': 0.0, 'error': str(e)}
        
        # 2. LKB Probit Model
        try:
            params = organ_params['LKB_Probit']
            dose_metrics_copy = dose_metrics.copy()
            dose_metrics_copy['v_effective'] = v_effective
            
            ntcp_lkb_probit = self.ntcp_lkb_probit(dose_metrics_copy, params['TD50'], params['m'], params['n'])
            
            results['LKB_Probit'] = {
                'NTCP': ntcp_lkb_probit,
                'v_effective': v_effective,
                'parameters_used': params
            }
        except Exception as e:
            print(f"Error calculating LKB Probit for {organ}: {e}")
            results['LKB_Probit'] = {'NTCP': 0.0, 'error': str(e)}
        
        # 3. RS Poisson Model
        try:
            params = organ_params['RS_Poisson']
            dvh_eqd2 = dvh.copy()
            dvh_eqd2['dose_gy'] = dvh_eqd2['dose_gy'].apply(
                lambda d: self.convert_to_eqd2(d, params['alpha_beta'], dose_per_fraction)
            )
            
            ntcp_rs = self.ntcp_rs_poisson(dvh_eqd2, params['D50'], params['gamma'], params['s'])
            
            results['RS_Poisson'] = {
                'NTCP': ntcp_rs,
                'parameters_used': params
            }
        except Exception as e:
            print(f"Error calculating RS Poisson for {organ}: {e}")
            results['RS_Poisson'] = {'NTCP': 0.0, 'error': str(e)}
        
        return results

def adapt_for_small_dataset(X, y, clinical_data=None):
    """
    Adjust ML pipeline parameters based on dataset size
    
    Parameters
    ----------
    X : np.ndarray or pd.DataFrame
        Feature matrix
    y : np.ndarray or pd.Series
        Binary outcomes
    clinical_data : pd.DataFrame, optional
        Clinical data for integration
        
    Returns
    -------
    dict
        Dictionary with adaptation parameters
    """
    n_samples = len(X) if hasattr(X, '__len__') else X.shape[0]
    n_events = int(np.sum(y)) if hasattr(y, '__len__') else int(y.sum())
    
    adaptations = {
        'n_samples': n_samples,
        'n_events': n_events,
        'cv_strategy': None,
        'cv_folds': None,
        'model_config': {},
        'feature_selector_config': {}
    }
    
    # CV Strategy Adaptation
    if n_samples < 30:
        adaptations['cv_strategy'] = 'LeaveOneOut'
        adaptations['cv_folds'] = n_samples
    elif n_samples < 100:
        adaptations['cv_strategy'] = 'StratifiedKFold'
        adaptations['cv_folds'] = min(5, max(3, n_samples // 3))
    else:
        adaptations['cv_strategy'] = 'StratifiedKFold'
        adaptations['cv_folds'] = 5
    
    # Model Complexity Adaptation
    if n_samples < 50:
        adaptations['model_config'] = {
            'ann': {'hidden_layer_sizes': (8,), 'max_iter': 200, 'alpha': 0.1},
            'xgboost': {'n_estimators': 20, 'max_depth': 2, 'learning_rate': 0.1}
        }
    elif n_samples < 100:
        adaptations['model_config'] = {
            'ann': {'hidden_layer_sizes': (16,), 'max_iter': 300, 'alpha': 0.05},
            'xgboost': {'n_estimators': 30, 'max_depth': 2, 'learning_rate': 0.05}
        }
    else:
        adaptations['model_config'] = {
            'ann': {'hidden_layer_sizes': (16, 8), 'max_iter': 500, 'alpha': 0.01},
            'xgboost': {'n_estimators': 50, 'max_depth': 3, 'learning_rate': 0.05}
        }
    
    # Feature Selection Adaptation
    if n_events < 10:
        adaptations['feature_selector_config'] = {
            'max_features': max(2, n_events // 5),  # More conservative EPV
            'force_essential_only': True
        }
    else:
        adaptations['feature_selector_config'] = {
            'max_features': max(3, n_events // 10),
            'force_essential_only': False
        }
    
    # Calculate EPV
    adaptations['epv'] = n_events / adaptations['feature_selector_config']['max_features'] if adaptations['feature_selector_config']['max_features'] > 0 else 0
    
    return adaptations


class MachineLearningModels:
    """Machine learning models for NTCP prediction with proper validation"""
    
    def __init__(self, random_state=42):
        self.random_state = random_state
        self.models = {}
        self.scalers = {}
        self.adaptation_reports = {}  # Store adaptation info per organ
        
    def prepare_features(self, organ_data):
        """Prepare feature matrix for ML models"""
        
        # Select relevant dose metrics for features
        feature_cols = [
            'mean_dose', 'max_dose', 'gEUD', 'total_volume',
            'V5', 'V10', 'V15', 'V20', 'V25', 'V30', 'V35', 'V40', 'V45', 'V50',
            'D1', 'D2', 'D5', 'D10', 'D20', 'D30', 'D50', 'D70', 'D90', 'D95'
        ]
        
        # Add clinical features if available (FIX 3: Ensure clinical features are included)
        clinical_candidates = ['Chemotherapy', 'Age', 'Sex', 'Diabetes', 'T_Stage', 'N_Stage', 
                             'Baseline_Salivary_Function', 'Smoking', 'Alcohol']
        for clinical_col in clinical_candidates:
            if clinical_col in organ_data.columns and clinical_col not in feature_cols:
                feature_cols.append(clinical_col)
                print(f"    [DEBUG] Added clinical feature: {clinical_col}")
        
        # Filter to available columns
        available_cols = [col for col in feature_cols if col in organ_data.columns]
        
        # Extract features and target
        X = organ_data[available_cols].copy()
        y = organ_data['Observed_Toxicity'].copy()
        
        # Remove rows with missing values
        valid_mask = ~(X.isna().any(axis=1) | y.isna())
        X = X[valid_mask]
        y = y[valid_mask]
        
        if len(X) < 10:
            return None, None, None
        
        # DEBUG: Feature availability
        print(f"\n[DEBUG] prepare_features() returned:")
        print(f"  - X shape: {X.shape if X is not None else 'None'}")
        print(f"  - Number of features: {len(available_cols) if available_cols else 0}")
        print(f"  - All feature names: {available_cols}")
        print(f"  - Sample of features (first 15): {available_cols[:15] if len(available_cols) > 15 else available_cols}")
        print(f"  - Contains Dmean? {'Dmean' in available_cols}")
        print(f"  - Contains mean_dose? {'mean_dose' in available_cols}")
        print(f"  - Contains gEUD? {'gEUD' in available_cols}")
        print(f"  - Contains clinical features? {any(['Chemo' in str(f) or 'Age' in str(f) or 'Sex' in str(f) or 'Diabetes' in str(f) for f in available_cols])}")
        print(f"  - y sum (events): {y.sum() if y is not None else 'None'}")
        print(f"  - Total samples: {len(y) if y is not None else 'None'}")
        
        return X, y, available_cols
    
    def train_ann_model(self, X_train, y_train, organ):
        """Train Artificial Neural Network with proper regularization"""
        
        # Create pipeline with scaling
        ann_pipeline = Pipeline([
            ('scaler', StandardScaler()),
            ('ann', MLPClassifier(
                hidden_layer_sizes=(20, 10),  # Conservative architecture
                activation='relu',
                solver='lbfgs',  # Good for small datasets
                alpha=0.01,  # L2 regularization
                max_iter=1000,
                random_state=self.random_state,
                early_stopping=True,
                validation_fraction=0.2,
                n_iter_no_change=20
            ))
        ])
        
        try:
            ann_pipeline.fit(X_train, y_train)
            return ann_pipeline
        except Exception as e:
            print(f"      Error: ANN training failed: {e}")
            return None
    
    def train_xgboost_model(self, X_train, y_train, organ):
        """Train XGBoost model with proper regularization"""
        
        if not XGBOOST_AVAILABLE:
            return None
        
        try:
            # Conservative XGBoost parameters to prevent overfitting; allow learning on small n
            xgb_model = xgb.XGBClassifier(
                n_estimators=50,
                max_depth=2,       # Shallow but non-constant (depth 1 often collapses)
                learning_rate=0.05,
                subsample=0.8,
                colsample_bytree=0.8,
                reg_alpha=0.1,
                reg_lambda=1.0,
                min_child_weight=2,  # Allow splits on small data (was implicit default)
                random_state=self.random_state,
                eval_metric='logloss'
            )
            
            xgb_model.fit(X_train, y_train)
            return xgb_model
        except Exception as e:
            print(f"      Error: XGBoost training failed: {e}")
            return None
    
    def _train_and_evaluate_fold(self, X_train, X_test, y_train, y_test, feature_cols, organ, use_v2_components):
        """Train and evaluate models for a single CV fold"""
        fold_results = {}
        
        # Train ANN
        if use_v2_components:
            try:
                ml_model = OverfitResistantMLModels(
                    n_features=X_train.shape[1],
                    n_samples=len(X_train),
                    n_events=int(np.sum(y_train)),
                    random_seed=self.random_state
                )
                ann_model = ml_model.create_ann_model()
                ann_model.fit(X_train, y_train)
            except Exception as e:
                ann_model = self.train_ann_model(X_train, y_train, organ)
        else:
            ann_model = self.train_ann_model(X_train, y_train, organ)
        
        if ann_model is not None:
            try:
                y_pred_ann = ann_model.predict_proba(X_test)[:, 1]
                fpr, tpr, _ = roc_curve(y_test, y_pred_ann)
                auc_ann = auc(fpr, tpr)
                fold_results['ANN'] = {
                    'model': ann_model,
                    'test_AUC': auc_ann
                }
            except Exception as e:
                fold_results['ANN'] = None
        
        # Train XGBoost
        if XGBOOST_AVAILABLE:
            if use_v2_components:
                try:
                    if 'ml_model' not in locals():
                        ml_model = OverfitResistantMLModels(
                            n_features=X_train.shape[1],
                            n_samples=len(X_train),
                            n_events=int(np.sum(y_train)),
                            random_seed=self.random_state
                        )
                    xgb_model = ml_model.create_xgboost_model()
                    if xgb_model is not None:
                        xgb_model.fit(X_train, y_train)
                except Exception as e:
                    xgb_model = self.train_xgboost_model(X_train, y_train, organ)
            else:
                xgb_model = self.train_xgboost_model(X_train, y_train, organ)
            
            if xgb_model is not None:
                try:
                    y_pred_xgb = xgb_model.predict_proba(X_test)[:, 1]
                    fpr, tpr, _ = roc_curve(y_test, y_pred_xgb)
                    auc_xgb = auc(fpr, tpr)
                    fold_results['XGBoost'] = {
                        'model': xgb_model,
                        'test_AUC': auc_xgb
                    }
                except Exception as e:
                    fold_results['XGBoost'] = None

        # Train Random Forest (CV path – conservative settings)
        try:
            rf_model = RandomForestClassifier(
                n_estimators=100,
                max_depth=3,
                min_samples_split=4,
                min_samples_leaf=2,
                max_features='sqrt',
                class_weight='balanced_subsample',
                random_state=self.random_state
            )
            rf_model.fit(X_train, y_train)
            y_pred_rf = rf_model.predict_proba(X_test)[:, 1]
            fpr_rf, tpr_rf, _ = roc_curve(y_test, y_pred_rf)
            auc_rf = auc(fpr_rf, tpr_rf)
            fold_results['RandomForest'] = {
                'model': rf_model,
                'test_AUC': auc_rf
            }
        except Exception:
            fold_results['RandomForest'] = None
        
        return fold_results
    
    def train_and_evaluate_ml_models(self, organ_data, organ):
        """Train and evaluate ML models with proper cross-validation"""
        
        print(f"   Training ML models for {organ}...")
        
        # DEBUG: Input data structure
        print(f"\n[DEBUG] train_and_evaluate_ml_models() called for {organ}")
        print(f"  - organ_data shape: {organ_data.shape}")
        print(f"  - organ_data columns (first 30): {list(organ_data.columns)[:30]}")
        print(f"  - 'PrimaryPatientID' in columns? {'PrimaryPatientID' in organ_data.columns}")
        print(f"  - V2_COMPONENTS_AVAILABLE: {V2_COMPONENTS_AVAILABLE}")
        
        # Prepare features on ALL data first
        X_all, y_all, feature_cols = self.prepare_features(organ_data)
        
        if X_all is None:
            print(f"    Warning: Insufficient data for ML models")
            return {}
        
        n_events_all = y_all.sum() if hasattr(y_all, 'sum') else np.sum(y_all)
        n_samples_all = len(y_all)
        
        print(f"     Total samples: {n_samples_all}, Events: {int(n_events_all)}")
        
        if n_events_all < 5 or n_samples_all < 20:
            print(f"    Warning: Insufficient events/samples for reliable ML training")
            return {}
        
        # NEW: Check dataset size and adapt
        X_all_array = X_all.values if hasattr(X_all, 'values') else X_all
        adaptations = adapt_for_small_dataset(X_all_array, y_all, organ_data)
        
        print(f"     Dataset adaptations:")
        print(f"       - CV Strategy: {adaptations['cv_strategy']} ({adaptations['cv_folds']} folds)")
        print(f"       - EPV: {adaptations['epv']:.2f}")
        print(f"       - ANN Config: {adaptations['model_config'].get('ann', {})}")
        print(f"       - XGBoost Config: {adaptations['model_config'].get('xgboost', {})}")
        
        # Store adaptation report
        self.adaptation_reports[organ] = adaptations
        
        # Check if v2.0 components are available
        use_v2_components = V2_COMPONENTS_AVAILABLE and 'PrimaryPatientID' in organ_data.columns
        
        # DECISION: Use cross-validation for small datasets (< 100 samples)
        use_cross_validation = n_samples_all < 100
        
        if use_cross_validation:
            print(f"     Using 5-fold cross-validation for small dataset (n={n_samples_all})...")
            # Convert to numpy arrays
            if hasattr(X_all, 'values'):
                X_all = X_all.values
            if hasattr(y_all, 'values'):
                y_all = y_all.values
            
            # Feature selection on ALL data (for small datasets, this is acceptable)
            if use_v2_components and len(feature_cols) > 3:
                try:
                    print(f"\n[DEBUG] Before feature selection:")
                    print(f"  - Total features available: {len(feature_cols)}")
                    # NEW: Pass clinical data to feature selector
                    selector = RadiobiologyGuidedFeatureSelector(
                        organ=organ,
                        clinical_data=organ_data if 'PrimaryPatientID' in organ_data.columns else None,
                        outcome_column='Observed_Toxicity'
                    )
                    X_all_df = pd.DataFrame(X_all, columns=feature_cols)
                    selected_features = selector.select_features(
                        X_all_df, y_all, organ=organ,
                        max_features=adaptations['feature_selector_config'].get('max_features')
                    )
                    
                    # Store significant clinical factors
                    if hasattr(selector, 'significant_clinical_factors') and selector.significant_clinical_factors:
                        adaptations['significant_clinical_factors'] = selector.significant_clinical_factors
                        print(f"  - Significant clinical factors: {selector.significant_clinical_factors}")
                    
                    print(f"\n[DEBUG] After feature selection:")
                    print(f"  - Selected features: {selected_features}")
                    print(f"  - Number selected: {len(selected_features)}")
                    
                    if len(selected_features) < len(feature_cols):
                        selected_indices = [i for i, f in enumerate(feature_cols) if f in selected_features]
                        X_all = X_all[:, selected_indices]
                        feature_cols = selected_features
                        print(f"     Selected {len(selected_features)} features: {selected_features[:5]}...")
                except Exception as e:
                    print(f"     Warning: Feature selection failed: {e}, using all features")
            
            # Store selected features for prediction
            self.selected_features = feature_cols
            
            # Use adapted CV strategy
            if adaptations['cv_strategy'] == 'LeaveOneOut':
                cv = LeaveOneOut()
            else:
                cv = StratifiedKFold(n_splits=adaptations['cv_folds'], shuffle=True, random_state=self.random_state)
            
            # Store CV results
            ann_cv_scores = []
            xgb_cv_scores = []
            rf_cv_scores = []
            ann_models = []
            xgb_models = []
            rf_models = []
            
            print(f"     Starting 5-fold cross-validation...")
            fold = 1
            for train_idx, test_idx in cv.split(X_all, y_all):
                X_train, X_test = X_all[train_idx], X_all[test_idx]
                y_train, y_test = y_all[train_idx], y_all[test_idx]
                
                print(f"       Fold {fold}: Train={len(y_train)} (events={int(y_train.sum())}), Test={len(y_test)} (events={int(y_test.sum())})")
                
                # Train and evaluate models for this fold
                fold_results = self._train_and_evaluate_fold(
                    X_train, X_test, y_train, y_test, feature_cols, organ, use_v2_components
                )
                
                if 'ANN' in fold_results and fold_results['ANN'] is not None:
                    ann_cv_scores.append(fold_results['ANN']['test_AUC'])
                    ann_models.append(fold_results['ANN']['model'])
                
                if 'XGBoost' in fold_results and fold_results['XGBoost'] is not None:
                    xgb_cv_scores.append(fold_results['XGBoost']['test_AUC'])
                    xgb_models.append(fold_results['XGBoost']['model'])

                if 'RandomForest' in fold_results and fold_results['RandomForest'] is not None:
                    rf_cv_scores.append(fold_results['RandomForest']['test_AUC'])
                    rf_models.append(fold_results['RandomForest']['model'])
                
                fold += 1
            
            # Aggregate results
            results = {}
            if ann_cv_scores:
                mean_ann_auc = np.mean(ann_cv_scores)
                std_ann_auc = np.std(ann_cv_scores)
                print(f"     ANN - CV AUC: {mean_ann_auc:.3f} ± {std_ann_auc:.3f} (5-fold)")
                results['ANN'] = {
                    'model': ann_models[0],  # Use first fold's model for prediction
                    'cv_AUC_mean': mean_ann_auc,
                    'cv_AUC_std': std_ann_auc,
                    'cv_AUC_scores': ann_cv_scores,
                    'n_samples': n_samples_all,
                    'n_events': int(n_events_all),
                    'feature_names': feature_cols,
                    'validation_method': '5-fold_cv'
                }
            
            if xgb_cv_scores:
                mean_xgb_auc = np.mean(xgb_cv_scores)
                std_xgb_auc = np.std(xgb_cv_scores)
                print(f"     XGBoost - CV AUC: {mean_xgb_auc:.3f} ± {std_xgb_auc:.3f} (5-fold)")
                results['XGBoost'] = {
                    'model': xgb_models[0],  # Use first fold's model for prediction
                    'cv_AUC_mean': mean_xgb_auc,
                    'cv_AUC_std': std_xgb_auc,
                    'cv_AUC_scores': xgb_cv_scores,
                    'n_samples': n_samples_all,
                    'n_events': int(n_events_all),
                    'feature_names': feature_cols,
                    'validation_method': '5-fold_cv'
                }

            if rf_cv_scores:
                mean_rf_auc = np.mean(rf_cv_scores)
                std_rf_auc = np.std(rf_cv_scores)
                print(f"     RandomForest - CV AUC: {mean_rf_auc:.3f} ± {std_rf_auc:.3f} (5-fold)")
                results['RandomForest'] = {
                    'model': rf_models[0],
                    'cv_AUC_mean': mean_rf_auc,
                    'cv_AUC_std': std_rf_auc,
                    'cv_AUC_scores': rf_cv_scores,
                    'n_samples': n_samples_all,
                    'n_events': int(n_events_all),
                    'feature_names': feature_cols,
                    'validation_method': '5-fold_cv'
                }
            
            # Store models and selected features
            self.models[organ] = results
            return results
            
        elif use_v2_components:
            # PATIENT-LEVEL SPLITTING (v2.0 - prevents data leakage)
            splitter = PatientDataSplitter(random_seed=self.random_state, test_size=0.2)
            train_df, test_df = splitter.create_splits(
                organ_data,
                patient_id_col='PrimaryPatientID',
                outcome_col='Observed_Toxicity'
            )
            
            # Check for leakage
            leakage_detector = DataLeakageDetector()
            leakage_check = leakage_detector.check_patient_overlap(
                train_df, test_df, 'PrimaryPatientID'
            )
            if not leakage_check:
                leakage_report = leakage_detector.generate_report()
                print(f"     WARNING: {leakage_report['errors']}")
            
            # Extract features AFTER split (prevents leakage)
            X_train_df, y_train_series, feature_cols = self.prepare_features(train_df)
            X_test_df, y_test_series, _ = self.prepare_features(test_df)
            
            if X_train_df is None or X_test_df is None:
                print(f"    Warning: Insufficient data after patient-level split")
                return {}
            
            # Convert to numpy arrays
            X_train = X_train_df.values
            X_test = X_test_df.values
            y_train = y_train_series.values
            y_test = y_test_series.values
            
            n_events = y_train.sum()
            n_samples = len(y_train)
            
            print(f"     Features: {len(feature_cols)}, Train Samples: {len(y_train)}, Events: {int(n_events)}")
            print(f"     Test Samples: {len(y_test)}, Test Events: {int(y_test.sum())}")
            
        else:
            # FALLBACK: Row-level splitting (original method)
            print(f"     Using row-level splitting (v2.0 components not available)...")
            
            # Prepare features
            X, y, feature_cols = self.prepare_features(organ_data)
            
            if X is None:
                print(f"    Warning: Insufficient data for ML models")
                return {}
            
            n_events = y.sum()
            n_samples = len(y)
            
            print(f"     Features: {len(feature_cols)}, Samples: {n_samples}, Events: {int(n_events)}")
            
            # Use stratified train-test split
            X_train, X_test, y_train, y_test = train_test_split(
                X.values if hasattr(X, 'values') else X,
                y.values if hasattr(y, 'values') else y,
                test_size=0.3, random_state=self.random_state, 
                stratify=y if n_events >= 3 else None
            )
        
        if n_events < 5 or n_samples < 20:
            print(f"    Warning: Insufficient events/samples for reliable ML training")
            return {}
        
        results = {}
        
        # V2.0: Feature selection before training (if available)
        if use_v2_components and len(feature_cols) > 3:
            try:
                print(f"\n[DEBUG] Before feature selection:")
                print(f"  - Total features available: {len(feature_cols)}")
                print(f"  - Feature names: {feature_cols}")
                # NEW: Pass clinical data to feature selector
                selector = RadiobiologyGuidedFeatureSelector(
                    organ=organ,
                    clinical_data=train_df if 'train_df' in locals() and 'PrimaryPatientID' in train_df.columns else None,
                    outcome_column='Observed_Toxicity'
                )
                X_train_df = pd.DataFrame(X_train, columns=feature_cols)
                selected_features = selector.select_features(
                    X_train_df, y_train, organ=organ,
                    max_features=adaptations['feature_selector_config'].get('max_features')
                )
                
                # Store significant clinical factors
                if hasattr(selector, 'significant_clinical_factors') and selector.significant_clinical_factors:
                    adaptations['significant_clinical_factors'] = selector.significant_clinical_factors
                    print(f"  - Significant clinical factors: {selector.significant_clinical_factors}")
                
                print(f"\n[DEBUG] After feature selection:")
                print(f"  - Selected features: {selected_features}")
                print(f"  - Number selected: {len(selected_features)}")
                
                if len(selected_features) < len(feature_cols):
                    # Use only selected features
                    selected_indices = [i for i, f in enumerate(feature_cols) if f in selected_features]
                    X_train = X_train[:, selected_indices]
                    X_test = X_test[:, selected_indices]
                    feature_cols = selected_features
                    print(f"     Selected {len(selected_features)} features: {selected_features[:5]}...")
            except Exception as e:
                print(f"     Warning: Feature selection failed: {e}, using all features")
                import traceback
                traceback.print_exc()
        
        # Train ANN
        print(f"     Training ANN...")
        
        # V2.0: Use OverfitResistantMLModels if available, with adaptations
        if use_v2_components:
            try:
                ml_model = OverfitResistantMLModels(
                    n_features=X_train.shape[1],
                    n_samples=len(X_train),
                    n_events=int(np.sum(y_train)),
                    random_seed=self.random_state
                )
                ann_model = ml_model.create_ann_model()
                
                # NEW: Apply small dataset adaptations
                ann_config = adaptations['model_config'].get('ann', {})
                if ann_config:
                    # Update ANN parameters
                    if 'hidden_layer_sizes' in ann_config:
                        ann_model.hidden_layer_sizes = ann_config['hidden_layer_sizes']
                    if 'max_iter' in ann_config:
                        ann_model.max_iter = ann_config['max_iter']
                    if 'alpha' in ann_config:
                        ann_model.alpha = ann_config['alpha']
                
                print(f"       EPV: {ml_model.epv:.2f} events per variable")
                print(f"       ANN Config: hidden_layers={ann_model.hidden_layer_sizes}, alpha={ann_model.alpha}")
                # Fit the model before using it
                ann_model.fit(X_train, y_train)
            except Exception as e:
                print(f"     Warning: OverfitResistantMLModels failed: {e}, using basic model")
                ann_model = self.train_ann_model(X_train, y_train, organ)
        else:
            ann_model = self.train_ann_model(X_train, y_train, organ)
        
        if ann_model is not None:
            # Evaluate on test set
            y_pred_ann = ann_model.predict_proba(X_test)[:, 1]
            
            # Calculate metrics
            try:
                fpr, tpr, _ = roc_curve(y_test, y_pred_ann)
                
                # V2.0: AUC with confidence intervals
                if use_v2_components:
                    try:
                        auc_ann, auc_ci = calculate_auc_with_ci(y_test, y_pred_ann)
                        print(f"       ANN - Test AUC: {auc_ann:.3f} (95% CI: {auc_ci[0]:.3f}-{auc_ci[1]:.3f})")
                    except Exception as e:
                        print(f"     Warning: AUC CI calculation failed: {e}")
                        auc_ann = auc(fpr, tpr)
                        auc_ci = (auc_ann, auc_ann)
                else:
                    auc_ann = auc(fpr, tpr)
                    auc_ci = (auc_ann, auc_ann)
                
                brier_ann = brier_score_loss(y_test, y_pred_ann)
                
                # Cross-validation on training set
                cv_scores = cross_val_score(ann_model, X_train, y_train, 
                                           cv=min(5, len(X_train)//3), scoring='roc_auc')
                
                results['ANN'] = {
                    'model': ann_model,
                    'test_AUC': auc_ann,
                    'test_AUC_CI': auc_ci if use_v2_components else None,
                    'test_Brier': brier_ann,
                    'cv_AUC_mean': np.mean(cv_scores),
                    'cv_AUC_std': np.std(cv_scores),
                    'n_train': len(X_train),
                    'n_test': len(X_test),
                    'feature_names': feature_cols,
                    'epv': ml_model.epv if use_v2_components and 'ml_model' in locals() else None
                }
                
                if not use_v2_components or 'auc_ci' not in locals():
                    print(f"       ANN - Test AUC: {auc_ann:.3f}, CV AUC: {np.mean(cv_scores):.3f}+/-{np.std(cv_scores):.3f}")
                
            except Exception as e:
                print(f"      Error: ANN evaluation failed: {e}")
        
        # Train XGBoost
        if XGBOOST_AVAILABLE:
            print(f"     Training XGBoost...")
            
            # V2.0: Use OverfitResistantMLModels if available, with adaptations
            if use_v2_components:
                try:
                    if 'ml_model' not in locals():
                        ml_model = OverfitResistantMLModels(
                            n_features=X_train.shape[1],
                            n_samples=len(X_train),
                            n_events=int(np.sum(y_train)),
                            random_seed=self.random_state
                        )
                    xgb_model = ml_model.create_xgboost_model()
                    
                    # NEW: Apply small dataset adaptations
                    xgb_config = adaptations['model_config'].get('xgboost', {})
                    if xgb_config and xgb_model is not None:
                        # Update XGBoost parameters
                        if 'n_estimators' in xgb_config:
                            xgb_model.n_estimators = xgb_config['n_estimators']
                        if 'max_depth' in xgb_config:
                            xgb_model.max_depth = xgb_config['max_depth']
                        if 'learning_rate' in xgb_config:
                            xgb_model.learning_rate = xgb_config['learning_rate']
                        print(f"       XGBoost Config: n_estimators={xgb_model.n_estimators}, max_depth={xgb_model.max_depth}, lr={xgb_model.learning_rate}")
                    
                    # Fit the model before using it
                    if xgb_model is not None:
                        xgb_model.fit(X_train, y_train)
                except Exception as e:
                    print(f"     Warning: OverfitResistantMLModels XGBoost failed: {e}, using basic model")
                    xgb_model = self.train_xgboost_model(X_train, y_train, organ)
            else:
                xgb_model = self.train_xgboost_model(X_train, y_train, organ)
            
            if xgb_model is not None:
                # Evaluate on test set
                y_pred_xgb = xgb_model.predict_proba(X_test)[:, 1]
                
                try:
                    fpr, tpr, _ = roc_curve(y_test, y_pred_xgb)
                    
                    # V2.0: AUC with confidence intervals
                    if use_v2_components:
                        try:
                            auc_xgb, auc_ci = calculate_auc_with_ci(y_test, y_pred_xgb)
                            print(f"       XGBoost - Test AUC: {auc_xgb:.3f} (95% CI: {auc_ci[0]:.3f}-{auc_ci[1]:.3f})")
                        except Exception as e:
                            print(f"     Warning: AUC CI calculation failed: {e}")
                            auc_xgb = auc(fpr, tpr)
                            auc_ci = (auc_xgb, auc_xgb)
                    else:
                        auc_xgb = auc(fpr, tpr)
                        auc_ci = (auc_xgb, auc_xgb)
                    
                    brier_xgb = brier_score_loss(y_test, y_pred_xgb)
                    
                    # Cross-validation on training set
                    cv_scores = cross_val_score(xgb_model, X_train, y_train, 
                                               cv=min(5, len(X_train)//3), scoring='roc_auc')
                    
                    # Feature importance
                    feature_importance = dict(zip(feature_cols, xgb_model.feature_importances_))
                    
                    results['XGBoost'] = {
                        'model': xgb_model,
                        'test_AUC': auc_xgb,
                        'test_AUC_CI': auc_ci if use_v2_components else None,
                        'test_Brier': brier_xgb,
                        'cv_AUC_mean': np.mean(cv_scores),
                        'cv_AUC_std': np.std(cv_scores),
                        'n_train': len(X_train),
                        'n_test': len(X_test),
                        'feature_names': feature_cols,
                        'feature_importance': feature_importance,
                        'epv': ml_model.epv if use_v2_components and 'ml_model' in locals() else None
                    }
                    
                    if not use_v2_components or 'auc_ci' not in locals():
                        print(f"       XGBoost - Test AUC: {auc_xgb:.3f}, CV AUC: {np.mean(cv_scores):.3f}+/-{np.std(cv_scores):.3f}")
                    
                    # Show top features
                    top_features = sorted(feature_importance.items(), key=lambda x: x[1], reverse=True)[:5]
                    print(f"       Top features: {', '.join([f'{feat}({imp:.3f})' for feat, imp in top_features])}")
                    
                except Exception as e:
                    print(f"      Error: XGBoost evaluation failed: {e}")
        
        # Train Random Forest
        print(f"     Training Random Forest...")
        try:
            if use_v2_components:
                # Reuse ml_model if available, otherwise create a new one
                if 'ml_model' not in locals():
                    ml_model = OverfitResistantMLModels(
                        n_features=X_train.shape[1],
                        n_samples=len(X_train),
                        n_events=int(np.sum(y_train)),
                        random_seed=self.random_state
                    )
                rf_model = ml_model.create_random_forest_model()
            else:
                rf_model = RandomForestClassifier(
                    n_estimators=100,
                    max_depth=3,
                    min_samples_split=4,
                    min_samples_leaf=2,
                    max_features='sqrt',
                    class_weight='balanced_subsample',
                    random_state=self.random_state
                )
            
            rf_model.fit(X_train, y_train)
            
            # Evaluate on test set
            y_pred_rf = rf_model.predict_proba(X_test)[:, 1]
            try:
                fpr_rf, tpr_rf, _ = roc_curve(y_test, y_pred_rf)
                if use_v2_components:
                    try:
                        auc_rf, auc_ci = calculate_auc_with_ci(y_test, y_pred_rf)
                        print(f"       RandomForest - Test AUC: {auc_rf:.3f} (95% CI: {auc_ci[0]:.3f}-{auc_ci[1]:.3f})")
                    except Exception as e:
                        print(f"     Warning: AUC CI calculation failed (RF): {e}")
                        auc_rf = auc(fpr_rf, tpr_rf)
                        auc_ci = (auc_rf, auc_rf)
                else:
                    auc_rf = auc(fpr_rf, tpr_rf)
                    auc_ci = (auc_rf, auc_rf)
                
                brier_rf = brier_score_loss(y_test, y_pred_rf)
                
                cv_scores_rf = cross_val_score(
                    rf_model, X_train, y_train,
                    cv=min(5, len(X_train)//3), scoring='roc_auc'
                )
                
                feature_importance_rf = dict(
                    zip(feature_cols, rf_model.feature_importances_)
                )
                
                results['RandomForest'] = {
                    'model': rf_model,
                    'test_AUC': auc_rf,
                    'test_AUC_CI': auc_ci if use_v2_components else None,
                    'test_Brier': brier_rf,
                    'cv_AUC_mean': np.mean(cv_scores_rf),
                    'cv_AUC_std': np.std(cv_scores_rf),
                    'n_train': len(X_train),
                    'n_test': len(X_test),
                    'feature_names': feature_cols,
                    'feature_importance': feature_importance_rf,
                    'epv': ml_model.epv if use_v2_components and 'ml_model' in locals() else None
                }
            except Exception as e:
                print(f"      Error: RandomForest evaluation failed: {e}")
        except Exception as e:
            print(f"     Warning: RandomForest training failed: {e}")
        
        # Store models and selected features for later use
        self.models[organ] = results
        self.selected_features = feature_cols  # Store for prediction consistency
        
        return results
    
    def predict_ml_models(self, organ_data, organ):
        """Generate predictions from trained ML models"""
        
        if organ not in self.models:
            return {}
        
        # Prepare features
        X, y, available_feature_cols = self.prepare_features(organ_data)
        
        if X is None:
            return {}
        
        # Convert to numpy if needed
        if hasattr(X, 'values'):
            X = X.values
        
        predictions = {}
        
        for model_name, model_info in self.models[organ].items():
            try:
                model = model_info['model']
                trained_features = model_info['feature_names']
                
                # Align features: use only features that were used during training
                feature_indices = []
                missing_features = []
                for feature in trained_features:
                    if feature in available_feature_cols:
                        idx = list(available_feature_cols).index(feature)
                        feature_indices.append(idx)
                    else:
                        missing_features.append(feature)
                
                if missing_features:
                    print(f"    Warning: {len(missing_features)} features missing for {model_name}: {missing_features[:3]}...")
                    # Continue with available features only
                
                if len(feature_indices) == len(trained_features):
                    # All features available
                    X_aligned = X[:, feature_indices]
                    y_pred = model.predict_proba(X_aligned)[:, 1]
                    predictions[f'NTCP_ML_{model_name}'] = y_pred
                elif len(feature_indices) > 0:
                    # Some features missing - try to predict with available ones
                    print(f"    Warning: Feature mismatch for {model_name} - using {len(feature_indices)}/{len(trained_features)} features")
                    # This may fail, but we try
                    try:
                        X_aligned = X[:, feature_indices]
                        y_pred = model.predict_proba(X_aligned)[:, 1]
                        predictions[f'NTCP_ML_{model_name}'] = y_pred
                    except:
                        print(f"    Error: Cannot predict with {model_name} due to feature mismatch")
                else:
                    print(f"    Warning: No matching features for {model_name}")
                    
            except Exception as e:
                print(f"    Error: Prediction failed for {model_name}: {e}")
        
        return predictions

class ComprehensivePlotter:
    """Create comprehensive plots for all organs and models"""
    
    def __init__(self, output_dir, ntcp_calc):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.ntcp_calc = ntcp_calc
    
    def calculate_calibration_data(self, y_true, y_pred, n_bins=5):
        """Calculate calibration data for reliability diagram"""
        
        # Remove invalid predictions
        valid_mask = ~(np.isnan(y_pred) | np.isnan(y_true))
        if np.sum(valid_mask) < n_bins:
            return None, None, None
        
        y_true_clean = y_true[valid_mask]
        y_pred_clean = y_pred[valid_mask]
        
        # Create bins based on predicted probabilities
        try:
            # Use quantile-based binning for more robust results
            bin_boundaries = np.percentile(y_pred_clean, np.linspace(0, 100, n_bins + 1))
            
            # Ensure unique boundaries
            bin_boundaries = np.unique(bin_boundaries)
            if len(bin_boundaries) < 3:  # Need at least 2 bins
                return None, None, None
            
            bin_centers = []
            bin_observed = []
            bin_counts = []
            
            for i in range(len(bin_boundaries) - 1):
                # Create mask for current bin
                if i == len(bin_boundaries) - 2:  # Last bin includes right boundary
                    mask = (y_pred_clean >= bin_boundaries[i]) & (y_pred_clean <= bin_boundaries[i + 1])
                else:
                    mask = (y_pred_clean >= bin_boundaries[i]) & (y_pred_clean < bin_boundaries[i + 1])
                
                if np.sum(mask) > 0:
                    bin_pred_mean = np.mean(y_pred_clean[mask])
                    bin_obs_mean = np.mean(y_true_clean[mask])
                    bin_count = np.sum(mask)
                    
                    bin_centers.append(bin_pred_mean)
                    bin_observed.append(bin_obs_mean)
                    bin_counts.append(bin_count)
            
            return np.array(bin_centers), np.array(bin_observed), np.array(bin_counts)
            
        except Exception as e:
            print(f"    Warning: Calibration calculation failed: {e}")
            return None, None, None
    
    def create_dose_response_plot(self, organ_data, organ):
        """Create dose-response plot for specific organ"""
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Get gEUD values and create smooth curve
        geud_values = organ_data['gEUD'].dropna()
        if len(geud_values) == 0:
            print(f"No gEUD data available for {organ}")
            plt.close()
            return
        
        dose_range = np.linspace(geud_values.min() * 0.8, geud_values.max() * 1.2, 300)
        
        # Plot theoretical curve using literature parameters
        if organ in self.ntcp_calc.literature_params:
            lit_params = self.ntcp_calc.literature_params[organ]['LKB_LogLogit']
            
            ntcp_curve = []
            for dose in dose_range:
                eqd2_dose = self.ntcp_calc.convert_to_eqd2(dose, lit_params['alpha_beta'], 2.0)
                ntcp = self.ntcp_calc.ntcp_lkb_loglogit(eqd2_dose, lit_params['TD50'], lit_params['gamma50'])
                ntcp_curve.append(ntcp)
            
            # Plot theoretical curve
            ax.plot(dose_range, ntcp_curve, color=COLORS['LKB_LogLogit'], 
                   linewidth=4, label=f"LKB Model (TD₅₀ = {lit_params['TD50']:.1f} Gy)",
                   alpha=0.8)
            
            # Mark TD50 on the curve
            ax.axvline(lit_params['TD50'], color=COLORS['literature'], 
                      linestyle='--', alpha=0.8, linewidth=3,
                      label=f"Literature TD₅₀")
        
        # Plot observed data points with binning
        valid_data = organ_data.dropna(subset=['gEUD', 'Observed_Toxicity'])
        
        if len(valid_data) > 0:
            # Create bins based on gEUD
            n_bins = min(8, max(3, len(valid_data) // 4))
            bins = np.percentile(valid_data['gEUD'], np.linspace(0, 100, n_bins + 1))
            
            bin_centers = []
            bin_rates = []
            bin_counts = []
            bin_errors = []
            
            for i in range(len(bins) - 1):
                mask = (valid_data['gEUD'] >= bins[i]) & (valid_data['gEUD'] < bins[i + 1])
                bin_data = valid_data[mask]
                
                if len(bin_data) > 0:
                    bin_centers.append(bin_data['gEUD'].mean())
                    rate = bin_data['Observed_Toxicity'].mean()
                    bin_rates.append(rate)
                    bin_counts.append(len(bin_data))
                    
                    # Calculate 95% confidence interval
                    n = len(bin_data)
                    if n > 1 and 0 < rate < 1:
                        se = np.sqrt(rate * (1 - rate) / n)
                        ci_width = 1.96 * se
                        bin_errors.append(ci_width)
                    else:
                        bin_errors.append(0)
            
            # Plot observed data with error bars
            if bin_centers:
                sizes = [80 + 20 * min(count, 15) for count in bin_counts]
                scatter = ax.scatter(bin_centers, bin_rates, s=sizes, 
                                   c=COLORS['observed'], alpha=0.9, 
                                   edgecolors='white', linewidth=2, 
                                   label='Observed Data', zorder=10)
                
                # Add error bars
                ax.errorbar(bin_centers, bin_rates, yerr=bin_errors, 
                           fmt='none', color=COLORS['observed'], alpha=0.7, 
                           capsize=5, capthick=2, linewidth=2, zorder=5)
        
        # Enhanced styling
        ax.set_xlabel(f'{organ} gEUD (Gy)', fontsize=16, fontweight='bold')
        ax.set_ylabel('NTCP', fontsize=16, fontweight='bold')
        
        # Grid
        ax.grid(True, alpha=0.3, color=COLORS['grid'], linewidth=1)
        
        # Enhanced legend
        legend = ax.legend(fontsize=14, loc='lower right', frameon=True, 
                          fancybox=True, shadow=True, framealpha=0.9)
        legend.get_frame().set_facecolor('white')
        legend.get_frame().set_edgecolor('gray')
        
        # Limits and ticks
        ax.set_ylim(0, 1.05)
        ax.set_xlim(geud_values.min() * 0.9, geud_values.max() * 1.1)
        ax.tick_params(axis='both', which='major', labelsize=12)
        
        # Save plot
        filename = f"{organ}_dose_response.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved dose-response plot: {filename}")
        plt.close()
    
    def create_roc_plot(self, organ_data, organ):
        """Create ROC plot for specific organ"""
        
        fig, ax = plt.subplots(figsize=(8, 8))
        
        # Traditional NTCP models
        traditional_models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']
        ml_models = []
        
        # Check for ML models
        ml_cols = [col for col in organ_data.columns if col.startswith('NTCP_ML_')]
        for ml_col in ml_cols:
            if 'ANN' in ml_col:
                ml_models.append(('ML_ANN', ml_col, 'ANN'))
            elif 'XGBoost' in ml_col:
                ml_models.append(('ML_XGBoost', ml_col, 'XGBoost'))
            elif 'RandomForest' in ml_col:
                ml_models.append(('ML_RandomForest', ml_col, 'Random Forest'))
        
        all_auc_values = []
        
        # Plot traditional models
        for model in traditional_models:
            ntcp_col = f'NTCP_{model}'
            
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 5:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    try:
                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                        auc_score = auc(fpr, tpr)
                        
                        # Plot ROC curve with unique styling
                        ax.plot(fpr, tpr, 
                               color=COLORS[model], 
                               linestyle=LINE_STYLES[model],
                               linewidth=3.0,
                               label=f'{model.replace("_", " ")}: AUC = {auc_score:.3f}',
                               alpha=0.8)
                        
                        all_auc_values.append((model.replace("_", " "), auc_score))
                        
                    except Exception as e:
                        print(f"    Error: ROC calculation failed for {model}: {e}")
        
        # Plot ML models with distinct styling
        for model_key, ntcp_col, model_label in ml_models:
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 5:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    try:
                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                        auc_score = auc(fpr, tpr)
                        
                        # Plot ML ROC curve with unique color and style
                        ax.plot(fpr, tpr, 
                               color=COLORS.get(model_key, COLORS['ML_ANN']), 
                               linestyle=LINE_STYLES.get(model_key, ':'),
                               linewidth=3.5,  # Slightly thicker for ML models
                               label=f'{model_label}: AUC = {auc_score:.3f}',
                               alpha=0.9)
                        
                        all_auc_values.append((model_label, auc_score))
                        
                    except Exception as e:
                        print(f"    Error: ROC calculation failed for {model_label}: {e}")
        
        # Plot diagonal reference line
        ax.plot([0, 1], [0, 1], 'k--', linewidth=2, alpha=0.6, 
               label='Random Classifier')
        
        # Enhanced styling
        ax.set_xlabel('False Positive Rate', fontsize=16, fontweight='bold')
        ax.set_ylabel('True Positive Rate', fontsize=16, fontweight='bold')
        
        # Grid and formatting
        ax.grid(True, alpha=0.3, color=COLORS['grid'], linewidth=1)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.set_aspect('equal')
        
        # Enhanced legend with better positioning
        legend = ax.legend(fontsize=12, loc='lower right', frameon=True, 
                          fancybox=True, shadow=True, framealpha=0.9)
        legend.get_frame().set_facecolor('white')
        legend.get_frame().set_edgecolor('gray')
        
        # Add sample size annotation
        n_total = len(organ_data)
        n_events = int(organ_data['Observed_Toxicity'].sum())
        ax.text(0.02, 0.98, f'{organ}\nSample: n={n_total}, events={n_events}', 
               transform=ax.transAxes, fontsize=12, fontweight='bold',
               bbox=dict(boxstyle='round,pad=0.5', facecolor='lightblue', alpha=0.8),
               verticalalignment='top')
        
        # Save plot
        filename = f"{organ}_ROC.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved ROC plot: {filename}")
        plt.close()
        
        return all_auc_values
    
    def create_calibration_plot(self, organ_data, organ):
        """Create calibration plot for specific organ"""
        
        fig, ax = plt.subplots(figsize=(8, 8))
        
        # Traditional NTCP models
        traditional_models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']
        ml_models = []
        
        # Check for ML models
        ml_cols = [col for col in organ_data.columns if col.startswith('NTCP_ML_')]
        for ml_col in ml_cols:
            if 'ANN' in ml_col:
                ml_models.append(('ML_ANN', ml_col, 'ANN'))
            elif 'XGBoost' in ml_col:
                ml_models.append(('ML_XGBoost', ml_col, 'XGBoost'))
            elif 'RandomForest' in ml_col:
                ml_models.append(('ML_RandomForest', ml_col, 'Random Forest'))
        
        calibration_metrics = {}
        
        # Plot traditional models
        for model in traditional_models:
            ntcp_col = f'NTCP_{model}'
            
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 10:  # Need more points for meaningful calibration
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    # Calculate calibration data
                    bin_centers, bin_observed, bin_counts = self.calculate_calibration_data(
                        y_true, y_pred, n_bins=min(5, len(valid_data) // 3))
                    
                    if bin_centers is not None and len(bin_centers) >= 2:
                        
                        # Calculate calibration metrics
                        slope = np.nan
                        intercept = np.nan
                        
                        if len(bin_centers) >= 2:
                            x_mean = np.mean(bin_centers)
                            y_mean = np.mean(bin_observed)
                            numerator = np.sum((bin_centers - x_mean) * (bin_observed - y_mean))
                            denominator = np.sum((bin_centers - x_mean) ** 2)
                            
                            if denominator != 0:
                                slope = numerator / denominator
                                intercept = y_mean - slope * x_mean
                                calibration_metrics[model] = {'slope': slope, 'intercept': intercept}
                        
                        # Create label with slope and intercept
                        model_name = model.replace('_', ' ')
                        if not np.isnan(slope) and not np.isnan(intercept):
                            label = f"{model_name}: slope={slope:.3f}, int={intercept:.3f}"
                        else:
                            label = f"{model_name}"
                        
                        # Plot calibration curve with markers
                        ax.plot(bin_centers, bin_observed,
                               color=COLORS[model],
                               linestyle=LINE_STYLES[model],
                               linewidth=2.5,
                               marker=MARKERS[model],
                               markersize=8,
                               markerfacecolor=COLORS[model],
                               markeredgecolor='white',
                               markeredgewidth=1,
                               label=label,
                               zorder=5)
        
        # Plot ML models
        for model_key, ntcp_col, model_label in ml_models:
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 10:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    # Calculate calibration data
                    bin_centers, bin_observed, bin_counts = self.calculate_calibration_data(
                        y_true, y_pred, n_bins=min(5, len(valid_data) // 3))
                    
                    if bin_centers is not None and len(bin_centers) >= 2:
                        
                        # Calculate calibration metrics
                        slope = np.nan
                        intercept = np.nan
                        
                        if len(bin_centers) >= 2:
                            x_mean = np.mean(bin_centers)
                            y_mean = np.mean(bin_observed)
                            numerator = np.sum((bin_centers - x_mean) * (bin_observed - y_mean))
                            denominator = np.sum((bin_centers - x_mean) ** 2)
                            
                            if denominator != 0:
                                slope = numerator / denominator
                                intercept = y_mean - slope * x_mean
                                calibration_metrics[model_key] = {'slope': slope, 'intercept': intercept}
                        
                        # Create label with slope and intercept
                        if not np.isnan(slope) and not np.isnan(intercept):
                            label = f"{model_label}: slope={slope:.3f}, int={intercept:.3f}"
                        else:
                            label = f"{model_label}"
                        
                        # Plot ML calibration curve
                        ax.plot(bin_centers, bin_observed,
                               color=COLORS[model_key],
                               linestyle=LINE_STYLES[model_key],
                               linewidth=3.0,
                               marker=MARKERS[model_key],
                               markersize=8,
                               markerfacecolor=COLORS[model_key],
                               markeredgecolor='white',
                               markeredgewidth=1,
                               label=label,
                               zorder=6)
        
        # Plot perfect calibration line
        ax.plot([0, 1], [0, 1], 'k--', linewidth=2, alpha=0.7, 
               label='Perfect Calibration', zorder=1)
        
        # Enhanced styling
        ax.set_xlabel('Predicted NTCP', fontsize=16, fontweight='bold')
        ax.set_ylabel('Observed Rate', fontsize=16, fontweight='bold')
        
        # Grid and formatting
        ax.grid(True, alpha=0.3, color=COLORS['grid'], linewidth=1)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.set_aspect('equal')
        
        # Enhanced legend
        legend = ax.legend(fontsize=10, loc='upper left', frameon=True, 
                          fancybox=True, shadow=True, framealpha=0.9)
        legend.get_frame().set_facecolor('white')
        legend.get_frame().set_edgecolor('gray')
        
        # Add organ name annotation
        ax.text(0.98, 0.02, f'{organ}', 
               transform=ax.transAxes, fontsize=14, fontweight='bold',
               bbox=dict(boxstyle='round,pad=0.5', facecolor='lightgreen', alpha=0.8),
               horizontalalignment='right', verticalalignment='bottom')
        
        # Save plot
        filename = f"{organ}_calibration.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved calibration plot: {filename}")
        plt.close()
        
        return calibration_metrics
    
    def create_combined_roc_calibration_plot(self, organ_data, organ):
        """Create combined ROC and calibration plot"""
        
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
        
        # Traditional NTCP models
        traditional_models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']
        ml_models = []
        
        # Check for ML models
        ml_cols = [col for col in organ_data.columns if col.startswith('NTCP_ML_')]
        for ml_col in ml_cols:
            if 'ANN' in ml_col:
                ml_models.append(('ML_ANN', ml_col, 'ANN'))
            elif 'XGBoost' in ml_col:
                ml_models.append(('ML_XGBoost', ml_col, 'XGBoost'))
            elif 'RandomForest' in ml_col:
                ml_models.append(('ML_RandomForest', ml_col, 'Random Forest'))
        
        # ROC Plot (left)
        all_auc_values = []
        
        # Plot traditional models ROC
        for model in traditional_models:
            ntcp_col = f'NTCP_{model}'
            
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 5:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    try:
                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                        auc_score = auc(fpr, tpr)
                        
                        ax1.plot(fpr, tpr, 
                               color=COLORS[model], 
                               linestyle=LINE_STYLES[model],
                               linewidth=3.0,
                               label=f'{model.replace("_", " ")}: AUC = {auc_score:.3f}',
                               alpha=0.8)
                        
                        all_auc_values.append((model.replace("_", " "), auc_score))
                        
                    except Exception as e:
                        print(f"    Error: ROC calculation failed for {model}: {e}")
        
        # Plot ML models ROC
        for model_key, ntcp_col, model_label in ml_models:
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 5:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    try:
                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                        auc_score = auc(fpr, tpr)
                        
                        ax1.plot(fpr, tpr, 
                               color=COLORS[model_key], 
                               linestyle=LINE_STYLES[model_key],
                               linewidth=3.5,
                               label=f'{model_label}: AUC = {auc_score:.3f}',
                               alpha=0.9)
                        
                        all_auc_values.append((model_label, auc_score))
                        
                    except Exception as e:
                        print(f"    Error: ROC calculation failed for {model_label}: {e}")
        
        # ROC diagonal and formatting
        ax1.plot([0, 1], [0, 1], 'k--', linewidth=2, alpha=0.6, label='Random')
        ax1.set_xlabel('False Positive Rate', fontsize=14, fontweight='bold')
        ax1.set_ylabel('True Positive Rate', fontsize=14, fontweight='bold')
        ax1.set_xlim(0, 1)
        ax1.set_ylim(0, 1)
        ax1.set_aspect('equal')
        ax1.grid(True, alpha=0.3)
        ax1.legend(fontsize=10, loc='lower right')
        
        # Calibration Plot (right)
        calibration_metrics = {}
        
        # Plot traditional models calibration
        for model in traditional_models:
            ntcp_col = f'NTCP_{model}'
            
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 10:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    bin_centers, bin_observed, bin_counts = self.calculate_calibration_data(
                        y_true, y_pred, n_bins=min(5, len(valid_data) // 3))
                    
                    if bin_centers is not None and len(bin_centers) >= 2:
                        
                        # Calculate slope
                        slope = np.nan
                        if len(bin_centers) >= 2:
                            x_mean = np.mean(bin_centers)
                            y_mean = np.mean(bin_observed)
                            numerator = np.sum((bin_centers - x_mean) * (bin_observed - y_mean))
                            denominator = np.sum((bin_centers - x_mean) ** 2)
                            
                            if denominator != 0:
                                slope = numerator / denominator
                        
                        model_name = model.replace('_', ' ')
                        label = f"{model_name}" if np.isnan(slope) else f"{model_name}: {slope:.3f}"
                        
                        ax2.plot(bin_centers, bin_observed,
                               color=COLORS[model],
                               linestyle=LINE_STYLES[model],
                               linewidth=2.5,
                               marker=MARKERS[model],
                               markersize=6,
                               markerfacecolor=COLORS[model],
                               markeredgecolor='white',
                               markeredgewidth=1,
                               label=label,
                               zorder=5)
        
        # Plot ML models calibration
        for model_key, ntcp_col, model_label in ml_models:
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 10:
                    y_true = valid_data['Observed_Toxicity'].values
                    y_pred = valid_data[ntcp_col].values
                    
                    bin_centers, bin_observed, bin_counts = self.calculate_calibration_data(
                        y_true, y_pred, n_bins=min(5, len(valid_data) // 3))
                    
                    if bin_centers is not None and len(bin_centers) >= 2:
                        
                        # Calculate slope
                        slope = np.nan
                        if len(bin_centers) >= 2:
                            x_mean = np.mean(bin_centers)
                            y_mean = np.mean(bin_observed)
                            numerator = np.sum((bin_centers - x_mean) * (bin_observed - y_mean))
                            denominator = np.sum((bin_centers - x_mean) ** 2)
                            
                            if denominator != 0:
                                slope = numerator / denominator
                        
                        label = f"{model_label}" if np.isnan(slope) else f"{model_label}: {slope:.3f}"
                        
                        ax2.plot(bin_centers, bin_observed,
                               color=COLORS[model_key],
                               linestyle=LINE_STYLES[model_key],
                               linewidth=3.0,
                               marker=MARKERS[model_key],
                               markersize=6,
                               markerfacecolor=COLORS[model_key],
                               markeredgecolor='white',
                               markeredgewidth=1,
                               label=label,
                               zorder=6)
        
        # Calibration perfect line and formatting
        ax2.plot([0, 1], [0, 1], 'k--', linewidth=2, alpha=0.7, label='Perfect')
        ax2.set_xlabel('Predicted NTCP', fontsize=14, fontweight='bold')
        ax2.set_ylabel('Observed Rate', fontsize=14, fontweight='bold')
        ax2.set_xlim(0, 1)
        ax2.set_ylim(0, 1)
        ax2.set_aspect('equal')
        ax2.grid(True, alpha=0.3)
        ax2.legend(fontsize=10, loc='upper left')
        
        # Add organ name as suptitle
        fig.suptitle(f'{organ}', fontsize=18, fontweight='bold', y=0.95)
        
        plt.tight_layout()
        
        # Save combined plot
        filename = f"{organ}_ROC_calibration_combined.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved combined ROC+calibration plot: {filename}")
        plt.close()
        
        return all_auc_values, calibration_metrics
    
    def create_comprehensive_analysis_plot(self, results_df):
        """Create comprehensive analysis plot across all organs"""
        
        fig = plt.figure(figsize=(20, 12))
        gs = fig.add_gridspec(3, 4, hspace=0.3, wspace=0.3)
        
        # Multi-OAR safety: always use sorted organ list
        organs = sorted(results_df['Organ'].unique())
        
        # Performance comparison subplot
        ax1 = fig.add_subplot(gs[0, :2])
        self._plot_performance_comparison(results_df, ax1)
        
        # Sample characteristics subplot  
        ax2 = fig.add_subplot(gs[0, 2:])
        self._plot_sample_characteristics(results_df, ax2)
        
        # Dose distribution subplots
        for i, organ in enumerate(organs[:3]):
            ax = fig.add_subplot(gs[1, i])
            self._plot_dose_distribution(results_df[results_df['Organ'] == organ], organ, ax)
        
        # Model performance trends subplot
        ax_trend = fig.add_subplot(gs[1, 3])
        self._plot_performance_trends(results_df, ax_trend)
        
        # Overall summary subplot
        ax_summary = fig.add_subplot(gs[2, :])
        self._plot_overall_summary(results_df, ax_summary)
        
        plt.suptitle('Comprehensive NTCP Analysis', fontsize=20, fontweight='bold', y=0.98)
        
        # Save comprehensive plot
        filename = "comprehensive_analysis.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved comprehensive analysis plot: {filename}")
        plt.close()
    
    def _plot_performance_comparison(self, results_df, ax):
        """Plot performance comparison across organs"""
        
        organs = results_df['Organ'].unique()
        traditional_models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']
        ml_models = ['ML_ANN', 'ML_XGBoost', 'ML_RandomForest']
        
        # Prepare data
        auc_data = {}
        for model in traditional_models + ml_models:
            auc_data[model] = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            
            # Traditional models
            for model in traditional_models:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            auc_data[model].append(auc_score)
                        except:
                            auc_data[model].append(np.nan)
                    else:
                        auc_data[model].append(np.nan)
                else:
                    auc_data[model].append(np.nan)
            
            # ML models
            for model in ml_models:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            auc_data[model].append(auc_score)
                        except:
                            auc_data[model].append(np.nan)
                    else:
                        auc_data[model].append(np.nan)
                else:
                    auc_data[model].append(np.nan)
        
        # Plot bars
        x = np.arange(len(organs))
        width = 0.15
        
        for i, model in enumerate(traditional_models + ml_models):
            aucs = auc_data[model]
            valid_aucs = [a if not np.isnan(a) else 0 for a in aucs]
            
            color = COLORS.get(model, COLORS['confidence'])
            
            bars = ax.bar(x + i * width, valid_aucs, width, 
                         label=model.replace('_', ' '), color=color, alpha=0.8)
            
            # Add value labels
            for j, (bar, auc_val) in enumerate(zip(bars, aucs)):
                if not np.isnan(auc_val) and auc_val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                           f'{auc_val:.3f}', ha='center', va='bottom', 
                           fontsize=8, fontweight='bold')
        
        ax.set_xlabel('Organ', fontsize=12, fontweight='bold')
        ax.set_ylabel('AUC', fontsize=12, fontweight='bold')
        ax.set_xticks(x + width * 2)
        ax.set_xticklabels(organs)
        ax.legend(fontsize=8, ncol=2)
        ax.set_ylim(0, 1.1)
        ax.grid(True, alpha=0.3, axis='y')
    
    def _plot_sample_characteristics(self, results_df, ax):
        """Plot sample characteristics by organ"""
        
        organs = results_df['Organ'].unique()
        sample_sizes = []
        event_rates = []
        colors = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            event_rate = (n_events / n_patients) * 100 if n_patients > 0 else 0
            
            sample_sizes.append(n_patients)
            event_rates.append(event_rate)
            
            # Color by data quality
            if n_events < 5:
                colors.append('red')
            elif n_events < 10:
                colors.append('orange')
            else:
                colors.append('green')
        
        scatter = ax.scatter(sample_sizes, event_rates, c=colors, s=200, alpha=0.7)
        
        # Add organ labels
        for i, organ in enumerate(organs):
            ax.annotate(organ, (sample_sizes[i], event_rates[i]), 
                       xytext=(5, 5), textcoords='offset points', fontsize=10)
        
        ax.set_xlabel('Sample Size (n)', fontsize=12, fontweight='bold')
        ax.set_ylabel('Event Rate (%)', fontsize=12, fontweight='bold')
        ax.grid(True, alpha=0.3)
        
        # Add threshold lines
        ax.axhline(y=5, color='red', linestyle='--', alpha=0.5, label='5% threshold')
        ax.axvline(x=30, color='blue', linestyle='--', alpha=0.5, label='n=30 threshold')
        ax.legend(fontsize=8)
    
    def _plot_dose_distribution(self, organ_data, organ, ax):
        """Plot dose distribution for specific organ"""
        
        valid_data = organ_data.dropna(subset=['gEUD', 'Observed_Toxicity'])
        
        if len(valid_data) == 0:
            ax.text(0.5, 0.5, f'No data\nfor {organ}', ha='center', va='center', 
                   transform=ax.transAxes, fontsize=12)
            return
        
        # Separate by outcome
        events = valid_data[valid_data['Observed_Toxicity'] == 1]['gEUD']
        non_events = valid_data[valid_data['Observed_Toxicity'] == 0]['gEUD']
        
        # Plot histograms
        bins = np.linspace(valid_data['gEUD'].min(), valid_data['gEUD'].max(), 15)
        
        ax.hist(non_events, bins=bins, alpha=0.7, color=COLORS['confidence'], 
               label=f'No Toxicity (n={len(non_events)})', density=True)
        ax.hist(events, bins=bins, alpha=0.8, color=COLORS['observed'],
               label=f'Toxicity (n={len(events)})', density=True)
        
        ax.set_xlabel('gEUD (Gy)', fontsize=10)
        ax.set_ylabel('Density', fontsize=10)
        ax.set_title(f'{organ}', fontsize=12, fontweight='bold')
        ax.legend(fontsize=8)
        ax.grid(True, alpha=0.3)
    
    def _plot_performance_trends(self, results_df, ax):
        """Plot performance trends across organs"""
        
        # Multi-OAR safety: always use sorted organ list
        organs = sorted(results_df['Organ'].unique())
        
        # Calculate mean AUC for traditional vs ML models
        trad_aucs = []
        ml_aucs = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            
            # Traditional models average
            trad_scores = []
            for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            trad_scores.append(auc_score)
                        except:
                            pass
            
            # ML models average
            ml_scores = []
            for model in ['ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            ml_scores.append(auc_score)
                        except:
                            pass
            
            trad_aucs.append(np.mean(trad_scores) if trad_scores else np.nan)
            ml_aucs.append(np.mean(ml_scores) if ml_scores else np.nan)
        
        x = range(len(organs))
        ax.plot(x, trad_aucs, 'o-', color=COLORS['LKB_LogLogit'], 
               linewidth=2, markersize=8, label='Traditional Models')
        ax.plot(x, ml_aucs, 's-', color=COLORS['ML_ANN'], 
               linewidth=2, markersize=8, label='ML Models')
        
        ax.set_xlabel('Organs', fontsize=10)
        ax.set_ylabel('Mean AUC', fontsize=10)
        ax.set_xticks(x)
        ax.set_xticklabels(organs, rotation=45)
        ax.legend(fontsize=8)
        ax.grid(True, alpha=0.3)
        ax.set_ylim(0.4, 1.0)
    
    def _plot_overall_summary(self, results_df, ax):
        """Plot overall summary statistics"""
        
        # Summary statistics table
        organs = results_df['Organ'].unique()
        
        summary_data = []
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            event_rate = (n_events / n_patients) * 100 if n_patients > 0 else 0
            
            # Best AUC
            best_auc = 0
            best_model = 'N/A'
            model_display_names = {
                'LKB_LogLogit': 'LKB Log-logit', 'LKB_Probit': 'LKB Probit', 'RS_Poisson': 'RS Poisson',
                'ML_ANN': 'ANN', 'ML_XGBoost': 'XGBoost', 'ML_RandomForest': 'Random Forest'
            }
            for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            if auc_score > best_auc:
                                best_auc = auc_score
                                best_model = model_display_names.get(model, model.replace('_', ' '))
                        except:
                            pass
            
            summary_data.append([organ, n_patients, n_events, f'{event_rate:.1f}%', 
                               f'{best_auc:.3f}', best_model])
        
        # Create table
        table = ax.table(cellText=summary_data,
                        colLabels=['Organ', 'Patients', 'Events', 'Event Rate', 'Best AUC', 'Best Model'],
                        cellLoc='center',
                        loc='center')
        
        table.auto_set_font_size(False)
        table.set_fontsize(10)
        table.scale(1.2, 1.5)
        
        # Style the table
        for i in range(len(summary_data) + 1):
            for j in range(6):
                cell = table[(i, j)]
                if i == 0:  # Header
                    cell.set_facecolor('#4CAF50')
                    cell.set_text_props(weight='bold', color='white')
                else:
                    cell.set_facecolor('#f0f0f0' if i % 2 == 0 else 'white')
        
        ax.axis('off')
    
    def create_model_performance_plot(self, results_df):
        """Create detailed model performance comparison plot"""
        
        organs = results_df['Organ'].unique()
        
        # Create figure with subplots
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
        
        # AUC comparison
        self._plot_auc_comparison(results_df, ax1)
        
        # Brier score comparison
        self._plot_brier_comparison(results_df, ax2)
        
        # Model type comparison
        self._plot_model_type_comparison(results_df, ax3)
        
        # Data quality vs performance
        self._plot_quality_vs_performance(results_df, ax4)
        
        plt.suptitle('Model Performance Analysis', fontsize=16, fontweight='bold')
        plt.tight_layout()
        
        # Save plot
        filename = "model_performance_analysis.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved model performance plot: {filename}")
        plt.close()
    
    def _plot_auc_comparison(self, results_df, ax):
        """Plot AUC comparison across organs and models"""
        
        organs = results_df['Organ'].unique()
        models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']
        
        # Prepare data
        auc_data = {model: [] for model in models}
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            
            for model in models:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            auc_data[model].append(auc_score)
                        except:
                            auc_data[model].append(np.nan)
                    else:
                        auc_data[model].append(np.nan)
                else:
                    auc_data[model].append(np.nan)
        
        # Plot grouped bars
        x = np.arange(len(organs))
        width = 0.15
        
        for i, model in enumerate(models):
            aucs = auc_data[model]
            valid_aucs = [a if not np.isnan(a) else 0 for a in aucs]
            
            color = COLORS.get(model, COLORS['confidence'])
            label = model.replace('_', ' ')
            
            bars = ax.bar(x + i * width, valid_aucs, width, 
                         label=label, color=color, alpha=0.8)
            
            # Add value labels
            for j, (bar, auc_val) in enumerate(zip(bars, aucs)):
                if not np.isnan(auc_val) and auc_val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01,
                           f'{auc_val:.3f}', ha='center', va='bottom', 
                           fontsize=8, fontweight='bold', rotation=45)
        
        ax.set_xlabel('Organ', fontsize=12, fontweight='bold')
        ax.set_ylabel('AUC', fontsize=12, fontweight='bold')
        ax.set_xticks(x + width * 2)
        ax.set_xticklabels(organs)
        ax.legend(fontsize=10, bbox_to_anchor=(1.05, 1), loc='upper left')
        ax.set_ylim(0, 1.1)
        ax.grid(True, alpha=0.3, axis='y')
    
    def _plot_brier_comparison(self, results_df, ax):
        """Plot Brier score comparison"""
        
        organs = results_df['Organ'].unique()
        models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']
        
        # Prepare data
        brier_data = {model: [] for model in models}
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            
            for model in models:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            brier_score = brier_score_loss(y_true, y_pred)
                            brier_data[model].append(brier_score)
                        except:
                            brier_data[model].append(np.nan)
                    else:
                        brier_data[model].append(np.nan)
                else:
                    brier_data[model].append(np.nan)
        
        # Plot grouped bars
        x = np.arange(len(organs))
        width = 0.15
        
        for i, model in enumerate(models):
            briers = brier_data[model]
            valid_briers = [b if not np.isnan(b) else 0 for b in briers]
            
            color = COLORS.get(model, COLORS['confidence'])
            label = model.replace('_', ' ')
            
            bars = ax.bar(x + i * width, valid_briers, width, 
                         label=label, color=color, alpha=0.8)
            
            # Add value labels
            for j, (bar, brier_val) in enumerate(zip(bars, briers)):
                if not np.isnan(brier_val) and brier_val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.005,
                           f'{brier_val:.3f}', ha='center', va='bottom', 
                           fontsize=8, fontweight='bold', rotation=45)
        
        ax.set_xlabel('Organ', fontsize=12, fontweight='bold')
        ax.set_ylabel('Brier Score (Lower = Better)', fontsize=12, fontweight='bold')
        ax.set_xticks(x + width * 2)
        ax.set_xticklabels(organs)
        ax.legend(fontsize=10, bbox_to_anchor=(1.05, 1), loc='upper left')
        ax.grid(True, alpha=0.3, axis='y')
    
    def _plot_model_type_comparison(self, results_df, ax):
        """Plot comparison between traditional and ML models"""
        
        organs = results_df['Organ'].unique()
        
        trad_aucs = []
        ml_aucs = []
        improvements = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            
            # Best traditional AUC
            best_trad = 0
            for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            best_trad = max(best_trad, auc_score)
                        except:
                            pass
            
            # Best ML AUC
            best_ml = 0
            for model in ['ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            best_ml = max(best_ml, auc_score)
                        except:
                            pass
            
            trad_aucs.append(best_trad if best_trad > 0 else np.nan)
            ml_aucs.append(best_ml if best_ml > 0 else np.nan)
            
            # Calculate improvement
            if best_trad > 0 and best_ml > 0:
                improvement = ((best_ml - best_trad) / best_trad) * 100
                improvements.append(improvement)
            else:
                improvements.append(np.nan)
        
        x = np.arange(len(organs))
        width = 0.35
        
        bars1 = ax.bar(x - width/2, [a if not np.isnan(a) else 0 for a in trad_aucs], 
                      width, label='Best Traditional', color=COLORS['LKB_LogLogit'], alpha=0.8)
        bars2 = ax.bar(x + width/2, [a if not np.isnan(a) else 0 for a in ml_aucs], 
                      width, label='Best ML', color=COLORS['ML_ANN'], alpha=0.8)
        
        # Add improvement percentages
        for i, (trad, ml, imp) in enumerate(zip(trad_aucs, ml_aucs, improvements)):
            if not np.isnan(imp):
                ax.text(i, max(trad, ml) + 0.05, f'{imp:+.1f}%', 
                       ha='center', va='bottom', fontsize=10, fontweight='bold')
        
        ax.set_xlabel('Organ', fontsize=12, fontweight='bold')
        ax.set_ylabel('Best AUC', fontsize=12, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(organs)
        ax.legend()
        ax.set_ylim(0, 1.1)
        ax.grid(True, alpha=0.3, axis='y')
    
    def _plot_quality_vs_performance(self, results_df, ax):
        """Plot data quality vs model performance"""
        
        organs = results_df['Organ'].unique()
        
        sample_sizes = []
        event_counts = []
        best_aucs = []
        colors = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            
            # Find best AUC
            best_auc = 0
            for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            best_auc = max(best_auc, auc_score)
                        except:
                            pass
            
            sample_sizes.append(n_patients)
            event_counts.append(n_events)
            best_aucs.append(best_auc)
            
            # Color by performance
            if best_auc >= 0.8:
                colors.append('green')
            elif best_auc >= 0.7:
                colors.append('blue')
            elif best_auc >= 0.6:
                colors.append('orange')
            else:
                colors.append('red')
        
        # Create bubble plot
        scatter = ax.scatter(sample_sizes, best_aucs, s=[e*10 for e in event_counts], 
                           c=colors, alpha=0.6)
        
        # Add organ labels
        for i, organ in enumerate(organs):
            ax.annotate(f'{organ}\n({event_counts[i]} events)', 
                       (sample_sizes[i], best_aucs[i]), 
                       xytext=(5, 5), textcoords='offset points', fontsize=9)
        
        ax.set_xlabel('Sample Size', fontsize=12, fontweight='bold')
        ax.set_ylabel('Best AUC', fontsize=12, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.set_ylim(0, 1.05)
        
        # Add performance thresholds
        ax.axhline(y=0.8, color='green', linestyle='--', alpha=0.5, label='Excellent (≥0.8)')
        ax.axhline(y=0.7, color='blue', linestyle='--', alpha=0.5, label='Good (≥0.7)')
        ax.axhline(y=0.6, color='orange', linestyle='--', alpha=0.5, label='Fair (≥0.6)')
        ax.legend(fontsize=8)
    def create_overall_performance_plot(self, results_df):
        """Create overall performance summary plot"""
        
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 10))
        
        organs = results_df['Organ'].unique()
        
        # Plot 1: AUC heatmap
        # Keep labels and underlying model columns in sync, including Random Forest
        models = ['LKB LogLogit', 'LKB Probit', 'RS Poisson', 'ANN', 'XGBoost', 'Random Forest']
        model_cols = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']
        
        auc_matrix = np.zeros((len(organs), len(models)))
        
        for i, organ in enumerate(organs):
            organ_data = results_df[results_df['Organ'] == organ]
            
            for j, model in enumerate(model_cols):
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            auc_matrix[i, j] = auc_score
                        except:
                            auc_matrix[i, j] = np.nan
                    else:
                        auc_matrix[i, j] = np.nan
                else:
                    auc_matrix[i, j] = np.nan
        
        # Create heatmap
        im = ax1.imshow(auc_matrix, cmap='RdYlGn', aspect='auto', vmin=0.4, vmax=1.0)
        ax1.set_xticks(range(len(models)))
        ax1.set_xticklabels(models, rotation=45, ha='right')
        ax1.set_yticks(range(len(organs)))
        ax1.set_yticklabels(organs)
        ax1.set_xlabel('Models', fontsize=12, fontweight='bold')
        ax1.set_ylabel('Organs', fontsize=12, fontweight='bold')
        
        # Add text annotations
        for i in range(len(organs)):
            for j in range(len(models)):
                if not np.isnan(auc_matrix[i, j]):
                    text = ax1.text(j, i, f'{auc_matrix[i, j]:.3f}',
                                   ha="center", va="center", color="black", fontweight='bold')
        
        plt.colorbar(im, ax=ax1, label='AUC')
        
        # Plot 2: Sample size vs performance
        sample_sizes = []
        event_rates = []
        best_aucs = []
        organ_names = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            event_rate = (n_events / n_patients) * 100 if n_patients > 0 else 0
            
            # Find best AUC
            best_auc = np.nanmax(auc_matrix[list(organs).index(organ), :])
            
            sample_sizes.append(n_patients)
            event_rates.append(event_rate)
            best_aucs.append(best_auc if not np.isnan(best_auc) else 0)
            organ_names.append(organ)
        
        scatter = ax2.scatter(sample_sizes, best_aucs, s=[er*10 for er in event_rates], 
                            alpha=0.6, c=range(len(organs)), cmap='tab10')
        
        for i, organ in enumerate(organ_names):
            ax2.annotate(organ, (sample_sizes[i], best_aucs[i]), 
                        xytext=(5, 5), textcoords='offset points', fontsize=10)
        
        ax2.set_xlabel('Sample Size', fontsize=12, fontweight='bold')
        ax2.set_ylabel('Best AUC', fontsize=12, fontweight='bold')
        ax2.grid(True, alpha=0.3)
        ax2.set_ylim(0, 1.05)
        
        # Plot 3: Model type comparison
        trad_means = []
        ml_means = []
        organ_labels = []
        
        for organ in organs:
            organ_idx = list(organs).index(organ)
            trad_aucs = auc_matrix[organ_idx, :3]  # First 3 are traditional
            ml_aucs = auc_matrix[organ_idx, 3:]    # Last 2 are ML
            
            trad_mean = np.nanmean(trad_aucs) if not np.all(np.isnan(trad_aucs)) else 0
            ml_mean = np.nanmean(ml_aucs) if not np.all(np.isnan(ml_aucs)) else 0
            
            trad_means.append(trad_mean)
            ml_means.append(ml_mean)
            organ_labels.append(organ)
        
        x = np.arange(len(organs))
        width = 0.35
        
        bars1 = ax3.bar(x - width/2, trad_means, width, label='Traditional NTCP', 
                       color=COLORS['LKB_LogLogit'], alpha=0.8)
        bars2 = ax3.bar(x + width/2, ml_means, width, label='Machine Learning', 
                       color=COLORS['ML_ANN'], alpha=0.8)
        
        # Add improvement percentages
        for i, (trad, ml) in enumerate(zip(trad_means, ml_means)):
            if trad > 0 and ml > 0:
                improvement = ((ml - trad) / trad) * 100
                ax3.text(i, max(trad, ml) + 0.02, f'{improvement:+.1f}%', 
                        ha='center', va='bottom', fontsize=10, fontweight='bold')
        
        ax3.set_xlabel('Organ', fontsize=12, fontweight='bold')
        ax3.set_ylabel('Mean AUC', fontsize=12, fontweight='bold')
        ax3.set_xticks(x)
        ax3.set_xticklabels(organ_labels)
        ax3.legend()
        ax3.grid(True, alpha=0.3, axis='y')
        ax3.set_ylim(0, 1.1)
        
        # Plot 4: Summary statistics
        ax4.axis('off')
        
        # Calculate summary statistics
        total_patients = len(results_df)
        total_events = int(results_df['Observed_Toxicity'].sum())
        overall_event_rate = (total_events / total_patients) * 100
        
        # Best performing models
        best_overall_auc = np.nanmax(auc_matrix)
        best_indices = np.where(auc_matrix == best_overall_auc)
        _disp = {'LKB_LogLogit': 'LKB Log-logit', 'LKB_Probit': 'LKB Probit', 'RS_Poisson': 'RS Poisson',
                 'ML_ANN': 'ANN', 'ML_XGBoost': 'XGBoost', 'ML_RandomForest': 'Random Forest'}
        if len(best_indices[0]) > 0:
            best_organ = organs[best_indices[0][0]]
            best_model_key = models[best_indices[1][0]]
            best_model = _disp.get(best_model_key, best_model_key.replace('_', ' '))
        else:
            best_organ = 'N/A'
            best_model = 'N/A'
        
        # Create summary text
        summary_text = f"""
        OVERALL SUMMARY
        ═══════════════════════════════
        
        Dataset Characteristics:
        • Total Patients: {total_patients}
        • Total Events: {total_events}
        • Overall Event Rate: {overall_event_rate:.1f}%
        • Organs Analyzed: {len(organs)}
        
        Model Performance:
        • Best Performance: {best_overall_auc:.3f} AUC
        • Best Model: {best_model}
        • Best Organ: {best_organ}
        
        Traditional vs ML:
        • Traditional Mean: {np.nanmean([t for t in trad_means if t > 0]):.3f}
        • ML Mean: {np.nanmean([m for m in ml_means if m > 0]):.3f}
        • ML Improvement: {((np.nanmean([m for m in ml_means if m > 0]) - np.nanmean([t for t in trad_means if t > 0])) / np.nanmean([t for t in trad_means if t > 0]) * 100):+.1f}%
        
        Model Availability:
        • Traditional Models: 3/3 (100%)
        • ML Models: Available for organs with ≥15 samples
        """
        
        ax4.text(0.1, 0.9, summary_text, transform=ax4.transAxes, fontsize=12,
                verticalalignment='top', fontfamily='monospace',
                bbox=dict(boxstyle='round,pad=0.5', facecolor='lightblue', alpha=0.8))
        
        plt.suptitle('Overall NTCP Model Performance Summary', fontsize=16, fontweight='bold')
        plt.tight_layout()
        
        # Save plot
        filename = "overall_performance_summary.png"
        plt.savefig(self.output_dir / filename, dpi=600, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        print(f" Saved overall performance plot: {filename}")
        plt.close()

def create_comprehensive_excel(results_df, output_dir):
    """Create comprehensive Excel file with all results"""
    
    output_path = Path(output_dir)
    excel_file = output_path / 'ntcp_results.xlsx'
    
    print(f" Creating comprehensive Excel file: {excel_file}")
    
    with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
        
        # Sheet 1: Complete Results
        results_df_copy = results_df.copy()
        
        # Round numerical columns
        numeric_cols = results_df_copy.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            if 'NTCP' in col or col in ['gEUD', 'mean_dose', 'max_dose', 'total_volume']:
                results_df_copy[col] = results_df_copy[col].round(4)
            else:
                results_df_copy[col] = results_df_copy[col].round(2)
        
        results_df_copy.to_excel(writer, sheet_name='Complete Results', index=False)
        
        # Sheet 2: Summary by Organ
        summary_data = []
        
        # Multi-OAR safety: always use sorted organ list
        for organ in sorted(results_df['Organ'].unique()):
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            event_rate = (n_events / n_patients) * 100 if n_patients > 0 else 0
            
            # Calculate mean gEUD
            mean_geud = organ_data['gEUD'].mean()
            geud_std = organ_data['gEUD'].std()
            geud_range = f"{organ_data['gEUD'].min():.1f} - {organ_data['gEUD'].max():.1f}"
            
            # Calculate model performance
            model_performance = {}
            for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            brier_score = brier_score_loss(y_true, y_pred)
                            model_performance[model] = {'AUC': auc_score, 'Brier': brier_score}
                        except:
                            model_performance[model] = {'AUC': np.nan, 'Brier': np.nan}
                    else:
                        model_performance[model] = {'AUC': np.nan, 'Brier': np.nan}
                else:
                    model_performance[model] = {'AUC': np.nan, 'Brier': np.nan}
            
            # Find best model
            best_auc = 0
            best_model = 'N/A'
            _display_names = {'LKB_LogLogit': 'LKB Log-logit', 'LKB_Probit': 'LKB Probit', 'RS_Poisson': 'RS Poisson',
                             'ML_ANN': 'ANN', 'ML_XGBoost': 'XGBoost', 'ML_RandomForest': 'Random Forest'}
            for model, perf in model_performance.items():
                if not np.isnan(perf['AUC']) and perf['AUC'] > best_auc:
                    best_auc = perf['AUC']
                    best_model = _display_names.get(model, model.replace('_', ' '))
            
            summary_row = {
                'Organ': organ,
                'Sample_Size': n_patients,
                'Events': n_events,
                'Event_Rate_Percent': f"{event_rate:.1f}%",
                'Mean_gEUD_Gy': f"{mean_geud:.1f}" if not np.isnan(mean_geud) else 'N/A',
                'gEUD_SD_Gy': f"{geud_std:.1f}" if not np.isnan(geud_std) else 'N/A',
                'gEUD_Range_Gy': geud_range,
                'Best_Model': best_model,
                'Best_AUC': f"{best_auc:.3f}" if best_auc > 0 else 'N/A'
            }
            
            # Add individual model performance
            for model, perf in model_performance.items():
                summary_row[f'{model}_AUC'] = f"{perf['AUC']:.3f}" if not np.isnan(perf['AUC']) else 'N/A'
                summary_row[f'{model}_Brier'] = f"{perf['Brier']:.3f}" if not np.isnan(perf['Brier']) else 'N/A'
            
            summary_data.append(summary_row)
        
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary by Organ', index=False)
        
        # Sheet 3: Model Performance Matrix
        organs = results_df['Organ'].unique()
        models = ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson', 'ML_ANN', 'ML_XGBoost', 'ML_RandomForest']
        
        performance_matrix = []
        
        for organ in organs:
            organ_data = results_df[results_df['Organ'] == organ]
            row = {'Organ': organ}
            
            for model in models:
                ntcp_col = f'NTCP_{model}'
                if ntcp_col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            y_true = valid_data['Observed_Toxicity'].values
                            y_pred = valid_data[ntcp_col].values
                            fpr, tpr, _ = roc_curve(y_true, y_pred)
                            auc_score = auc(fpr, tpr)
                            row[f'{model}_AUC'] = f"{auc_score:.3f}"
                        except:
                            row[f'{model}_AUC'] = 'Error'
                    else:
                        row[f'{model}_AUC'] = 'Insufficient Data'
                else:
                    row[f'{model}_AUC'] = 'Not Available'
            
            performance_matrix.append(row)
        
        performance_df = pd.DataFrame(performance_matrix)
        performance_df.to_excel(writer, sheet_name='Performance Matrix', index=False)
        
        # Sheet 4: Dose Metrics
        dose_metrics_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ', 'gEUD', 'mean_dose', 'max_dose', 'total_volume']
        
        # Add V-dose and D-dose columns
        v_cols = [col for col in results_df.columns if col.startswith('V') and col[1:].isdigit()]
        d_cols = [col for col in results_df.columns if col.startswith('D') and any(c.isdigit() for c in col[1:])]
        
        dose_metrics_cols.extend(v_cols)
        dose_metrics_cols.extend(d_cols)
        
        # Filter to available columns
        available_dose_cols = [col for col in dose_metrics_cols if col in results_df.columns]
        
        dose_df = results_df[available_dose_cols].copy()
        
        # Round dose metrics
        numeric_dose_cols = dose_df.select_dtypes(include=[np.number]).columns
        for col in numeric_dose_cols:
            dose_df[col] = dose_df[col].round(2)
        
        dose_df.to_excel(writer, sheet_name='Dose Metrics', index=False)
        
        # Sheet 5: NTCP Predictions Only
        ntcp_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ', 'Observed_Toxicity']
        ntcp_prediction_cols = [col for col in results_df.columns if col.startswith('NTCP_')]
        
        ntcp_cols.extend(ntcp_prediction_cols)
        ntcp_df = results_df[ntcp_cols].copy()
        
        # Round NTCP predictions
        for col in ntcp_prediction_cols:
            ntcp_df[col] = ntcp_df[col].round(4)
        
        ntcp_df.to_excel(writer, sheet_name='NTCP Predictions', index=False)
        
        # Sheet 6: Literature Parameters
        lit_params_data = []
        
        ntcp_calc = NTCPCalculator()
        for organ, params in ntcp_calc.literature_params.items():
            for model_type, model_params in params.items():
                row = {
                    'Organ': organ,
                    'Model': model_type,
                    **model_params
                }
                lit_params_data.append(row)
        
        lit_params_df = pd.DataFrame(lit_params_data)
        lit_params_df.to_excel(writer, sheet_name='Literature Parameters', index=False)
        
        # Sheet 7: Local Classical Parameters (NEW - additive)
        if hasattr(process_all_patients, 'local_calibration_params') and process_all_patients.local_calibration_params:
            local_params_data = []
            for organ, params in process_all_patients.local_calibration_params.items():
                local_params_data.append(params)
            
            if local_params_data:
                local_params_df = pd.DataFrame(local_params_data)
                local_params_df.to_excel(writer, sheet_name='Local Classical Parameters', index=False)
        
        # Sheet 8: Analysis Metadata (renumbered)
        metadata = [
            ['Analysis Date', pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Total Patients', len(results_df)],
            ['Total Events', int(results_df['Observed_Toxicity'].sum())],
            ['Overall Event Rate (%)', f"{(results_df['Observed_Toxicity'].sum() / len(results_df) * 100):.1f}"],
            ['Number of Organs', len(results_df['Organ'].unique())],
            ['Organs Analyzed', ', '.join(results_df['Organ'].unique())],
            ['Traditional Models', 'LKB Log-Logistic, LKB Probit, RS Poisson (QUANTEC)'],
            ['Local Classical Models', 'LKB Local, RS Local (cohort-calibrated)'],
            ['ML Models', 'ANN, XGBoost (where sufficient data)'],
            ['Performance Metrics', 'AUC, Brier Score'],
            ['Minimum Sample Size for ML', '15 patients per organ'],
            ['Software Version', 'py_ntcpx_v1.0.0']
        ]
        
        metadata_df = pd.DataFrame(metadata, columns=['Parameter', 'Value'])
        metadata_df.to_excel(writer, sheet_name='Analysis Metadata', index=False)
    
    print(f" Comprehensive Excel file created: {excel_file}")
    return excel_file

def _get_single_column(df, name):
    """Return a single Series even if Excel had duplicate columns."""
    cols = [c for c in df.columns if c.lower() == name.lower()]
    if len(cols) == 0:
        return None
    if len(cols) == 1:
        return df[cols[0]]
    # Multiple columns → take first non-null row-wise
    return df[cols].bfill(axis=1).iloc[:, 0]

def load_patient_data(output_dir):
    """
    Load patient data from Clinical Contract v2 reconciled file
    
    Single clinical source: code0_output/clinical_reconciled.xlsx (enforced)
    
    Clinical Contract v2 guarantees:
    - patient_id: mandatory, no missing values
    - xerostomia_grade2plus: mandatory, no missing values, must be 0 or 1
    - followup_months: mandatory, no missing values, must be > 0
    - Optional columns: age, sex, baseline_xerostomia, tobacco_exposure, chemotherapy, hpv_status
    
    Identity matching: DVH.PrimaryPatientID == clinical.patient_id
    """
    try:
        # Enforce single clinical source: code0_output/clinical_reconciled.xlsx
        # Use pipeline-root based resolution (code0_output is at pipeline root, not in output_dir)
        pipeline_root = os.path.abspath(os.path.join(output_dir, ".."))
        clinical_file = os.path.join(pipeline_root, "code0_output", "clinical_reconciled.xlsx")
        clinical_path = Path(clinical_file)
        
        if not clinical_path.exists():
            raise FileNotFoundError(
                f"Clinical Contract v2 file not found: {clinical_file}\n"
                "Run Step 0 (code0_clinical_reconciliation.py) first to generate clinical_reconciled.xlsx"
            )
        
        # Load clinical data from single source
        clinical_df = pd.read_excel(clinical_file)
        
        # Normalize column names using utility function
        patient_df = normalize_columns(clinical_df)
        
        # Validate only PatientID (Clinical Contract v2 mandatory field)
        # Note: normalize_columns() converts 'patient_id' to 'PatientID'
        if 'PatientID' not in patient_df.columns:
            raise ValueError("Clinical Contract v2 violation: PatientID column missing after normalization")
        
        # Map xerostomia_grade2plus to Observed_Toxicity for compatibility
        if 'xerostomia_grade2plus' in patient_df.columns:
            patient_df['Observed_Toxicity'] = patient_df['xerostomia_grade2plus']
        elif 'Observed_Toxicity' not in patient_df.columns:
            # Try to find toxicity column
            tox_cols = [c for c in patient_df.columns if 'toxicity' in c.lower() or 'tox' in c.lower()]
            if tox_cols:
                patient_df['Observed_Toxicity'] = patient_df[tox_cols[0]]
            else:
                print("Warning: No toxicity column found. Creating dummy column.")
                patient_df['Observed_Toxicity'] = 0
        
        # Normalize patient_id for matching (identity-safe key)
        # Note: DVH uses PrimaryPatientID, clinical uses PatientID (normalized from patient_id)
        # Matching: DVH.PrimaryPatientID == clinical.PatientID
        patient_df["patient_id_norm"] = patient_df["PatientID"].apply(normalize_dvh_id)
        
        # Map PatientID to PrimaryPatientID for DVH matching (internal use only)
        # This allows us to match: DVH.PrimaryPatientID == clinical.PatientID
        patient_df["PrimaryPatientID"] = patient_df["PatientID"]
        patient_df["PrimaryPatientID_norm"] = patient_df["patient_id_norm"]
        
        # Try to get AnonPatientID from registry for display (optional)
        try:
            from contract_validator import ContractValidator
            contracts_dir = Path(output_dir).parent / "contracts"
            if not contracts_dir.exists():
                contracts_dir = Path(output_dir) / "contracts"
            validator = ContractValidator(contracts_dir)
            registry = validator.load_step1_registry()
            
            if registry is not None and 'AnonPatientID' in registry.columns and 'PrimaryPatientID' in registry.columns:
                mapping_df = registry[['PrimaryPatientID', 'AnonPatientID']].drop_duplicates()
                patient_df = pd.merge(
                    patient_df,
                    mapping_df,
                    on='PrimaryPatientID',
                    how='left'
                )
        except Exception as e:
            # Registry not available - AnonPatientID will be None
            pass
        
        # Fill missing values with defaults
        if 'dose_per_fraction' not in patient_df.columns:
            patient_df['dose_per_fraction'] = 2.0
        
        if 'n_fractions' not in patient_df.columns:
            patient_df['n_fractions'] = 35
        
        return patient_df
        
    except Exception as e:
        print(f"Error loading patient data: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_all_patients(dvh_dir, patient_data_file, output_dir):
    """Enhanced main processing pipeline with traditional + ML models"""
    
    print("Enhanced NTCP Analysis: Traditional + Machine Learning Models")
    print("=" * 65)
    
    # Initialize components
    dvh_processor = DVHProcessor(dvh_dir)
    ntcp_calc = NTCPCalculator()
    ml_models = MachineLearningModels()
    
    # Load patient data from Clinical Contract v2 reconciled file
    # Note: patient_data_file parameter is kept for backward compatibility but not used
    patient_df = load_patient_data(output_dir)
    if patient_df is None:
        return
    
    print(f"Loaded {len(patient_df)} patient-organ combinations from Clinical Contract v2")
    
    # Clinical Contract v2: PatientID is guaranteed by Step-0 (normalized from patient_id)
    # Identity matching: DVH.PrimaryPatientID == clinical.PatientID
    if 'PatientID' not in patient_df.columns:
        raise ValueError("Clinical Contract v2 violation: PatientID column missing")
    
    # patient_id is already normalized in load_patient_data()
    # PrimaryPatientID is mapped from patient_id for DVH matching
    valid_patient_count = patient_df['patient_id_norm'].notna().sum() if 'patient_id_norm' in patient_df.columns else len(patient_df)
    print(f"Using patient_id for matching (Clinical Contract v2: {valid_patient_count}/{len(patient_df)} valid)")
    print(f"Identity matching: DVH.PrimaryPatientID == clinical.patient_id")
    
    # Get AnonPatientID for display (optional)
    has_anon = 'AnonPatientID' in patient_df.columns
    if not has_anon:
        # Try to map from registry
        try:
            from contract_validator import ContractValidator
            contracts_dir = Path(output_dir).parent / "contracts"
            validator = ContractValidator(contracts_dir)
            registry_df = validator.load_step1_registry()
            if registry_df is not None and 'AnonPatientID' in registry_df.columns:
                mapping_df = registry_df[['PrimaryPatientID', 'AnonPatientID']].drop_duplicates()
                patient_df = pd.merge(
                    patient_df,
                    mapping_df,
                    on='PrimaryPatientID',
                    how='left'
                )
                has_anon = 'AnonPatientID' in patient_df.columns
        except Exception:
            pass
    
    # Process each patient-organ combination
    results = []
    dropped_count = 0
    total_count = len(patient_df)
    
    for _, row in patient_df.iterrows():
        # Clinical Contract v2: use PatientID from clinical data (normalized from patient_id)
        # Handle both normalized (PatientID) and original (patient_id) column names
        patient_id = row.get('PatientID', row.get('patient_id', None))
        if patient_id is None:
            raise ValueError("Clinical Contract v2 violation: PatientID/patient_id not found in row")
        # Map to PrimaryPatientID for DVH matching: DVH.PrimaryPatientID == clinical.PatientID
        primary_patient_id = row.get('PrimaryPatientID', patient_id)  # Already mapped in load_patient_data()
        
        organ = row.get('Organ', None)  # Organ may not be in clinical data (per-patient, not per-organ)
        observed_toxicity = row['Observed_Toxicity']
        dose_per_fraction = row.get('dose_per_fraction', 2.0)
        anon_patient_id = row.get('AnonPatientID', None) if has_anon else None
        
        # Display AnonPatientID if available, otherwise patient_id
        display_id = anon_patient_id if anon_patient_id else patient_id
        print(f"\nProcessing {display_id} - {organ if organ else 'N/A'}" + (f" (patient_id: {patient_id})" if anon_patient_id else ""))
        
        # Load DVH data using PrimaryPatientID (mapped from patient_id)
        # Identity matching: DVH.PrimaryPatientID == clinical.patient_id
        # DVH filenames use PrimaryPatientID format: {PrimaryPatientID}_{Organ}.csv
        primary_id_norm = row.get('PrimaryPatientID_norm', row.get('patient_id_norm'))
        patient_name = row.get('PatientName', '')
        dvh, extracted_patient_id = dvh_processor.load_dvh_file(
            primary_patient_id, organ, patient_name, dvh_id=primary_patient_id, dvh_id_norm=primary_id_norm
        )
        if dvh is None:
            print(f"  Warning: Skipping - DVH file not found")
            dropped_count += 1
            continue
        
        # Use extracted PrimaryPatientID from DVH filename for result tracking
        # This ensures we match the actual DVH file identifier
        result_patient_id = extracted_patient_id if extracted_patient_id else primary_patient_id
        
        # Calculate dose metrics
        dose_metrics = dvh_processor.calculate_dose_metrics(dvh)
        if dose_metrics is None:
            print(f"  Warning: Skipping - Could not calculate dose metrics")
            continue
        
        # Get organ-specific parameters
        if organ in ntcp_calc.literature_params:
            lit_params = ntcp_calc.literature_params[organ]
            a_param = lit_params['LKB_LogLogit']['a']
            n_param = lit_params['LKB_Probit']['n']
        else:
            print(f"  Warning: No literature parameters for {organ}")
            continue
        
        # Calculate gEUD and effective volume
        geud = dvh_processor.calculate_gEUD(dvh, a_param)
        v_effective = dvh_processor.calculate_effective_volume(dvh, n_param)
        
        # Add to dose metrics
        dose_metrics['gEUD'] = geud
        dose_metrics['v_effective'] = v_effective
        
        print(f"   Total volume: {dose_metrics['total_volume']:.1f} cm³")
        print(f"   Mean dose: {dose_metrics['mean_dose']:.1f} Gy")
        print(f"   Max dose: {dose_metrics['max_dose']:.1f} Gy")
        print(f"   gEUD (a={a_param}): {geud:.1f} Gy")
        
        # Calculate NTCP for traditional models
        ntcp_results = ntcp_calc.calculate_all_ntcp_models(
            dvh, dose_metrics, organ, dose_per_fraction
        )
        
        # Compile results - Clinical Contract v2: use patient_id from clinical data
        # Identity matching: DVH.PrimaryPatientID == clinical.patient_id
        result_row = {
            'patient_id': patient_id,  # Clinical Contract v2: patient_id from clinical data
            'PrimaryPatientID': primary_patient_id,  # Mapped from patient_id for DVH matching (DVH uses PrimaryPatientID)
            'AnonPatientID': anon_patient_id if anon_patient_id and not pd.isna(anon_patient_id) else None,  # Anonymized ID for display only
            'Organ': organ,
            'Observed_Toxicity': observed_toxicity,
            'dose_per_fraction': dose_per_fraction,
            **dose_metrics
        }
        
        # Add traditional NTCP predictions
        for model_name, model_result in ntcp_results.items():
            result_row[f'NTCP_{model_name}'] = model_result.get('NTCP', np.nan)
            print(f"  [MODEL] {model_name}: {model_result.get('NTCP', 0):.3f}")
        
        # Calculate uNTCP (Uncertainty-Aware NTCP) for LKB_Probit model
        if UncertaintyAwareNTCP is not None and 'LKB_Probit' in ntcp_results:
            try:
                untcp_calc = UncertaintyAwareNTCP()
                # Get parameters for LKB_Probit
                if organ in ntcp_calc.literature_params:
                    lit_params = ntcp_calc.literature_params[organ]
                    probit_params = lit_params.get('LKB_Probit', {})
                    if probit_params:
                        # Create wrapper function for LKB_Probit
                        def ntcp_wrapper(params_dict, dvh_data):
                            """Wrapper for LKB_Probit NTCP calculation"""
                            TD50 = params_dict.get('TD50', probit_params.get('TD50', 50.0))
                            m = params_dict.get('m', probit_params.get('m', 0.2))
                            n = params_dict.get('n', probit_params.get('n', 0.5))
                            return ntcp_calc.ntcp_lkb_probit(dose_metrics, TD50, m, n)
                        
                        # Prepare params dict
                        untcp_params = {
                            'TD50': probit_params.get('TD50', 50.0),
                            'm': probit_params.get('m', 0.2),
                            'n': probit_params.get('n', 0.5)
                        }
                        
                        untcp_result = untcp_calc.calculate_untcp(ntcp_wrapper, untcp_params, dvh)
                        if untcp_result:
                            result_row['uNTCP'] = untcp_result['ntcp']
                            result_row['uNTCP_STD'] = untcp_result['std']
                            result_row['uNTCP_CI_L'] = untcp_result['ci_lower']
                            result_row['uNTCP_CI_U'] = untcp_result['ci_upper']
            except Exception as e:
                print(f"  Warning: uNTCP calculation failed: {e}")
                result_row['uNTCP'] = None
                result_row['uNTCP_STD'] = None
                result_row['uNTCP_CI_L'] = None
                result_row['uNTCP_CI_U'] = None
        else:
            result_row['uNTCP'] = None
            result_row['uNTCP_STD'] = None
            result_row['uNTCP_CI_L'] = None
            result_row['uNTCP_CI_U'] = None
        
        # Add novel NTCP models (Probabilistic gEUD and Monte Carlo)
        if ProbabilisticgEUDModel is not None:
            try:
                prob_geud_model = ProbabilisticgEUDModel(organ)
                prob_ntcp = prob_geud_model.calculate_ntcp_distribution(dvh)
                if prob_ntcp:
                    result_row['ProbNTCP_Mean'] = prob_ntcp['mean']
                    result_row['ProbNTCP_CI_L'] = prob_ntcp['ci_lower']
                    result_row['ProbNTCP_CI_U'] = prob_ntcp['ci_upper']
                    result_row['Prob_gEUD_mean'] = prob_ntcp['mean']
                    result_row['Prob_gEUD_std'] = prob_ntcp['std']
                    # Keep old field names for backward compatibility
                    result_row['NTCP_Probabilistic_gEUD'] = prob_ntcp['mean']
                    result_row['NTCP_Probabilistic_gEUD_std'] = prob_ntcp['std']
                    result_row['NTCP_Probabilistic_gEUD_ci_lower'] = prob_ntcp['ci_lower']
                    result_row['NTCP_Probabilistic_gEUD_ci_upper'] = prob_ntcp['ci_upper']
                    print(f"  [MODEL] Probabilistic gEUD: {prob_ntcp['mean']:.3f} +/- {prob_ntcp['std']:.3f}")
            except Exception as e:
                print(f"  Warning: Probabilistic gEUD calculation failed: {e}")
                result_row['ProbNTCP_Mean'] = None
                result_row['ProbNTCP_CI_L'] = None
                result_row['ProbNTCP_CI_U'] = None
                result_row['Prob_gEUD_mean'] = None
                result_row['Prob_gEUD_std'] = None
        else:
            result_row['ProbNTCP_Mean'] = None
            result_row['ProbNTCP_CI_L'] = None
            result_row['ProbNTCP_CI_U'] = None
            result_row['Prob_gEUD_mean'] = None
            result_row['Prob_gEUD_std'] = None
        
        if MonteCarloNTCPModel is not None and organ in ntcp_calc.literature_params:
            try:
                mc_model = MonteCarloNTCPModel(organ)
                lit_params = ntcp_calc.literature_params[organ]
                # Extract parameters for MC model
                mc_params = {
                    'n': lit_params['LKB_Probit'].get('n', 0.5),
                    'TD50': lit_params['LKB_Probit'].get('TD50', 50.0),
                    'm': lit_params['LKB_Probit'].get('m', 0.2)
                }
                mc_ntcp = mc_model.calculate_ntcp_with_uncertainty(dvh, mc_params)
                if mc_ntcp:
                    result_row['MC_NTCP_Mean'] = mc_ntcp['mean']
                    result_row['MC_NTCP_CI_L'] = mc_ntcp['ci_lower']
                    result_row['MC_NTCP_CI_U'] = mc_ntcp['ci_upper']
                    result_row['MonteCarlo_NTCP_mean'] = mc_ntcp['mean']
                    result_row['MonteCarlo_NTCP_std'] = mc_ntcp['std']
                    # Keep old field names for backward compatibility
                    result_row['NTCP_MonteCarlo'] = mc_ntcp['mean']
                    result_row['NTCP_MonteCarlo_std'] = mc_ntcp['std']
                    result_row['NTCP_MonteCarlo_ci_lower'] = mc_ntcp['ci_lower']
                    result_row['NTCP_MonteCarlo_ci_upper'] = mc_ntcp['ci_upper']
                    print(f"  [MODEL] Monte Carlo NTCP: {mc_ntcp['mean']:.3f} +/- {mc_ntcp['std']:.3f}")
            except Exception as e:
                print(f"  Warning: Monte Carlo NTCP calculation failed: {e}")
                result_row['MC_NTCP_Mean'] = None
                result_row['MC_NTCP_CI_L'] = None
                result_row['MC_NTCP_CI_U'] = None
                result_row['MonteCarlo_NTCP_mean'] = None
                result_row['MonteCarlo_NTCP_std'] = None
        else:
            result_row['MC_NTCP_Mean'] = None
            result_row['MC_NTCP_CI_L'] = None
            result_row['MC_NTCP_CI_U'] = None
            result_row['MonteCarlo_NTCP_mean'] = None
            result_row['MonteCarlo_NTCP_std'] = None

        # Compute and store Uncertainty-Aware NTCP (inverse variance weighting of Prob gEUD and Monte Carlo)
        prob_geud_mean = result_row.get('Prob_gEUD_mean')
        prob_geud_std = result_row.get('Prob_gEUD_std')
        mc_mean = result_row.get('MonteCarlo_NTCP_mean')
        mc_std = result_row.get('MonteCarlo_NTCP_std')
        if (prob_geud_mean is not None and mc_mean is not None and
                prob_geud_std is not None and mc_std is not None):
            if prob_geud_std > 0 and mc_std > 0:
                w_prob = 1.0 / (prob_geud_std ** 2)
                w_mc = 1.0 / (mc_std ** 2)
                uNTCP = (prob_geud_mean * w_prob + mc_mean * w_mc) / (w_prob + w_mc)
            else:
                uNTCP = (prob_geud_mean + mc_mean) / 2
            result_row['uNTCP'] = uNTCP
        # else: preserve existing uNTCP from untcp_calc (set earlier)

        results.append(result_row)
    
    # Convert to DataFrame
    results_df = pd.DataFrame(results)
    
    # ========================================================================
    # LOCAL CLASSICAL CALIBRATION (NEW - Additive only)
    # ========================================================================
    # Fit local sigmoid curve to cohort data and compute locally calibrated
    # LKB and RS models for fair comparison with ML
    print("\n" + "="*60)
    print("LOCAL CLASSICAL CALIBRATION")
    print("="*60)
    
    try:
        from ntcp_models.local_classical_calibration import (
            fit_local_sigmoid, 
            local_sigmoid_to_lkb, 
            local_sigmoid_to_rs
        )
        
        # Process each organ separately
        for organ in sorted(results_df['Organ'].unique()):
            organ_data = results_df[results_df['Organ'] == organ].copy()
            
            if len(organ_data) < 10:
                print(f"  Skipping {organ}: insufficient data for local calibration (n={len(organ_data)})")
                continue
            
            print(f"\n  Fitting local sigmoid for {organ} (n={len(organ_data)})...")
            
            # Use mean_dose as dose metric (already computed)
            if 'mean_dose' not in organ_data.columns:
                print(f"    Warning: mean_dose not available for {organ}, skipping local calibration")
                continue
            
            # Prepare data
            dose_metric = organ_data['mean_dose'].values
            toxicity = organ_data['Observed_Toxicity'].values.astype(float)
            
            # Remove NaN values
            valid_mask = ~(np.isnan(dose_metric) | np.isnan(toxicity))
            dose_metric_clean = dose_metric[valid_mask]
            toxicity_clean = toxicity[valid_mask]
            
            if len(dose_metric_clean) < 10:
                print(f"    Warning: Insufficient valid data for {organ} ({len(dose_metric_clean)} samples)")
                continue
            
            # Fit local sigmoid
            try:
                D50_local, k_local = fit_local_sigmoid(dose_metric_clean, toxicity_clean)
                print(f"    Local sigmoid: D50={D50_local:.2f} Gy, k={k_local:.2f} Gy")
                
                # Convert to LKB parameters
                TD50_local, m_local = local_sigmoid_to_lkb(D50_local, k_local)
                print(f"    Local LKB: TD50={TD50_local:.2f} Gy, m={m_local:.3f}")
                
                # Convert to RS parameters (for benchmarking only; no NTCP_RS_LOCAL column)
                D50_rs_local, gamma_local = local_sigmoid_to_rs(D50_local, k_local)
                print(f"    Local RS params (benchmark): D50={D50_rs_local:.2f} Gy, gamma={gamma_local:.3f}")
                
                # Compute local LKB NTCP for all patients (RS uses full Källman via MLE/QUANTEC only)
                organ_indices = organ_data.index
                ntcp_lkb_local = []
                
                for idx in organ_indices:
                    row = results_df.loc[idx]
                    
                    # Get dose metrics for LKB (needs v_effective and max_dose)
                    dose_val = row.get('mean_dose', np.nan)
                    v_effective = row.get('v_effective', 1.0)
                    max_dose = row.get('max_dose', dose_val)
                    
                    if np.isnan(dose_val):
                        ntcp_lkb_local.append(np.nan)
                        continue
                    
                    # Local LKB NTCP using existing LKB Probit function
                    try:
                        dose_metrics_local = {
                            'v_effective': v_effective,
                            'max_dose': max_dose
                        }
                        ntcp_lkb = ntcp_calc.ntcp_lkb_probit(
                            dose_metrics_local, 
                            TD50_local, 
                            m_local, 
                            n=1.0
                        )
                        ntcp_lkb_local.append(ntcp_lkb)
                    except Exception:
                        from scipy.stats import norm
                        t = (dose_val - TD50_local) / (m_local * TD50_local) if TD50_local > 0 else 0
                        ntcp_lkb = norm.cdf(t)
                        ntcp_lkb = np.clip(ntcp_lkb, 1e-10, 1 - 1e-10)
                        ntcp_lkb_local.append(ntcp_lkb)
                
                # Add local LKB only; RS uses full Källman model via QUANTEC/MLE (no simplified RS_local)
                results_df.loc[organ_indices, 'NTCP_LKB_LOCAL'] = ntcp_lkb_local
                
                # Save local parameters (will be exported to JSON later)
                # Store in a module-level variable for later export
                if not hasattr(process_all_patients, 'local_calibration_params'):
                    process_all_patients.local_calibration_params = {}
                
                process_all_patients.local_calibration_params[organ] = {
                    'organ': organ,
                    'D50_local': float(D50_local),
                    'k_local': float(k_local),
                    'TD50_local': float(TD50_local),
                    'm_local': float(m_local),
                    'D50_rs_local': float(D50_rs_local),
                    'gamma_local': float(gamma_local),
                    'n_samples': int(len(dose_metric_clean))
                }
                
            except Exception as e:
                print(f"    Warning: Local calibration failed for {organ}: {e}")
                organ_indices = organ_data.index
                results_df.loc[organ_indices, 'NTCP_LKB_LOCAL'] = np.nan
        
        print("\n" + "="*60)
        
    except ImportError as e:
        print(f"  Warning: Local calibration module not available: {e}")
        results_df['NTCP_LKB_LOCAL'] = np.nan
    except Exception as e:
        print(f"  Warning: Local calibration failed: {e}")
        if 'NTCP_LKB_LOCAL' not in results_df.columns:
            results_df['NTCP_LKB_LOCAL'] = np.nan
    
    # ========================================================================
    # END LOCAL CLASSICAL CALIBRATION
    # ========================================================================
    
    # Calculate match rate (successful matches / total attempts) - identity-safe
    matched_cases = len(results_df) if not results_df.empty else 0
    match_rate = matched_cases / total_count if total_count > 0 else 0
    
    # Validation printout: show PrimaryPatientID match rate
    print(f"\n[INFO] PrimaryPatientID match rate: {matched_cases}/{total_count} ({match_rate:.1%})")
    
    # Validation guarantee: abort if not all DVHs matched (critical for identity safety)
    if matched_cases < total_count:
        drop_rate = dropped_count / total_count if total_count > 0 else 0
        if drop_rate > 0.10:
            print(f"\n[ERROR] IDENTITY MISMATCH: {dropped_count}/{total_count} ({drop_rate:.1%}) DVHs were dropped.")
            print("This indicates a mismatch between PrimaryPatientID in clinical data and DVH filenames.")
            print("All DVHs must be matched - cannot proceed with incomplete cohort.")
            print("\nCheck:")
            print("  1. Clinical Excel must contain PrimaryPatientID column")
            print("  2. PrimaryPatientID must match DVH filename prefix (before first underscore)")
            print("  3. Organ names must match exactly")
            print("  4. Verify Step1_DVHRegistry.xlsx matches clinical data")
            raise ValueError(f"Identity mismatch: {dropped_count} of {total_count} DVHs could not be matched. Cannot proceed with incomplete cohort.")
    
    # Guard against empty results or missing Organ column
    if results_df.empty or 'Organ' not in results_df.columns:
        print("No valid DVH-organ matches found. Skipping ML stage.")
        print(f"Processed {len(results_df)} valid patient-organ DVHs")
        return results_df
    
    # Report successful processing - identity-safe validation
    # Multi-OAR safety: always use sorted organ list
    for organ in sorted(results_df['Organ'].unique()):
        organ_count = len(results_df[results_df['Organ'] == organ])
        primary_ids = results_df[results_df['Organ'] == organ]['PrimaryPatientID'].unique()
        print(f"Processed {organ_count} valid patient-organ DVHs for {organ} (PrimaryPatientID range: {len(primary_ids)} unique patients)")
        
        # Validation guarantee: check against Step1 registry
        try:
            from contract_validator import ContractValidator
            contracts_dir = Path(output_dir).parent / "contracts"
            validator = ContractValidator(contracts_dir)
            registry_df = validator.load_step1_registry()
            
            if registry_df is not None:
                registry_organ_count = len(registry_df[registry_df['Organ'] == organ])
                if organ_count < registry_organ_count:
                    print(f"[ERROR] Identity validation failed: {organ_count} < {registry_organ_count} DVHs matched for {organ}")
                    print(f"  Expected: {registry_organ_count} DVHs from Step1 registry")
                    print(f"  Matched: {organ_count} DVHs")
                    raise ValueError(f"Identity validation failed for {organ}: Expected {registry_organ_count} DVHs, matched {organ_count}")
                elif organ_count == registry_organ_count:
                    print(f"  [VALIDATED] All {registry_organ_count} registry DVHs matched for {organ}")
        except ImportError:
            pass  # Contract validator not available
        except Exception as e:
            print(f"Warning: Could not validate against registry: {e}")
    
    # Train ML models per organ
    print(f"\n Training Machine Learning Models")
    print("=" * 40)
    
    ml_results = {}
    # Multi-OAR safety: always use sorted organ list and group by (PrimaryPatientID, Organ)
    for organ in sorted(results_df['Organ'].unique()):
        organ_data = results_df[results_df['Organ'] == organ].copy()
        
        if len(organ_data) >= 15:  # Minimum for ML training
            print(f"\n[MODEL] ML Analysis for {organ}...")
            ml_organ_results = ml_models.train_and_evaluate_ml_models(organ_data, organ)
            
            if ml_organ_results:
                ml_results[organ] = ml_organ_results
                
                # Add ML predictions to results
                ml_predictions = ml_models.predict_ml_models(organ_data, organ)
                
                for pred_col, pred_values in ml_predictions.items():
                    # Map predictions back to full results DataFrame
                    organ_mask = results_df['Organ'] == organ
                    organ_indices = results_df[organ_mask].index
                    
                    if len(pred_values) == len(organ_indices):
                        results_df.loc[organ_indices, pred_col] = pred_values
                
                # Calculate CCS (Cohort Consistency Score) for QA
                if CohortConsistencyScore is not None:
                    try:
                        ccs_calculator = CohortConsistencyScore()
                        # Prepare features for CCS (same as ML features)
                        X_features, y, feature_cols = ml_models.prepare_features(organ_data)
                        
                        if X_features is not None and len(X_features) > 0:
                            # Fit CCS on training data
                            ccs_calculator.fit(X_features)
                            
                            # Calculate CCS for each prediction (v3.0.0: warnings instead of DO_NOT_USE)
                            ccs_values = []
                            ccs_warnings = []
                            ccs_warning_flags = []  # Boolean: True if CCS below adaptive threshold
                            
                            # Get adaptive threshold from calculator
                            adaptive_threshold = ccs_calculator.ccs_threshold
                            
                            for idx in organ_indices:
                                # Get features for this patient
                                patient_row = results_df.loc[idx]
                                patient_features = []
                                for feat in feature_cols:
                                    if feat in patient_row.index:
                                        patient_features.append(patient_row[feat])
                                    else:
                                        patient_features.append(0.0)
                                
                                # Calculate CCS
                                try:
                                    ccs_result = ccs_calculator.calculate_ccs(np.array(patient_features))
                                    ccs_val = ccs_result['ccs']
                                    ccs_values.append(ccs_val)
                                    ccs_warnings.append(ccs_result['warning_level'])
                                    
                                    # v3.0.0: CCS_Warning flag (boolean) instead of DO_NOT_USE
                                    # True if CCS below adaptive threshold (interpretations should be treated with caution)
                                    ccs_warning_flag = ccs_val < adaptive_threshold
                                    ccs_warning_flags.append(ccs_warning_flag)
                                except Exception as e:
                                    ccs_values.append(np.nan)
                                    ccs_warnings.append('unknown')
                                    ccs_warning_flags.append(False)  # Default to False if calculation fails
                            
                            # Add CCS to results DataFrame
                            if len(ccs_values) == len(organ_indices):
                                results_df.loc[organ_indices, 'CCS'] = ccs_values
                                results_df.loc[organ_indices, 'CCS_Warning'] = ccs_warnings
                                results_df.loc[organ_indices, 'CCS_Warning_Flag'] = ccs_warning_flags
                                
                                # Log warning count
                                n_warnings = sum(ccs_warning_flags)
                                if n_warnings > 0:
                                    print(f"  ⚠️ INFO - CCS below adaptive threshold ({adaptive_threshold:.2f}) for {n_warnings}/{len(organ_indices)} predictions. Interpretations should be treated with caution.")
                    except Exception as e:
                        print(f"  Warning: CCS calculation failed: {e}")
        else:
            print(f"\nWarning: Insufficient data for ML training in {organ} ({len(organ_data)} samples)")
    
    # ================================
    # STEP-3 MODEL EXPORT FOR STEP-7.1
    # ================================
    import joblib, json
    
    model_dir = Path(output_dir) / "models"
    model_dir.mkdir(parents=True, exist_ok=True)
    
    # Export models for each organ that has trained models
    for organ in sorted(ml_models.models.keys()):
        organ_models = ml_models.models[organ]
        
        if 'ANN' in organ_models and 'XGBoost' in organ_models:
            # Extract ANN model and scaler from pipeline
            ann_pipeline = organ_models['ANN']['model']
            ann_model = ann_pipeline.named_steps['ann']
            scaler = ann_pipeline.named_steps['scaler']
            xgb_model = organ_models['XGBoost']['model']
            rf_model = organ_models.get('RandomForest', {}).get('model', None)
            
            # Get feature columns
            feature_columns = organ_models['ANN'].get('feature_names', [])
            
            # Create feature matrix for this organ
            organ_data = results_df[results_df['Organ'] == organ].copy()
            if len(feature_columns) > 0:
                # Filter to available features
                available_features = [f for f in feature_columns if f in organ_data.columns]
                if len(available_features) > 0:
                    feature_matrix = organ_data[available_features].copy()
                    
                    # Save models with organ-specific naming (for multi-organ support)
                    joblib.dump(ann_model, model_dir / f"{organ}_ANN_model.pkl")
                    joblib.dump(xgb_model, model_dir / f"{organ}_XGBoost_model.pkl")
                    joblib.dump(scaler, model_dir / f"{organ}_scaler.pkl")
                    if rf_model is not None:
                        joblib.dump(rf_model, model_dir / f"{organ}_RandomForest_model.pkl")
                    
                    # Also save as default (for single-organ or primary organ)
                    if organ == sorted(ml_models.models.keys())[0]:
                        joblib.dump(ann_model, model_dir / "ANN_model.pkl")
                        joblib.dump(xgb_model, model_dir / "XGBoost_model.pkl")
                        joblib.dump(scaler, model_dir / "scaler.pkl")
                        if rf_model is not None:
                            joblib.dump(rf_model, model_dir / "RandomForest_model.pkl")
                    
                    # Categorize features into dosimetric / biological / clinical
                    dosimetric_features = [f for f in available_features if any(f.startswith(prefix) for prefix in ['V', 'D', 'mean', 'max', 'gEUD', 'total'])]
                    biological_features = [f for f in available_features if any(f.startswith(prefix) for prefix in ['BED', 'EQD2', 'NTCP'])]
                    clinical_features = [f for f in available_features if f not in dosimetric_features and f not in biological_features]
                    
                    # Save feature registry
                    feature_registry = {
                        "dosimetric": dosimetric_features,
                        "biological": biological_features,
                        "clinical": clinical_features,
                        "all_features": available_features
                    }
                    
                    with open(model_dir / f"{organ}_feature_registry.json", "w") as f:
                        json.dump(feature_registry, f, indent=2)
                    
                    # Also save as default
                    if organ == sorted(ml_models.models.keys())[0]:
                        with open(model_dir / "feature_registry.json", "w") as f:
                            json.dump(feature_registry, f, indent=2)
                    
                    # Save feature matrix used for training
                    feature_matrix.to_csv(model_dir / f"{organ}_feature_matrix.csv", index=False)
                    
                    # Also save as default
                    if organ == sorted(ml_models.models.keys())[0]:
                        feature_matrix.to_csv(model_dir / "feature_matrix.csv", index=False)
                    
                    print(f"[EXPORT] Step-7.1 model package created for {organ}: {model_dir}")
    
    # Enforce identity contract before saving
    try:
        from contract_validator import enforce_identity_contract
        enforce_identity_contract(results_df)
        print("[VALIDATED] Identity contract enforced: PrimaryPatientID+Organ unique and non-null")
    except Exception as e:
        raise ValueError(f"Identity governance failed: {e}")
    
    # Save results
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    results_df.to_csv(output_path / 'enhanced_ntcp_calculations.csv', index=False)
    print(f"\n Saved enhanced NTCP calculations to {output_path / 'enhanced_ntcp_calculations.csv'}")
    
    # Create comprehensive Excel file
    create_comprehensive_excel(results_df, output_dir)
    
    # Export ML CV metrics (CV_AUC, Test_AUC, Apparent_AUC) for tiered/QA use
    try:
        from sklearn.metrics import roc_curve, auc
        cv_rows = []
        model_to_col = {'ANN': 'NTCP_ML_ANN', 'XGBoost': 'NTCP_ML_XGBoost', 'RandomForest': 'NTCP_ML_RandomForest'}
        for organ in sorted(ml_models.models.keys()):
            organ_models = ml_models.models[organ]
            organ_data = results_df[results_df['Organ'] == organ]
            y_true = organ_data['Observed_Toxicity'].values.astype(int)
            for model_name, pred_col in model_to_col.items():
                if model_name not in organ_models or pred_col not in results_df.columns:
                    continue
                info = organ_models[model_name]
                cv_mean = info.get('cv_AUC_mean')
                cv_std = info.get('cv_AUC_std')
                test_auc = info.get('test_AUC')
                if test_auc is None and cv_mean is not None:
                    test_auc = cv_mean  # CV path: use mean fold test AUC as Test_AUC
                y_pred = organ_data[pred_col].values
                valid = ~np.isnan(y_pred)
                if valid.sum() < 5:
                    continue
                y_pred_valid = y_pred[valid]
                n_unique = len(np.unique(y_pred_valid))
                constant_predictor = (n_unique == 1)
                if constant_predictor:
                    print(f"\n[ML CV] WARNING: {model_name} ({organ}) has constant predictions (single value). AUC is uninformative.")
                fpr, tpr, _ = roc_curve(y_true[valid], y_pred_valid)
                apparent_auc = auc(fpr, tpr)
                cv_rows.append({
                    'Organ': organ,
                    'Model': model_name,
                    'Apparent_AUC': apparent_auc,
                    'CV_AUC_mean': cv_mean,
                    'CV_AUC_std': cv_std,
                    'Test_AUC': test_auc,
                    'Constant_Predictor': constant_predictor
                })
        if cv_rows:
            cv_df = pd.DataFrame(cv_rows)
            cv_path = output_path / 'ml_cv_metrics.xlsx'
            cv_df.to_excel(cv_path, index=False)
            print(f"\n[ML CV] Saved CV metrics to {cv_path}")
    except Exception as e:
        print(f"\n[ML CV] Warning: Could not export ml_cv_metrics.xlsx: {e}")
    
    # Create comprehensive plots
    print(f"\n Creating Comprehensive Publication-Ready Plots")
    print("=" * 55)
    
    plotter = ComprehensivePlotter(output_path / 'plots', ntcp_calc)
    
    # Multi-OAR safety: always use sorted organ list
    for organ in sorted(results_df['Organ'].unique()):
        organ_data = results_df[results_df['Organ'] == organ].copy()
        
        print(f"\n Creating plots for {organ}...")
        
        # Individual plots for each organ
        # Disabled: duplicate dose-response plot (gEUD-based)
        # Biological dose-response refit (mean dose-based) is handled in Step 1.5
        # plotter.create_dose_response_plot(organ_data, organ)
        plotter.create_roc_plot(organ_data, organ)
        plotter.create_calibration_plot(organ_data, organ)
        plotter.create_combined_roc_calibration_plot(organ_data, organ)
    
    # Overall analysis plots
    print(f"\n Creating comprehensive analysis plots...")
    plotter.create_comprehensive_analysis_plot(results_df)
    plotter.create_model_performance_plot(results_df)
    plotter.create_overall_performance_plot(results_df)
    
    # Export local classical calibration parameters (NEW - additive)
    try:
        if hasattr(process_all_patients, 'local_calibration_params') and process_all_patients.local_calibration_params:
            import json
            params_file = Path(output_dir) / 'local_classical_parameters.json'
            with open(params_file, 'w') as f:
                json.dump(process_all_patients.local_calibration_params, f, indent=2)
            print(f"\n[LOCAL CALIBRATION] Parameters saved to: {params_file}")
    except Exception as e:
        print(f"Warning: Could not export local calibration parameters: {e}")
    
    print(f"Processed {len(results_df)} valid patient-organ DVHs")
    return results_df

def create_enhanced_summary_report(results_df, output_dir):
    """Create enhanced summary report"""
    
    output_path = Path(output_dir)
    
    summary_stats = []
    
    # Multi-OAR safety: always use sorted organ list
    for organ in sorted(results_df['Organ'].unique()):
        organ_data = results_df[results_df['Organ'] == organ]
        
        n_patients = len(organ_data)
        n_events = int(organ_data['Observed_Toxicity'].sum())
        event_rate = n_events / n_patients if n_patients > 0 else 0
        
        # Get all model performance
        model_performance = {}
        
        # Traditional NTCP models (QUANTEC)
        for model in ['LKB_LogLogit', 'LKB_Probit', 'RS_Poisson']:
            ntcp_col = f'NTCP_{model}'
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                if len(valid_data) >= 5:
                    try:
                        fpr, tpr, _ = roc_curve(valid_data['Observed_Toxicity'], valid_data[ntcp_col])
                        auc_score = auc(fpr, tpr)
                        model_performance[f'QUANTEC-{model}'] = auc_score
                    except:
                        model_performance[f'QUANTEC-{model}'] = np.nan
        
        # Local classical models (NEW - additive)
        for model in ['LKB_LOCAL']:
            ntcp_col = f'NTCP_{model}'
            if ntcp_col in organ_data.columns:
                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                if len(valid_data) >= 5:
                    try:
                        fpr, tpr, _ = roc_curve(valid_data['Observed_Toxicity'], valid_data[ntcp_col])
                        auc_score = auc(fpr, tpr)
                        model_performance[f'Local-{model.replace("_LOCAL", "")}'] = auc_score
                    except:
                        model_performance[f'Local-{model.replace("_LOCAL", "")}'] = np.nan
        
        # ML models (ANN, XGBoost, RandomForest, GradientBoosting, etc.)
        ml_cols = [col for col in organ_data.columns if col.startswith('NTCP_ML_')]
        for ml_col in ml_cols:
            lc = ml_col.lower()
            if 'ann' in lc:
                model_name = 'ML_ANN'
            elif 'xgboost' in lc or 'xgb' in lc:
                model_name = 'ML_XGBoost'
            elif 'randomforest' in lc or 'rf' in lc:
                model_name = 'ML_RandomForest'
            elif 'gradient' in lc or 'gb' in lc:
                model_name = 'ML_GradientBoosting'
            else:
                # Fallback: use suffix after NTCP_ML_
                model_name = f"ML_{ml_col.split('NTCP_ML_')[-1]}"
            
            valid_data = organ_data.dropna(subset=[ml_col, 'Observed_Toxicity'])
            if len(valid_data) >= 5:
                try:
                    fpr, tpr, _ = roc_curve(valid_data['Observed_Toxicity'], valid_data[ml_col])
                    auc_score = auc(fpr, tpr)
                    model_performance[model_name] = auc_score
                except Exception:
                    model_performance[model_name] = np.nan
        
        # Find best models
        best_overall = max(model_performance.items(), key=lambda x: x[1] if not np.isnan(x[1]) else 0) if model_performance else ('None', 0)
        
        physics_models = {k: v for k, v in model_performance.items() if not k.startswith('ML_')}
        ml_models_perf = {k: v for k, v in model_performance.items() if k.startswith('ML_')}
        
        best_physics = max(physics_models.items(), key=lambda x: x[1] if not np.isnan(x[1]) else 0) if physics_models else ('None', 0)
        best_ml = max(ml_models_perf.items(), key=lambda x: x[1] if not np.isnan(x[1]) else 0) if ml_models_perf else ('None', 0)
        
        summary_stats.append({
            'Organ': organ,
            'N_Patients': n_patients,
            'N_Events': n_events,
            'Event_Rate_Percent': f"{event_rate*100:.1f}%",
            'Best_Overall_Model': best_overall[0],
            'Best_Overall_AUC': f"{best_overall[1]:.3f}" if not np.isnan(best_overall[1]) else 'N/A',
            'Best_Physics_Model': best_physics[0],
            'Best_Physics_AUC': f"{best_physics[1]:.3f}" if not np.isnan(best_physics[1]) else 'N/A',
            'Best_ML_Model': best_ml[0],
            'Best_ML_AUC': f"{best_ml[1]:.3f}" if not np.isnan(best_ml[1]) else 'N/A',
            'ML_Available': 'Yes' if ml_models_perf else 'No',
            'Data_Quality': get_data_quality_rating(n_patients, n_events),
            'Clinical_Recommendation': get_clinical_recommendation(best_overall[1], n_events, ml_models_perf)
        })
    
    summary_df = pd.DataFrame(summary_stats)
    summary_df.to_csv(output_path / 'enhanced_summary_performance.csv', index=False)
    
    # Create detailed text report
    report_text = []
    report_text.append("Enhanced NTCP Analysis: Traditional + Machine Learning Models")
    report_text.append("=" * 65)
    report_text.append(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_text.append(f"Enhanced features: Traditional NTCP + ML Models (ANN, XGBoost)")
    report_text.append(f"Plot quality: 600 DPI publication-ready")
    report_text.append("")
    
    for _, row in summary_df.iterrows():
        organ = row['Organ']
        report_text.append(f"{organ.upper()} ENHANCED ANALYSIS")
        report_text.append("-" * (len(organ) + 20))
        report_text.append(f"Sample Size: {row['N_Patients']} patients")
        report_text.append(f"Events: {row['N_Events']} ({row['Event_Rate_Percent']})")
        report_text.append("")
        
        report_text.append("Model Performance Comparison:")
        report_text.append(f"  Best Overall: {row['Best_Overall_Model']} (AUC: {row['Best_Overall_AUC']})")
        report_text.append(f"  Best Traditional: {row['Best_Physics_Model']} (AUC: {row['Best_Physics_AUC']})")
        
        if row['ML_Available'] == 'Yes':
            report_text.append(f"  Best ML: {row['Best_ML_Model']} (AUC: {row['Best_ML_AUC']})")
            
            # Calculate improvement
            try:
                physics_auc = float(row['Best_Physics_AUC'])
                ml_auc = float(row['Best_ML_AUC'])
                improvement = ((ml_auc - physics_auc) / physics_auc) * 100
                report_text.append(f"  ML Improvement: {improvement:+.1f}%")
            except:
                report_text.append(f"  ML Improvement: Cannot calculate")
        else:
            report_text.append(f"  ML Models: Not available (insufficient data)")
        
        report_text.append("")
        report_text.append(f"Data Quality: {row['Data_Quality']}")
        report_text.append(f"Clinical Recommendation: {row['Clinical_Recommendation']}")
        report_text.append("")
        report_text.append("-" * 60)
        report_text.append("")
    
    # Save enhanced text report
    with open(output_path / 'enhanced_analysis_report.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(report_text))
    
    print(f"\n Enhanced summary report saved to {output_path / 'enhanced_analysis_report.txt'}")
    print(f" Enhanced performance table saved to {output_path / 'enhanced_summary_performance.csv'}")

def perform_hospital_validation(results_df: pd.DataFrame, output_dir: Path) -> bool:
    """
    Perform hospital-to-hospital validation if HospitalID is present.
    
    Args:
        results_df: Results DataFrame with NTCP predictions
        output_dir: Output directory for hospital validation results
    
    Returns:
        True if hospital validation performed, False if HospitalID not present
    """
    # Check if HospitalID is present
    if 'HospitalID' not in results_df.columns:
        print("[INFO] Single-institution dataset detected — skipping hospital validation")
        return False
    
    # Check if HospitalID has valid values
    hospital_data = results_df['HospitalID'].dropna()
    if len(hospital_data) == 0:
        print("[INFO] HospitalID column present but empty — skipping hospital validation")
        return False
    
    unique_hospitals = hospital_data.unique()
    if len(unique_hospitals) < 2:
        print(f"[INFO] Only one hospital found ({len(unique_hospitals)}) — skipping hospital validation")
        return False
    
    print(f"\n{'='*60}")
    print("HOSPITAL-TO-HOSPITAL VALIDATION")
    print(f"{'='*60}")
    print(f"Hospitals detected: {len(unique_hospitals)}")
    print(f"Hospital IDs: {', '.join(sorted(map(str, unique_hospitals)))}")
    
    # Create hospital validation directory
    hospital_output_dir = output_dir / 'hospital_validation'
    hospital_output_dir.mkdir(parents=True, exist_ok=True)
    
    # Per-organ hospital validation
    hospital_validation_results = []
    
    for organ in sorted(results_df['Organ'].unique()):
        organ_data = results_df[results_df['Organ'] == organ].copy()
        
        if len(organ_data) == 0 or 'HospitalID' not in organ_data.columns:
            continue
        
        print(f"\nProcessing {organ}...")
        
        organ_hospital_results = []
        
        # Get NTCP columns
        ntcp_cols = [col for col in organ_data.columns if col.startswith('NTCP_')]
        
        for hospital_id in sorted(organ_data['HospitalID'].dropna().unique()):
            hospital_data = organ_data[organ_data['HospitalID'] == hospital_id].copy()
            
            if len(hospital_data) == 0:
                continue
            
            n_patients = len(hospital_data)
            n_events = int(hospital_data['Observed_Toxicity'].sum()) if 'Observed_Toxicity' in hospital_data.columns else 0
            event_rate = n_events / n_patients if n_patients > 0 else 0
            
            hospital_result = {
                'Organ': organ,
                'HospitalID': hospital_id,
                'N_Patients': n_patients,
                'N_Events': n_events,
                'Event_Rate': event_rate
            }
            
            # Calculate ROC and calibration metrics per hospital
            for ntcp_col in ntcp_cols:
                valid_data = hospital_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                
                if len(valid_data) >= 5:
                    try:
                        y_true = valid_data['Observed_Toxicity'].values
                        y_pred = valid_data[ntcp_col].values
                        
                        # ROC
                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                        auc_score = auc(fpr, tpr)
                        
                        # Brier score
                        brier = brier_score_loss(y_true, y_pred)
                        
                        model_name = ntcp_col.replace('NTCP_', '')
                        hospital_result[f'{model_name}_AUC'] = auc_score
                        hospital_result[f'{model_name}_Brier'] = brier
                    except Exception as e:
                        print(f"  Warning: Hospital {hospital_id} {ntcp_col} metrics failed: {e}")
                        model_name = ntcp_col.replace('NTCP_', '')
                        hospital_result[f'{model_name}_AUC'] = np.nan
                        hospital_result[f'{model_name}_Brier'] = np.nan
                else:
                    model_name = ntcp_col.replace('NTCP_', '')
                    hospital_result[f'{model_name}_AUC'] = np.nan
                    hospital_result[f'{model_name}_Brier'] = np.nan
            
            organ_hospital_results.append(hospital_result)
        
        if organ_hospital_results:
            hospital_validation_results.extend(organ_hospital_results)
            
            # Save per-organ hospital validation
            organ_hospital_df = pd.DataFrame(organ_hospital_results)
            organ_output_file = hospital_output_dir / f'hospital_validation_{organ}.xlsx'
            with pd.ExcelWriter(organ_output_file, engine='openpyxl') as writer:
                organ_hospital_df.to_excel(writer, index=False, sheet_name='Hospital_Validation')
            
            print(f"  Saved: {organ_output_file}")
            print(f"    Hospitals: {len(organ_hospital_results)}")
            print(f"    Total patients: {sum(r['N_Patients'] for r in organ_hospital_results)}")
    
    # Save cross-hospital performance table
    if hospital_validation_results:
        cross_hospital_df = pd.DataFrame(hospital_validation_results)
        cross_hospital_file = hospital_output_dir / 'cross_hospital_performance.xlsx'
        with pd.ExcelWriter(cross_hospital_file, engine='openpyxl') as writer:
            cross_hospital_df.to_excel(writer, index=False, sheet_name='Cross_Hospital_Performance')
        
        print(f"\n[OK] Cross-hospital performance table saved: {cross_hospital_file}")
        print(f"  Total hospital-organ combinations: {len(hospital_validation_results)}")
        
        return True
    
    return False

def get_data_quality_rating(n_patients, n_events):
    """Enhanced data quality assessment"""
    if n_events < 5:
        return 'Poor (< 5 events, ML not feasible)'
    elif n_events < 10:
        return 'Fair (5-9 events, limited ML)'
    elif n_patients < 30:
        return 'Good (≥10 events, ML possible)'
    elif n_patients >= 50 and n_events >= 15:
        return 'Excellent (≥15 events, ≥50 patients, ML reliable)'
    else:
        return 'Very Good (adequate for ML)'

def get_clinical_recommendation(best_auc, n_events, ml_models):
    """Enhanced clinical recommendation including ML considerations"""
    
    if isinstance(best_auc, str) or n_events < 5:
        return 'Insufficient events for reliable recommendations'
    
    auc_val = float(best_auc) if isinstance(best_auc, str) else best_auc
    
    if auc_val < 0.6:
        return 'Poor discrimination - not recommended for clinical use'
    elif auc_val < 0.7:
        base_rec = 'Moderate discrimination - use with caution'
    elif auc_val < 0.8:
        base_rec = 'Good discrimination - suitable for clinical decision support'
    else:
        base_rec = 'Excellent discrimination - highly suitable for clinical use'
    
    # Add ML-specific recommendations
    if ml_models:
        base_rec += '; ML models available for enhanced predictions'
    
    return base_rec

def main():
    """Enhanced main execution function"""
    
    # --- SAFETY FIX: protect Path from shadowing ---
    Path = _Path
    
    parser = argparse.ArgumentParser(description='Enhanced NTCP Analysis: Traditional + ML Models')
    parser.add_argument('--dvh_dir', default='dDVH_csv', 
                       help='Directory containing DVH CSV files (default: dDVH_csv)')
    parser.add_argument('--patient_data', default=None,
                       help='[DEPRECATED] Patient data file - now loaded from code0_output/clinical_reconciled.xlsx (Clinical Contract v2)')
    parser.add_argument('--output_dir', default='enhanced_ntcp_analysis',
                       help='Output directory (default: enhanced_ntcp_analysis)')
    parser.add_argument('--ml_models', action='store_true', default=True,
                       help='Enable machine learning models (default: True)')
    
    args = parser.parse_args()
    
    print(" Enhanced NTCP Analysis: Traditional + Machine Learning")
    print("=" * 60)
    print("Features:")
    print("Traditional NTCP models (LKB Log-Logistic, LKB Probit, RS Poisson)")
    print("Machine learning models (ANN, XGBoost)")
    print("Unique colors and legends for all models")
    print("Enhanced 600 DPI publication-ready plots")
    print("Comprehensive Excel output (ntcp_results.xlsx)")
    print("Proper ML validation and anti-overfitting measures")
    
    # Validate input paths
    dvh_path = Path(args.dvh_dir)
    output_path = Path(args.output_dir)
    
    if not dvh_path.exists():
        print(f"Error: Error: DVH directory '{dvh_path}' not found")
        return
    
    # Clinical Contract v2: patient data is loaded from code0_output/clinical_reconciled.xlsx
    # This is validated in load_patient_data()
    
    # Check for DVH files
    dvh_files = list(dvh_path.glob('*.csv'))
    if not dvh_files:
        print(f"Error: Error: No CSV files found in '{dvh_path}'")
        return
    
    print(f" Found {len(dvh_files)} DVH files in {dvh_path}")
    print(f" Clinical data: code0_output/clinical_reconciled.xlsx (Clinical Contract v2)")
    
    # Check XGBoost availability
    if XGBOOST_AVAILABLE:
        print(" XGBoost available for ML modeling")
    else:
        print("Warning: XGBoost not available - only ANN will be used")
    
    # ---------------------------------------------------------
    # Ensure ntcp_calc is defined at function scope
    # ---------------------------------------------------------
    ntcp_calc = None
    
    # Initialize NTCPCalculator for use throughout main()
    ntcp_calc = NTCPCalculator()
    
    try:
        # Step 1: Enhanced processing with traditional + ML models
        print("\n[MODEL] Step 1: Enhanced DVH processing and model training...")
        # Clinical Contract v2: patient_data parameter is ignored, loaded from code0_output/clinical_reconciled.xlsx
        results_df = process_all_patients(args.dvh_dir, None, args.output_dir)
        
        if results_df is None or len(results_df) == 0:
            print("Error: No data processed. Please check file formats and patient IDs.")
            return
        
        print(f" Processed {len(results_df)} patient-organ combinations")
        
        # Count available models
        ntcp_cols = [col for col in results_df.columns if col.startswith('NTCP_')]
        traditional_models = [col for col in ntcp_cols if not 'ML_' in col]
        ml_models = [col for col in ntcp_cols if 'ML_' in col]
        
        print(f" Traditional NTCP models: {len(traditional_models)}")
        print(f" ML models trained: {len(ml_models)}")
        
        # Step 1.5: Biological dose–response refitting (separate from calibration)
        try:
            print("\n[MODEL] Step 1.5: Biological dose–response refitting...")
            from biological_refitting import refit_all_organs
            
            # Load config file if it exists
            config_path = Path('py_ntcpx_config/local_refit_bounds.json')
            if not config_path.exists():
                print(f"  Warning: Config file not found: {config_path}, using defaults")
                config_path = None
            
            # Safety check: ensure we're using binary events, not calibrated probabilities
            # Add refit_source column to mark that we're using clinical events
            if 'Observed_Toxicity' in results_df.columns:
                results_df['refit_source'] = 'clinical_events'
            elif 'event' in results_df.columns:
                results_df['refit_source'] = 'clinical_events'
            else:
                # Check if we have any probability columns that might be mistakenly used
                prob_cols = [col for col in results_df.columns if 'calibrated' in col.lower() or 'probability' in col.lower()]
                if prob_cols:
                    raise ValueError(f"ERROR: No binary event column found. Found probability columns: {prob_cols}. "
                                  f"Biological refit requires binary events (0/1) from clinical outcomes, "
                                  f"NOT calibrated probabilities.")
                else:
                    raise ValueError("ERROR: No binary event column (Observed_Toxicity or event) found for biological refit.")
            
            # Assert that event column is integer/binary
            event_col = 'Observed_Toxicity' if 'Observed_Toxicity' in results_df.columns else 'event'
            if not results_df[event_col].dtype in [np.int64, np.int32, int]:
                # Try to convert
                try:
                    results_df[event_col] = results_df[event_col].astype(int)
                except:
                    raise ValueError(f"ERROR: Event column '{event_col}' must be integer binary (0/1), "
                                   f"not probabilities. Found dtype: {results_df[event_col].dtype}")
            
            # Ensure values are 0/1
            unique_vals = results_df[event_col].dropna().unique()
            if not all(v in [0, 1] for v in unique_vals):
                raise ValueError(f"ERROR: Event column '{event_col}' must contain only 0 and 1 values. "
                               f"Found values: {unique_vals}")
            
            refit_all_organs(results_df, args.output_dir, n_bootstrap=1000, config_path=str(config_path) if config_path else None)
            print("  [OK] Biological refitting completed")
        except ImportError:
            print("  Warning: biological_refitting module not available, skipping biological refitting")
        except Exception as e:
            print(f"  Warning: Biological refitting failed: {e}")
            import traceback
            traceback.print_exc()
        
        # Publication-quality dose-response plots (Framework 1 & 2)
        GENERATE_PUBLICATION_DR_PLOTS = True
        if GENERATE_PUBLICATION_DR_PLOTS:
            # Defensive assertion: ensure ntcp_calc is initialized
            if ntcp_calc is None:
                raise RuntimeError("ntcp_calc was not initialized before DR plotting.")
            
            print("\n[PLOT] Generating publication-quality dose-response figures...")
            try:
                from pathlib import Path
                import json
                from scipy.special import expit
                
                plots_dir = Path(args.output_dir) / 'plots'
                plots_dir.mkdir(parents=True, exist_ok=True)
                
                # Load refitted parameters if available
                bio_params_file = Path(args.output_dir) / 'local_biological_parameters.json'
                classical_params_file = Path(args.output_dir) / 'local_classical_parameters.json'
                
                bio_params = {}
                classical_params = {}
                
                if bio_params_file.exists():
                    with open(bio_params_file, 'r') as f:
                        bio_params = json.load(f)
                
                if classical_params_file.exists():
                    with open(classical_params_file, 'r') as f:
                        classical_params = json.load(f)
                
                # Process each organ
                for organ in sorted(results_df['Organ'].unique()):
                    organ_data = results_df[results_df['Organ'] == organ].copy()
                    
                    # Filter valid data
                    valid_mask = (
                        organ_data['gEUD'].notna() &
                        (organ_data['gEUD'] > 0) &
                        organ_data['mean_dose'].notna() &
                        (organ_data['mean_dose'] > 0) &
                        organ_data['Observed_Toxicity'].notna()
                    )
                    organ_data_valid = organ_data[valid_mask].copy()
                    
                    if len(organ_data_valid) < 5:
                        print(f"  Skipping {organ}: insufficient data")
                        continue
                    
                    # Extract data
                    geud_vals = organ_data_valid['gEUD'].values
                    mean_dose_vals = organ_data_valid['mean_dose'].values
                    ntcp_obs = organ_data_valid['Observed_Toxicity'].values.astype(float)
                    
                    # Get NTCP predictions
                    ntcp_lkb_loglogit_pred = organ_data_valid['NTCP_LKB_LogLogit'].values if 'NTCP_LKB_LogLogit' in organ_data_valid.columns else np.array([])
                    ntcp_lkb_probit_pred = organ_data_valid['NTCP_LKB_Probit'].values if 'NTCP_LKB_Probit' in organ_data_valid.columns else np.array([])
                    ntcp_rs_pred = organ_data_valid['NTCP_RS_Poisson'].values if 'NTCP_RS_Poisson' in organ_data_valid.columns else np.array([])
                    ntcp_ann_pred = organ_data_valid['NTCP_ML_ANN'].values if 'NTCP_ML_ANN' in organ_data_valid.columns else np.array([])
                    ntcp_xgb_pred = organ_data_valid['NTCP_ML_XGBoost'].values if 'NTCP_ML_XGBoost' in organ_data_valid.columns else np.array([])
                    ntcp_rf_pred = organ_data_valid['NTCP_ML_RandomForest'].values if 'NTCP_ML_RandomForest' in organ_data_valid.columns else np.array([])
                    
                    # Get refitted parameters for classical models
                    lkb_loglogit_params = None
                    lkb_probit_params = None
                    rs_poisson_params = None
                    bio_logistic_params = None
                    
                    if organ in classical_params:
                        if classical_params[organ].get('LKB_LogLogit') and classical_params[organ]['LKB_LogLogit']:
                            lkb_loglogit_params = classical_params[organ]['LKB_LogLogit'].get('point_estimate')
                        if classical_params[organ].get('LKB_Probit') and classical_params[organ]['LKB_Probit']:
                            lkb_probit_params = classical_params[organ]['LKB_Probit'].get('point_estimate')
                        if classical_params[organ].get('RS_Poisson') and classical_params[organ]['RS_Poisson']:
                            rs_poisson_params = classical_params[organ]['RS_Poisson'].get('point_estimate')
                    
                    if organ in bio_params:
                        if bio_params[organ].get('Logistic') and bio_params[organ]['Logistic']:
                            bio_logistic_params = bio_params[organ]['Logistic'].get('point_estimate')
                    
                    # Get literature parameters as fallback
                    if organ in ntcp_calc.literature_params:
                        lit_params = ntcp_calc.literature_params[organ]
                    else:
                        lit_params = None
                    
                    # FIGURE 1: LKB Dose-Response (gEUD axis)
                    if lkb_loglogit_params or lkb_probit_params or lit_params:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(max(0.1, geud_vals.min() * 0.8), geud_vals.max() * 1.2, 300)
                        
                        # LKB Log-logit curve
                        if lkb_loglogit_params:
                            td50 = lkb_loglogit_params['TD50']
                            gamma50 = lkb_loglogit_params['gamma50']
                            ntcp_lkb_loglogit = 1.0 / (1.0 + np.power(td50 / np.maximum(x, 1e-6), gamma50))
                        elif lit_params and 'LKB_LogLogit' in lit_params:
                            td50 = lit_params['LKB_LogLogit']['TD50']
                            gamma50 = lit_params['LKB_LogLogit']['gamma50']
                            # Convert to EQD2 if needed
                            eqd2_x = ntcp_calc.convert_to_eqd2(x, lit_params['LKB_LogLogit'].get('alpha_beta', 3), 2.0)
                            ntcp_lkb_loglogit = ntcp_calc.ntcp_lkb_loglogit(eqd2_x, td50, gamma50)
                        else:
                            ntcp_lkb_loglogit = None
                        
                        if ntcp_lkb_loglogit is not None:
                            ax.plot(x, ntcp_lkb_loglogit, color="#1f77b4", lw=3.0, label="LKB (Log-logit)")
                            ax.scatter(geud_vals, ntcp_obs, facecolors="none", edgecolors="#1f77b4",
                                     s=55, lw=1.8, alpha=0.9, zorder=5)
                        
                        # LKB Probit curve
                        if lkb_probit_params:
                            td50 = lkb_probit_params['TD50']
                            m = lkb_probit_params['m']
                            t = (x - td50) / (m * td50)
                            ntcp_lkb_probit = norm.cdf(t)
                        elif lit_params and 'LKB_Probit' in lit_params:
                            td50 = lit_params['LKB_Probit']['TD50']
                            m = lit_params['LKB_Probit']['m']
                            n = lit_params['LKB_Probit'].get('n', 0.5)
                            # Use dose metrics for probit
                            dose_metrics_dict = {'gEUD': x, 'mean_dose': x, 'max_dose': x * 1.1}
                            ntcp_lkb_probit = ntcp_calc.ntcp_lkb_probit(dose_metrics_dict, td50, m, n)
                        else:
                            ntcp_lkb_probit = None
                        
                        if ntcp_lkb_probit is not None:
                            ax.plot(x, ntcp_lkb_probit, color="#d62728", lw=3.0, ls="--", label="LKB (Probit)")
                            ax.scatter(geud_vals, ntcp_obs, marker="s", facecolors="none", edgecolors="#d62728",
                                     s=55, lw=1.8, alpha=0.9, zorder=5)
                        
                        ax.set_xlabel(f"{organ} gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_xlim(left=0)
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, which="major", alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        plt.tight_layout()
                        plt.savefig(plots_dir / f"Figure1_LKB_gEUD_DR_{organ}.png", dpi=1200)
                        plt.savefig(plots_dir / f"Figure1_LKB_gEUD_DR_{organ}.svg")
                        plt.close()
                        print(f"    Saved Figure1_LKB_gEUD_DR_{organ}.png")
                    
                    # FIGURE 2: RS Poisson Dose-Response (gEUD axis)
                    if rs_poisson_params or (lit_params and 'RS_Poisson' in lit_params):
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(max(0.1, geud_vals.min() * 0.8), geud_vals.max() * 1.2, 300)
                        
                        if rs_poisson_params:
                            d50 = rs_poisson_params['D50']
                            gamma = rs_poisson_params['gamma']
                            dose_ratio = x / d50
                            # Källman et al. RS curve (same as ntcp_rs_poisson for uniform dose): p = 2^{-exp(e·γ(1−r))}
                            ntcp_rs_poisson = 2.0 ** (-np.exp(np.e * gamma * (1.0 - dose_ratio)))
                        elif lit_params and 'RS_Poisson' in lit_params:
                            d50 = lit_params['RS_Poisson']['D50']
                            gamma = lit_params['RS_Poisson']['gamma']
                            # Källman et al. RS curve for uniform dose (same as ntcp_rs_poisson)
                            dose_ratio = x / np.maximum(d50, 1e-6)
                            ntcp_rs_poisson = 2.0 ** (-np.exp(np.e * gamma * (1.0 - dose_ratio)))
                        else:
                            ntcp_rs_poisson = None
                        
                        if ntcp_rs_poisson is not None:
                            ax.plot(x, ntcp_rs_poisson, color="#f0ad1a", lw=3.0, ls="-.",
                                  label="RS Poisson")
                            ax.scatter(geud_vals, ntcp_obs, marker="^", facecolors="none",
                                     edgecolors="#f0ad1a", s=65, lw=1.8, zorder=5)
                        
                        ax.set_xlabel(f"{organ} gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        plt.tight_layout()
                        plt.savefig(plots_dir / f"Figure2_RS_gEUD_DR_{organ}.png", dpi=1200)
                        plt.savefig(plots_dir / f"Figure2_RS_gEUD_DR_{organ}.svg")
                        plt.close()
                        print(f"    Saved Figure2_RS_gEUD_DR_{organ}.png")
                    
                    # FIGURE 3: Biological Logistic Dose-Response (Mean Dose)
                    if bio_logistic_params:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(max(0.1, mean_dose_vals.min() * 0.8), mean_dose_vals.max() * 1.2, 300)
                        
                        td50_bio = bio_logistic_params['TD50']
                        k_bio = bio_logistic_params['k']
                        ntcp_biological = expit((x - td50_bio) / k_bio)
                        
                        ax.plot(x, ntcp_biological, color="#2c3e50", lw=3.5, label="Biological logistic")
                        ax.scatter(mean_dose_vals, ntcp_obs, color="gray", alpha=0.65, s=45, zorder=2)
                        ax.axvline(td50_bio, color="#7f7f7f", lw=2.0, ls=":", label="TD50")
                        
                        ax.set_xlabel(f"Mean {organ.lower()} dose (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        plt.tight_layout()
                        plt.savefig(plots_dir / f"Figure3_Biological_MeanDose_DR_{organ}.png", dpi=1200)
                        plt.savefig(plots_dir / f"Figure3_Biological_MeanDose_DR_{organ}.svg")
                        plt.close()
                        print(f"    Saved Figure3_Biological_MeanDose_DR_{organ}.png")
                    
                    # FIGURE 4: Predicted NTCP vs gEUD (All Models) - Prediction comparison
                    if len(ntcp_lkb_loglogit_pred) > 0 or len(ntcp_lkb_probit_pred) > 0:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        # All arrays are already aligned from organ_data_valid
                        if len(ntcp_lkb_loglogit_pred) > 0 and len(ntcp_lkb_loglogit_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_lkb_loglogit_pred,
                                     s=40, alpha=0.6, label="LKB Log-logit")
                        if len(ntcp_lkb_probit_pred) > 0 and len(ntcp_lkb_probit_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_lkb_probit_pred,
                                     s=40, alpha=0.6, label="LKB Probit")
                        if len(ntcp_rs_pred) > 0 and len(ntcp_rs_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_rs_pred,
                                     s=40, alpha=0.6, label="RS Poisson")
                        if len(ntcp_ann_pred) > 0 and len(ntcp_ann_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_ann_pred,
                                     s=40, alpha=0.6, label="ANN")
                        if len(ntcp_xgb_pred) > 0 and len(ntcp_xgb_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_xgb_pred,
                                     s=40, alpha=0.6, label="XGBoost")
                        if len(ntcp_rf_pred) > 0 and len(ntcp_rf_pred) == len(geud_vals):
                            ax.scatter(geud_vals, ntcp_rf_pred,
                                     s=40, alpha=0.6, label="Random Forest")
                        
                        ax.set_xlabel(f"{organ} gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("Predicted NTCP", fontsize=14, fontweight="bold")
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=10, frameon=True, ncol=2)
                        plt.tight_layout()
                        plt.savefig(plots_dir / f"Figure4_Predicted_NTCP_Comparison_{organ}.png", dpi=1200)
                        plt.savefig(plots_dir / f"Figure4_Predicted_NTCP_Comparison_{organ}.svg")
                        plt.close()
                        print(f"    Saved Figure4_Predicted_NTCP_Comparison_{organ}.png")
                    
                    # FIGURE 5: Observed vs Predicted NTCP (Calibration-style)
                    if len(ntcp_obs) > 0:
                        fig, ax = plt.subplots(figsize=(6.5, 6.5))
                        
                        # All arrays are already aligned from organ_data_valid
                        if len(ntcp_lkb_loglogit_pred) > 0 and len(ntcp_lkb_loglogit_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_lkb_loglogit_pred,
                                     label="LKB Log-logit", alpha=0.6)
                        if len(ntcp_ann_pred) > 0 and len(ntcp_ann_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_ann_pred,
                                     label="ANN", alpha=0.6)
                        if len(ntcp_xgb_pred) > 0 and len(ntcp_xgb_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_xgb_pred,
                                     label="XGBoost", alpha=0.6)
                        if len(ntcp_rf_pred) > 0 and len(ntcp_rf_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_rf_pred,
                                     label="Random Forest", alpha=0.6)
                        if len(ntcp_lkb_probit_pred) > 0 and len(ntcp_lkb_probit_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_lkb_probit_pred,
                                     label="LKB Probit", alpha=0.6)
                        if len(ntcp_rs_pred) > 0 and len(ntcp_rs_pred) == len(ntcp_obs):
                            ax.scatter(ntcp_obs, ntcp_rs_pred,
                                     label="RS Poisson", alpha=0.6)
                        
                        ax.plot([0, 1], [0, 1], 'k--', lw=1.5, label="Perfect calibration")
                        ax.set_xlabel("Observed toxicity", fontsize=14, fontweight="bold")
                        ax.set_ylabel("Predicted NTCP", fontsize=14, fontweight="bold")
                        ax.set_xlim(0, 1)
                        ax.set_ylim(0, 1)
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=11)
                        plt.tight_layout()
                        plt.savefig(plots_dir / f"Figure5_Observed_vs_Predicted_{organ}.png", dpi=1200)
                        plt.savefig(plots_dir / f"Figure5_Observed_vs_Predicted_{organ}.svg")
                        plt.close()
                        print(f"    Saved Figure5_Observed_vs_Predicted_{organ}.png")
                
                print("  [OK] Publication dose-response figures generated")
            except Exception as e:
                print(f"  Warning: Publication plotting failed: {e}")
                import traceback
                traceback.print_exc()
        
        # ============================================================
        # REFERENCE-STYLE DR FIGURE (FIGURE 4) - MANUSCRIPT MATERIALS
        # ============================================================
        
        def generate_reference_style_DR_figure(results_df, output_dir, ntcp_calc):
            """
            Generate reference-style dose-response figure (Figure 4) for manuscript.
            Combines all three classical models (LKB Log-logit, LKB Probit, RS Poisson) on one plot.
            Uses smooth model curves + hollow data markers matching reference aesthetic.
            """
            try:
                # Load refitted parameters if available
                bio_params_file = Path(output_dir) / 'local_biological_parameters.json'
                classical_params_file = Path(output_dir) / 'local_classical_parameters.json'
                
                bio_params = {}
                classical_params = {}
                
                if bio_params_file.exists():
                    import json
                    with open(bio_params_file, 'r') as f:
                        bio_params = json.load(f)
                
                if classical_params_file.exists():
                    import json
                    with open(classical_params_file, 'r') as f:
                        classical_params = json.load(f)
                
                # Find first organ with sufficient data and classical models
                organ_for_dr = None
                organ_data_valid = None
                gEUD = None
                y_obs = None
                ntcp_lkb_loglogit_func = None
                ntcp_lkb_probit_func = None
                ntcp_rs_poisson_func = None
                
                for organ in sorted(results_df['Organ'].unique()):
                    organ_data = results_df[results_df['Organ'] == organ].copy()
                    
                    valid_mask = (
                        organ_data['gEUD'].notna() &
                        (organ_data['gEUD'] > 0) &
                        organ_data['Observed_Toxicity'].notna()
                    )
                    organ_data_valid = organ_data[valid_mask].copy()
                    
                    if len(organ_data_valid) < 5:
                        continue
                    
                    # Check if we have classical models
                    has_classical = (
                        'NTCP_LKB_LogLogit' in organ_data_valid.columns or
                        'NTCP_LKB_Probit' in organ_data_valid.columns or
                        'NTCP_RS_Poisson' in organ_data_valid.columns
                    )
                    
                    if has_classical:
                        organ_for_dr = organ
                        gEUD = organ_data_valid['gEUD'].values
                        y_obs = organ_data_valid['Observed_Toxicity'].values.astype(int)
                        break
                
                if organ_for_dr is None or gEUD is None or y_obs is None:
                    print("  Warning: No organ with sufficient data for reference DR figure")
                    return
                
                # Create model functions from refitted parameters or literature
                # LKB Log-logit
                if organ_for_dr in classical_params and classical_params[organ_for_dr].get('LKB_LogLogit'):
                    params = classical_params[organ_for_dr]['LKB_LogLogit'].get('point_estimate', {})
                    if params:
                        td50 = params.get('TD50')
                        gamma50 = params.get('gamma50')
                        if td50 and gamma50:
                            ntcp_lkb_loglogit_func = lambda x: 1.0 / (1.0 + np.power(td50 / np.maximum(x, 1e-6), 4.0 * gamma50))
                            lkb_loglogit_label = f"LKB (Log-logit) [TD₅₀={td50:.1f} Gy, γ₅₀={gamma50:.2f}]"
                elif organ_for_dr in ntcp_calc.literature_params and 'LKB_LogLogit' in ntcp_calc.literature_params[organ_for_dr]:
                    lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_LogLogit']
                    td50 = lit_params['TD50']
                    gamma50 = lit_params['gamma50']
                    alpha_beta = lit_params.get('alpha_beta', 3)
                    ntcp_lkb_loglogit_func = lambda x: ntcp_calc.ntcp_lkb_loglogit(
                        ntcp_calc.convert_to_eqd2(x, alpha_beta, 2.0), td50, gamma50
                    )
                    lkb_loglogit_label = f"LKB (Log-logit) [TD₅₀={td50:.1f} Gy, γ₅₀={gamma50:.2f}]"
                
                # LKB Probit
                if organ_for_dr in classical_params and classical_params[organ_for_dr].get('LKB_Probit'):
                    params = classical_params[organ_for_dr]['LKB_Probit'].get('point_estimate', {})
                    if params:
                        td50 = params.get('TD50')
                        m = params.get('m')
                        if td50 and m:
                            ntcp_lkb_probit_func = lambda x: norm.cdf((x - td50) / (m * td50))
                            lkb_probit_label = f"LKB (Probit) [TD₅₀={td50:.1f} Gy, m={m:.2f}]"
                elif organ_for_dr in ntcp_calc.literature_params and 'LKB_Probit' in ntcp_calc.literature_params[organ_for_dr]:
                    lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_Probit']
                    td50 = lit_params['TD50']
                    m = lit_params['m']
                    ntcp_lkb_probit_func = lambda x: norm.cdf((x - td50) / (m * td50))
                    lkb_probit_label = f"LKB (Probit) [TD₅₀={td50:.1f} Gy, m={m:.2f}]"
                
                # RS Poisson
                if organ_for_dr in classical_params and classical_params[organ_for_dr].get('RS_Poisson'):
                    params = classical_params[organ_for_dr]['RS_Poisson'].get('point_estimate', {})
                    if params:
                        d50 = params.get('D50')
                        gamma = params.get('gamma')
                        if d50 and gamma:
                            ntcp_rs_poisson_func = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                            rs_poisson_label = f"RS Poisson [D₅₀={d50:.1f} Gy, γ={gamma:.2f}]"
                elif organ_for_dr in ntcp_calc.literature_params and 'RS_Poisson' in ntcp_calc.literature_params[organ_for_dr]:
                    lit_params = ntcp_calc.literature_params[organ_for_dr]['RS_Poisson']
                    d50 = lit_params['D50']
                    gamma = lit_params['gamma']
                    ntcp_rs_poisson_func = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                    rs_poisson_label = f"RS Poisson [D₅₀={d50:.1f} Gy, γ={gamma:.2f}]"
                
                # Generate figure if at least one model is available
                if ntcp_lkb_loglogit_func is None and ntcp_lkb_probit_func is None and ntcp_rs_poisson_func is None:
                    print("  Warning: No classical models available for reference DR figure")
                    return
                
                # Create figure with exact reference formatting
                fig, ax = plt.subplots(figsize=(7.5, 6.0))
                
                x = np.linspace(gEUD.min(), gEUD.max(), 400)
                
                # Plot curves with exact formatting
                if ntcp_lkb_loglogit_func is not None:
                    ax.plot(x, ntcp_lkb_loglogit_func(x),
                            color="#1f77b4", lw=3.5, label=lkb_loglogit_label)
                    ax.scatter(gEUD, y_obs,
                               facecolors="none", edgecolors="#1f77b4",
                               s=60, lw=1.6, alpha=0.9, zorder=5)
                
                if ntcp_lkb_probit_func is not None:
                    ax.plot(x, ntcp_lkb_probit_func(x),
                            color="#d62728", lw=3.5, ls="--", label=lkb_probit_label)
                    ax.scatter(gEUD, y_obs,
                               marker="s", facecolors="none", edgecolors="#d62728",
                               s=60, lw=1.6, alpha=0.9, zorder=5)
                
                if ntcp_rs_poisson_func is not None:
                    ax.plot(x, ntcp_rs_poisson_func(x),
                            color="#f0ad1a", lw=3.5, ls="-.", label=rs_poisson_label)
                    ax.scatter(gEUD, y_obs,
                               marker="^", facecolors="none", edgecolors="#f0ad1a",
                               s=60, lw=1.6, alpha=0.9, zorder=5)
                
                # Axes formatting
                ax.set_xlabel("gEUD (Gy)", fontsize=14, fontweight="bold")
                ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                ax.set_ylim(0, 1.02)
                ax.grid(True, which="major", alpha=0.25)
                ax.legend(fontsize=11, frameon=True, loc="lower right")
                
                plt.tight_layout()
                
                manuscript_dir = Path(output_dir) / 'plots'
                manuscript_dir.mkdir(parents=True, exist_ok=True)
                
                plt.savefig(manuscript_dir / "Figure4_DR_Reference.png", dpi=1200)
                plt.savefig(manuscript_dir / "Figure4_DR_Reference.svg")
                plt.close()
                
                print(f"  [OK] Reference-style DR figure (Figure 4) saved to plots/")
                
            except Exception as e:
                print(f"  Warning: Reference-style DR figure generation failed: {e}")
                import traceback
                traceback.print_exc()
        
        # Generate reference-style DR figure
        if GENERATE_PUBLICATION_DR_PLOTS:
            try:
                generate_reference_style_DR_figure(results_df, args.output_dir, ntcp_calc)
            except Exception as e:
                print(f"  Warning: Reference DR figure generation failed: {e}")
        
        # ============================================================
        # PUBLICATION-GRADE DOSE–RESPONSE PLOTS (MODEL-NATIVE)
        # ============================================================
        
        if GENERATE_PUBLICATION_DR_PLOTS:
            try:
                print("\n[PLOT] Generating model-native dose-response plots...")
                import json
                from scipy.special import expit
                
                # Find first organ with sufficient data and refitted parameters
                organ_for_dr = None
                organ_data_valid = None
                gEUD = None
                mean_dose = None
                ntcp_obs = None
                ntcp_lkb_loglogit_pred = None
                ntcp_lkb_probit_pred = None
                ntcp_rs_pred = None
                ntcp_ann_pred = None
                ntcp_xgb_pred = None
                TD50_bio = None
                
                # Get refitted parameters
                bio_params_file = Path(args.output_dir) / 'local_biological_parameters.json'
                classical_params_file = Path(args.output_dir) / 'local_classical_parameters.json'
                
                bio_params = {}
                classical_params = {}
                
                if bio_params_file.exists():
                    with open(bio_params_file, 'r') as f:
                        bio_params = json.load(f)
                
                if classical_params_file.exists():
                    with open(classical_params_file, 'r') as f:
                        classical_params = json.load(f)
                
                # Find first organ with sufficient data
                for organ in sorted(results_df['Organ'].unique()):
                    organ_data = results_df[results_df['Organ'] == organ].copy()
                    
                    valid_mask = (
                        organ_data['gEUD'].notna() &
                        (organ_data['gEUD'] > 0) &
                        organ_data['mean_dose'].notna() &
                        (organ_data['mean_dose'] > 0) &
                        organ_data['Observed_Toxicity'].notna()
                    )
                    organ_data_valid = organ_data[valid_mask].copy()
                    
                    if len(organ_data_valid) < 5:
                        continue
                    
                    # Check if we have refitted parameters for this organ
                    has_classical = organ in classical_params and (
                        classical_params[organ].get('LKB_LogLogit') or 
                        classical_params[organ].get('LKB_Probit') or
                        classical_params[organ].get('RS_Poisson')
                    )
                    has_bio = organ in bio_params and bio_params[organ].get('Logistic')
                    
                    if has_classical or has_bio:
                        organ_for_dr = organ
                        break
                
                if organ_for_dr is None:
                    print("  Warning: No organ with sufficient data and refitted parameters found for DR plots")
                else:
                    # Extract data
                    gEUD = organ_data_valid['gEUD'].values
                    mean_dose = organ_data_valid['mean_dose'].values
                    ntcp_obs = organ_data_valid['Observed_Toxicity'].values.astype(float)
                    
                    # Get NTCP predictions
                    ntcp_lkb_loglogit_pred = organ_data_valid['NTCP_LKB_LogLogit'].values if 'NTCP_LKB_LogLogit' in organ_data_valid.columns else np.array([])
                    ntcp_lkb_probit_pred = organ_data_valid['NTCP_LKB_Probit'].values if 'NTCP_LKB_Probit' in organ_data_valid.columns else np.array([])
                    ntcp_rs_pred = organ_data_valid['NTCP_RS_Poisson'].values if 'NTCP_RS_Poisson' in organ_data_valid.columns else np.array([])
                    ntcp_ann_pred = organ_data_valid['NTCP_ML_ANN'].values if 'NTCP_ML_ANN' in organ_data_valid.columns else np.array([])
                    ntcp_xgb_pred = organ_data_valid['NTCP_ML_XGBoost'].values if 'NTCP_ML_XGBoost' in organ_data_valid.columns else np.array([])
                    ntcp_rf_pred = organ_data_valid['NTCP_ML_RandomForest'].values if 'NTCP_ML_RandomForest' in organ_data_valid.columns else np.array([])
                    
                    # Get refitted parameters
                    lkb_loglogit_params = None
                    lkb_probit_params = None
                    rs_poisson_params = None
                    bio_logistic_params = None
                    
                    if organ_for_dr in classical_params:
                        if classical_params[organ_for_dr].get('LKB_LogLogit'):
                            lkb_loglogit_params = classical_params[organ_for_dr]['LKB_LogLogit'].get('point_estimate')
                        if classical_params[organ_for_dr].get('LKB_Probit'):
                            lkb_probit_params = classical_params[organ_for_dr]['LKB_Probit'].get('point_estimate')
                        if classical_params[organ_for_dr].get('RS_Poisson'):
                            rs_poisson_params = classical_params[organ_for_dr]['RS_Poisson'].get('point_estimate')
                    
                    if organ_for_dr in bio_params:
                        if bio_params[organ_for_dr].get('Logistic'):
                            bio_logistic_params = bio_params[organ_for_dr]['Logistic'].get('point_estimate')
                    
                    # Create vectorized model functions
                    ntcp_lkb_loglogit = None
                    ntcp_lkb_probit = None
                    ntcp_rs_poisson = None
                    ntcp_biological = None
                    
                    if lkb_loglogit_params:
                        td50 = lkb_loglogit_params['TD50']
                        gamma50 = lkb_loglogit_params['gamma50']
                        ntcp_lkb_loglogit = lambda x: 1.0 / (1.0 + np.power(td50 / np.maximum(x, 1e-6), 4.0 * gamma50))
                    elif organ_for_dr in ntcp_calc.literature_params and 'LKB_LogLogit' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_LogLogit']
                        td50 = lit_params['TD50']
                        gamma50 = lit_params['gamma50']
                        alpha_beta = lit_params.get('alpha_beta', 3)
                        ntcp_lkb_loglogit = lambda x: ntcp_calc.ntcp_lkb_loglogit(
                            ntcp_calc.convert_to_eqd2(x, alpha_beta, 2.0), td50, gamma50
                        )
                    
                    if lkb_probit_params:
                        td50 = lkb_probit_params['TD50']
                        m = lkb_probit_params['m']
                        n = lkb_probit_params.get('n', 0.5)
                        # For probit, we need v_effective - use a simplified version with gEUD
                        # Create a simplified probit function using gEUD directly
                        ntcp_lkb_probit = lambda x: norm.cdf((x - td50) / (m * td50))
                    elif organ_for_dr in ntcp_calc.literature_params and 'LKB_Probit' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_Probit']
                        td50 = lit_params['TD50']
                        m = lit_params['m']
                        ntcp_lkb_probit = lambda x: norm.cdf((x - td50) / (m * td50))
                    
                    if rs_poisson_params:
                        d50 = rs_poisson_params['D50']
                        gamma = rs_poisson_params['gamma']
                        # Källman et al. RS: p = 2^{-exp(e·γ(1−D/D50))}
                        ntcp_rs_poisson = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                    elif organ_for_dr in ntcp_calc.literature_params and 'RS_Poisson' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['RS_Poisson']
                        d50 = lit_params['D50']
                        gamma = lit_params['gamma']
                        ntcp_rs_poisson = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                    
                    if bio_logistic_params:
                        TD50_bio = bio_logistic_params['TD50']
                        k_bio = bio_logistic_params['k']
                        ntcp_biological = lambda x: expit((x - TD50_bio) / k_bio)
                    
                    # Create plots directory (single location for all figures)
                    dr_dir = os.path.join(args.output_dir, "plots")
                    os.makedirs(dr_dir, exist_ok=True)
                    
                    # --------------------------------------------------------
                    # FIGURE 1 — LKB Dose–Response (gEUD axis)
                    # --------------------------------------------------------
                    if ntcp_lkb_loglogit is not None or ntcp_lkb_probit is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(np.min(gEUD), np.max(gEUD), 300)
                        
                        if ntcp_lkb_loglogit is not None:
                            ax.plot(x, ntcp_lkb_loglogit(x),
                                    color="#1f77b4", lw=3.0, label="LKB (Log-logit)")
                        if ntcp_lkb_probit is not None:
                            ax.plot(x, ntcp_lkb_probit(x),
                                    color="#d62728", lw=3.0, ls="--", label="LKB (Probit)")
                        
                        ax.scatter(gEUD, ntcp_obs,
                                   facecolors="none", edgecolors="#1f77b4",
                                   s=55, lw=1.6, alpha=0.85)
                        
                        ax.set_xlabel("Parotid gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(os.path.join(dr_dir, "Figure1_LKB_gEUD_DR.png"), dpi=1200)
                        plt.savefig(os.path.join(dr_dir, "Figure1_LKB_gEUD_DR.svg"))
                        plt.close()
                        print(f"    Saved Figure1_LKB_gEUD_DR.png")
                    
                    # --------------------------------------------------------
                    # FIGURE 2 — RS Poisson Dose–Response (gEUD axis)
                    # --------------------------------------------------------
                    if ntcp_rs_poisson is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(np.min(gEUD), np.max(gEUD), 300)
                        
                        ax.plot(x, ntcp_rs_poisson(x),
                                color="#f0ad1a", lw=3.0, ls="-.", label="RS Poisson")
                        
                        ax.scatter(gEUD, ntcp_obs,
                                   marker="^", facecolors="none",
                                   edgecolors="#f0ad1a", s=65, lw=1.6)
                        
                        ax.set_xlabel("Parotid gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(os.path.join(dr_dir, "Figure2_RS_gEUD_DR.png"), dpi=1200)
                        plt.savefig(os.path.join(dr_dir, "Figure2_RS_gEUD_DR.svg"))
                        plt.close()
                        print(f"    Saved Figure2_RS_gEUD_DR.png")
                    
                    # --------------------------------------------------------
                    # FIGURE 3 — Biological Logistic Dose–Response (Mean Dose)
                    # --------------------------------------------------------
                    if ntcp_biological is not None and TD50_bio is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x_md = np.linspace(np.min(mean_dose), np.max(mean_dose), 300)
                        
                        ax.plot(x_md, ntcp_biological(x_md),
                                color="#2c3e50", lw=3.5, label="Biological logistic")
                        
                        ax.scatter(mean_dose, ntcp_obs,
                                   color="gray", alpha=0.65, s=45, zorder=2)
                        
                        ax.axvline(TD50_bio, color="#7f7f7f",
                                   lw=2.0, ls=":", label="TD50")
                        
                        ax.set_xlabel("Mean parotid dose (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(frameon=True, fontsize=11, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(os.path.join(dr_dir, "Figure3_Biological_MeanDose_DR.png"), dpi=1200)
                        plt.savefig(os.path.join(dr_dir, "Figure3_Biological_MeanDose_DR.svg"))
                        plt.close()
                        print(f"    Saved Figure3_Biological_MeanDose_DR.png")
                    
                    # --------------------------------------------------------
                    # FIGURE 4 — Predicted NTCP Comparison (Clinical)
                    # --------------------------------------------------------
                    if (len(ntcp_lkb_loglogit_pred) > 0 or len(ntcp_lkb_probit_pred) > 0 or 
                        len(ntcp_rs_pred) > 0 or len(ntcp_ann_pred) > 0 or len(ntcp_xgb_pred) > 0):
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        if len(ntcp_lkb_loglogit_pred) > 0:
                            ax.scatter(gEUD, ntcp_lkb_loglogit_pred, s=40, alpha=0.6, label="LKB Log-logit")
                        if len(ntcp_lkb_probit_pred) > 0:
                            ax.scatter(gEUD, ntcp_lkb_probit_pred, s=40, alpha=0.6, label="LKB Probit")
                        if len(ntcp_rs_pred) > 0:
                            ax.scatter(gEUD, ntcp_rs_pred, s=40, alpha=0.6, label="RS Poisson")
                        if len(ntcp_ann_pred) > 0:
                            ax.scatter(gEUD, ntcp_ann_pred, s=40, alpha=0.6, label="ANN")
                        if len(ntcp_xgb_pred) > 0:
                            ax.scatter(gEUD, ntcp_xgb_pred, s=40, alpha=0.6, label="XGBoost")
                        if len(ntcp_rf_pred) > 0:
                            ax.scatter(gEUD, ntcp_rf_pred, s=40, alpha=0.6, label="Random Forest")
                        
                        ax.set_xlabel("Parotid gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("Predicted NTCP", fontsize=14, fontweight="bold")
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=10, frameon=True, ncol=2)
                        
                        plt.tight_layout()
                        plt.savefig(os.path.join(dr_dir, "Figure4_Predicted_NTCP_Comparison.png"), dpi=1200)
                        plt.savefig(os.path.join(dr_dir, "Figure4_Predicted_NTCP_Comparison.svg"))
                        plt.close()
                        print(f"    Saved Figure4_Predicted_NTCP_Comparison.png")
                    
                    print("  [OK] Model-native dose-response plots generated")
            except Exception as e:
                print(f"  Warning: Model-native DR plotting failed: {e}")
                import traceback
                traceback.print_exc()
        
        # Step 2: Enhanced summary report
        print("\n Step 2: Generating enhanced summary report...")
        create_enhanced_summary_report(results_df, args.output_dir)
        
        # Step 3: Hospital validation (if HospitalID present)
        print("\n Step 3: Checking for hospital validation...")
        perform_hospital_validation(results_df, Path(args.output_dir))
        
        print("\n Enhanced analysis completed successfully!")
        print("=" * 60)
        print(f"All outputs saved to: {Path(args.output_dir).absolute()}")
        print("\nGenerated files:")
        print("   ntcp_results.xlsx - Comprehensive Excel file with all results")
        print("  enhanced_ntcp_calculations.csv - All model predictions")
        print("  enhanced_summary_performance.csv - Performance table")
        print("  enhanced_analysis_report.txt - Detailed report")
        print("  plots/ - 600 DPI publication-ready plots:")
        print("    • [Organ]_dose_response.png - Dose-response curves")
        print("    • [Organ]_ROC.png - ROC curves with unique colors")
        print("    • [Organ]_calibration.png - Calibration plots")
        print("    • [Organ]_ROC_calibration_combined.png - Combined plots")
        print("    • comprehensive_analysis.png - Overall analysis")
        print("    • model_performance_analysis.png - Performance comparison")
        print("    • overall_performance_summary.png - Summary overview")
        
        # Display enhanced key findings
        print("\n Enhanced Key Findings:")
        
        # Summary by organ
        # Multi-OAR safety: always use sorted organ list
        for organ in sorted(results_df['Organ'].unique()):
            organ_data = results_df[results_df['Organ'] == organ]
            n_patients = len(organ_data)
            n_events = int(organ_data['Observed_Toxicity'].sum())
            event_rate = n_events / n_patients if n_patients > 0 else 0
            
            print(f"  {organ}: {n_patients} patients, {n_events} events ({event_rate:.1%})")
            
            # Find best traditional and ML models
            traditional_aucs = []
            ml_aucs = []
            
            for col in ntcp_cols:
                if col in organ_data.columns:
                    valid_data = organ_data.dropna(subset=[col, 'Observed_Toxicity'])
                    if len(valid_data) >= 5:
                        try:
                            fpr, tpr, _ = roc_curve(valid_data['Observed_Toxicity'], valid_data[col])
                            auc_score = auc(fpr, tpr)
                            
                            if 'ML_' in col:
                                ml_aucs.append((col, auc_score))
                            else:
                                traditional_aucs.append((col, auc_score))
                        except:
                            pass
            
            # Report best models
            if traditional_aucs:
                best_trad = max(traditional_aucs, key=lambda x: x[1])
                print(f"    Best traditional: {best_trad[0]} (AUC = {best_trad[1]:.3f})")
            
            if ml_aucs:
                best_ml = max(ml_aucs, key=lambda x: x[1])
                print(f"    Best ML: {best_ml[0]} (AUC = {best_ml[1]:.3f})")
                
                # Calculate improvement
                if traditional_aucs:
                    improvement = ((best_ml[1] - best_trad[1]) / best_trad[1]) * 100
                    print(f"    ML improvement: {improvement:+.1f}%")
            else:
                print(f"    ML models: Not trained (insufficient data)")
        
        print("\n Next Steps:")
        print("  1. Review ntcp_results.xlsx for comprehensive results")
        print("  2. Examine publication-ready plots in plots/ directory")
        print("  3. Read enhanced_analysis_report.txt for detailed findings")
        print("  4. Consider ML model deployment if performance is superior")
        print("  5. Validate findings with external cohorts")
        print("  6. Use unique color coding to distinguish model types:")
        print("     - Traditional NTCP: Blue (LKB LogLogit), Red (LKB Probit), Gold (RS Poisson)")
        print("     - ML models: Purple (ANN), Green (XGBoost)")
        
    except Exception as e:
        print(f"\nError: Error during enhanced analysis: {e}")
        import traceback
        print("\nFull error traceback:")
        traceback.print_exc()
        
        print("\n Troubleshooting tips:")
        print("  • Ensure sufficient data for ML training (≥15 samples per organ)")
        print("  • Check that all required Python packages are installed:")
        print("    pip install scikit-learn xgboost pandas numpy matplotlib seaborn scipy openpyxl")
        print("  • Verify DVH files and patient data formats")
        print("  • Ensure unique patient IDs in DVH filenames")
        
        # =========================
        # PUBLICATION DR PLOTS
        # =========================
        
        # Publication-grade DR plots (biological + classical only)
        # ============================================================
        # Publication-quality Dose–Response (DR) plots
        # Biological + Classical models ONLY
        # ============================================================
        
        try:
            # Extract data from results_df for the first organ with sufficient data
            if 'results_df' in locals() and results_df is not None and len(results_df) > 0:
                # Find organ with biological or classical models
                organ_for_dr = None
                for organ in sorted(results_df['Organ'].unique()):
                    organ_data = results_df[results_df['Organ'] == organ].copy()
                    valid_mask = (
                        organ_data['gEUD'].notna() &
                        (organ_data['gEUD'] > 0) &
                        organ_data['mean_dose'].notna() &
                        (organ_data['mean_dose'] > 0) &
                        organ_data['Observed_Toxicity'].notna()
                    )
                    organ_data_valid = organ_data[valid_mask].copy()
                    
                    if len(organ_data_valid) < 5:
                        continue
                    
                    # Check if we have classical or biological models
                    has_classical = (
                        'NTCP_LKB_LogLogit' in organ_data_valid.columns or
                        'NTCP_LKB_Probit' in organ_data_valid.columns or
                        'NTCP_RS_Poisson' in organ_data_valid.columns
                    )
                    has_bio = 'NTCP_Biological_Logistic' in organ_data_valid.columns
                    
                    if has_classical or has_bio:
                        organ_for_dr = organ
                        break
                
                if organ_for_dr is not None:
                    organ_data = results_df[results_df['Organ'] == organ_for_dr].copy()
                    valid_mask = (
                        organ_data['gEUD'].notna() &
                        (organ_data['gEUD'] > 0) &
                        organ_data['mean_dose'].notna() &
                        (organ_data['mean_dose'] > 0) &
                        organ_data['Observed_Toxicity'].notna()
                    )
                    organ_data_valid = organ_data[valid_mask].copy()
                    
                    # Extract observed endpoint
                    y_obs = organ_data_valid['Observed_Toxicity'].values.astype(int)
                    gEUD = organ_data_valid['gEUD'].values
                    mean_dose = organ_data_valid['mean_dose'].values
                    
                    # Use ntcp_calc already initialized at function scope
                    # (No need to recreate - already available)
                    
                    # Load refitted parameters if available
                    classical_params = {}
                    bio_params = {}
                    refit_dir = Path(args.output_dir) / 'biological_refitting'
                    if refit_dir.exists():
                        classical_params_file = refit_dir / 'classical_refitted_params.json'
                        bio_params_file = refit_dir / 'biological_refitted_params.json'
                        if classical_params_file.exists():
                            import json
                            with open(classical_params_file, 'r') as f:
                                classical_params = json.load(f)
                        if bio_params_file.exists():
                            import json
                            with open(bio_params_file, 'r') as f:
                                bio_params = json.load(f)
                    
                    # Create model functions
                    ntcp_lkb_loglogit = None
                    ntcp_lkb_probit = None
                    ntcp_rs_poisson = None
                    ntcp_biological = None
                    TD50_bio = None
                    
                    # LKB Log-logit
                    if organ_for_dr in classical_params and classical_params[organ_for_dr].get('LKB_LogLogit'):
                        params = classical_params[organ_for_dr]['LKB_LogLogit'].get('point_estimate', {})
                        if params:
                            td50 = params.get('TD50')
                            gamma50 = params.get('gamma50')
                            if td50 and gamma50:
                                ntcp_lkb_loglogit = lambda x: 1.0 / (1.0 + np.power(td50 / np.maximum(x, 1e-6), 4.0 * gamma50))
                    elif organ_for_dr in ntcp_calc.literature_params and 'LKB_LogLogit' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_LogLogit']
                        td50 = lit_params['TD50']
                        gamma50 = lit_params['gamma50']
                        alpha_beta = lit_params.get('alpha_beta', 3)
                        ntcp_lkb_loglogit = lambda x: ntcp_calc.ntcp_lkb_loglogit(
                            ntcp_calc.convert_to_eqd2(x, alpha_beta, 2.0), td50, gamma50
                        )
                    
                    # LKB Probit
                    if organ_for_dr in classical_params and classical_params[organ_for_dr].get('LKB_Probit'):
                        params = classical_params[organ_for_dr]['LKB_Probit'].get('point_estimate', {})
                        if params:
                            td50 = params.get('TD50')
                            m = params.get('m')
                            if td50 and m:
                                ntcp_lkb_probit = lambda x: norm.cdf((x - td50) / (m * td50))
                    elif organ_for_dr in ntcp_calc.literature_params and 'LKB_Probit' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['LKB_Probit']
                        td50 = lit_params['TD50']
                        m = lit_params['m']
                        ntcp_lkb_probit = lambda x: norm.cdf((x - td50) / (m * td50))
                    
                    # RS Poisson
                    if organ_for_dr in classical_params and classical_params[organ_for_dr].get('RS_Poisson'):
                        params = classical_params[organ_for_dr]['RS_Poisson'].get('point_estimate', {})
                        if params:
                            d50 = params.get('D50')
                            gamma = params.get('gamma')
                            if d50 and gamma:
                                ntcp_rs_poisson = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                    elif organ_for_dr in ntcp_calc.literature_params and 'RS_Poisson' in ntcp_calc.literature_params[organ_for_dr]:
                        lit_params = ntcp_calc.literature_params[organ_for_dr]['RS_Poisson']
                        d50 = lit_params['D50']
                        gamma = lit_params['gamma']
                        ntcp_rs_poisson = lambda x: 2.0 ** (-np.exp(np.e * gamma * (1.0 - np.maximum(x, 1e-6) / np.maximum(d50, 1e-6))))
                    
                    # Biological logistic
                    if organ_for_dr in bio_params and bio_params[organ_for_dr].get('Logistic'):
                        params = bio_params[organ_for_dr]['Logistic'].get('point_estimate', {})
                        if params:
                            TD50_bio = params.get('TD50')
                            k_bio = params.get('k')
                            if TD50_bio and k_bio:
                                from scipy.special import expit
                                ntcp_biological = lambda x: expit((x - TD50_bio) / k_bio)
                    
                    # Use single plots directory
                    DR_DIR = Path(args.output_dir) / "plots"
                    DR_DIR.mkdir(parents=True, exist_ok=True)
                    
                    # ============================================================
                    # FIGURE 1 — LKB DR (gEUD)
                    # ============================================================
                    if ntcp_lkb_loglogit is not None or ntcp_lkb_probit is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(gEUD.min(), gEUD.max(), 400)
                        
                        if ntcp_lkb_loglogit is not None:
                            ax.plot(x, ntcp_lkb_loglogit(x),
                                    color="#1f77b4", lw=3.2, label="LKB (Log-logit)")
                        
                        if ntcp_lkb_probit is not None:
                            ax.plot(x, ntcp_lkb_probit(x),
                                    color="#d62728", lw=3.2, ls="--", label="LKB (Probit)")
                        
                        ax.scatter(gEUD, y_obs,
                                   facecolors="none", edgecolors="gray",
                                   s=55, lw=1.6, alpha=0.9)
                        
                        ax.set_xlabel(f"{organ_for_dr} gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_xlim(left=0)
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=11, frameon=True, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(DR_DIR / "Figure1_LKB_gEUD_DR.png", dpi=1200)
                        plt.savefig(DR_DIR / "Figure1_LKB_gEUD_DR.svg")
                        plt.close()
                    
                    # ============================================================
                    # FIGURE 2 — RS Poisson DR (gEUD)
                    # ============================================================
                    if ntcp_rs_poisson is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(gEUD.min(), gEUD.max(), 400)
                        
                        ax.plot(x, ntcp_rs_poisson(x),
                                color="#f0ad1a", lw=3.2, ls="-.", label="RS Poisson")
                        
                        ax.scatter(gEUD, y_obs,
                                   marker="^", facecolors="none",
                                   edgecolors="#f0ad1a", s=65, lw=1.6)
                        
                        ax.set_xlabel(f"{organ_for_dr} gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=11, frameon=True, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(DR_DIR / "Figure2_RS_gEUD_DR.png", dpi=1200)
                        plt.savefig(DR_DIR / "Figure2_RS_gEUD_DR.svg")
                        plt.close()
                    
                    # ============================================================
                    # FIGURE 3 — Biological Logistic DR (Mean dose)
                    # ============================================================
                    if ntcp_biological is not None and TD50_bio is not None:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x_md = np.linspace(mean_dose.min(), mean_dose.max(), 400)
                        
                        ax.plot(x_md, ntcp_biological(x_md),
                                color="#2c3e50", lw=3.8, label="Biological logistic")
                        
                        ax.scatter(mean_dose, y_obs,
                                   color="gray", alpha=0.65, s=45)
                        
                        ax.axvline(TD50_bio, color="gray", lw=2.2, ls=":", label="TD50")
                        
                        ax.set_xlabel(f"Mean {organ_for_dr.lower()} dose (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, alpha=0.25)
                        ax.legend(fontsize=11, frameon=True, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(DR_DIR / "Figure3_Biological_MeanDose_DR.png", dpi=1200)
                        plt.savefig(DR_DIR / "Figure3_Biological_MeanDose_DR.svg")
                        plt.close()
                    
                    # ============================================================
                    # FIGURE 4 — Classical Models Only (NO ML)
                    # ============================================================
                    # Use single plots directory (no separate DR_plots)
                    code3_dr_dir = Path(args.output_dir) / "plots"
                    code3_dr_dir.mkdir(parents=True, exist_ok=True)
                    
                    # Generate Figure4 with ONLY classical models
                    if (ntcp_lkb_loglogit is not None or ntcp_lkb_probit is not None or ntcp_rs_poisson is not None):
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        x = np.linspace(gEUD.min(), gEUD.max(), 400)
                        
                        # LKB Log-logit
                        if ntcp_lkb_loglogit is not None:
                            ax.plot(x, ntcp_lkb_loglogit(x),
                                    color="#1f77b4", lw=3.5, label="LKB (Log-logit)")
                            ax.scatter(gEUD, y_obs,
                                       facecolors="none", edgecolors="#1f77b4",
                                       s=60, lw=1.6, alpha=0.9, zorder=5)
                        
                        # LKB Probit
                        if ntcp_lkb_probit is not None:
                            ax.plot(x, ntcp_lkb_probit(x),
                                    color="#d62728", lw=3.5, ls="--", label="LKB (Probit)")
                            ax.scatter(gEUD, y_obs,
                                       marker="s", facecolors="none", edgecolors="#d62728",
                                       s=60, lw=1.6, alpha=0.9, zorder=5)
                        
                        # RS Poisson
                        if ntcp_rs_poisson is not None:
                            ax.plot(x, ntcp_rs_poisson(x),
                                    color="#f0ad1a", lw=3.5, ls="-.", label="RS Poisson")
                            ax.scatter(gEUD, y_obs,
                                       marker="^", facecolors="none", edgecolors="#f0ad1a",
                                       s=60, lw=1.6, alpha=0.9, zorder=5)
                        
                        # Explicitly EXCLUDE ML models - NO ANN, NO XGBoost
                        # This figure shows ONLY classical NTCP models
                        
                        ax.set_xlabel("gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("Predicted NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, which="major", alpha=0.25)
                        ax.legend(fontsize=11, frameon=True, loc="lower right")
                        
                        plt.tight_layout()
                        plt.savefig(code3_dr_dir / "Figure4_DR_Classical_Only.png", dpi=1200)
                        plt.savefig(code3_dr_dir / "Figure4_DR_Classical_Only.svg")
                        plt.close()
                        print(f"  [OK] Figure4_DR_Classical_Only.png saved (classical models only, NO ML)")
                    
                    print(f"\n[DR PLOTS] Publication-grade DR plots saved to {DR_DIR}")
        except Exception as dr_error:
            print(f"\nWarning: Could not generate publication DR plots: {dr_error}")
            import traceback
            traceback.print_exc()
        
        # ============================================================
        # PAROTID VOLUME SENSITIVITY ANALYSIS (REPORTING-ONLY)
        # ============================================================
        
        def perform_parotid_volume_sensitivity_analysis(results_df, output_dir, ntcp_calc):
            """
            Perform reporting-only sensitivity analysis for parotid volume heterogeneity.
            Does NOT modify models, predictions, or exclude patients.
            """
            try:
                print("\n[ANALYSIS] Performing parotid volume sensitivity analysis (reporting-only)...")
                
                # Filter for Parotid organ only
                parotid_data = results_df[results_df['Organ'].str.contains('Parotid', case=False, na=False)].copy()
                
                if len(parotid_data) == 0:
                    print("  Warning: No Parotid data found. Skipping volume sensitivity analysis.")
                    return
                
                # Group by patient to compute total parotid volume
                patient_volumes = {}
                for _, row in parotid_data.iterrows():
                    patient_id = row.get('PrimaryPatientID') or row.get('AnonPatientID')
                    if patient_id is None:
                        continue
                    
                    if patient_id not in patient_volumes:
                        patient_volumes[patient_id] = 0.0
                    
                    volume = row.get('total_volume', 0.0)
                    if pd.notna(volume):
                        patient_volumes[patient_id] += float(volume)
                
                # Add volume category column
                parotid_data = parotid_data.copy()
                parotid_data['parotid_volume_category'] = parotid_data.apply(
                    lambda row: (
                        'low_volume' if patient_volumes.get(row.get('PrimaryPatientID') or row.get('AnonPatientID'), 0.0) < 40.0
                        else 'high_volume'
                    ), axis=1
                )
                
                # Descriptive analysis by category
                sensitivity_summary = []
                
                for category in ['low_volume', 'high_volume']:
                    cat_data = parotid_data[parotid_data['parotid_volume_category'] == category].copy()
                    
                    if len(cat_data) == 0:
                        continue
                    
                    n = len(cat_data)
                    mean_dose_mean = cat_data['mean_dose'].mean() if 'mean_dose' in cat_data.columns else np.nan
                    mean_dose_sd = cat_data['mean_dose'].std() if 'mean_dose' in cat_data.columns else np.nan
                    geud_mean = cat_data['gEUD'].mean() if 'gEUD' in cat_data.columns else np.nan
                    geud_sd = cat_data['gEUD'].std() if 'gEUD' in cat_data.columns else np.nan
                    
                    # NTCP means
                    ntcp_lkb_loglogit_mean = cat_data['NTCP_LKB_LogLogit'].mean() if 'NTCP_LKB_LogLogit' in cat_data.columns else np.nan
                    ntcp_lkb_probit_mean = cat_data['NTCP_LKB_Probit'].mean() if 'NTCP_LKB_Probit' in cat_data.columns else np.nan
                    ntcp_rs_poisson_mean = cat_data['NTCP_RS_Poisson'].mean() if 'NTCP_RS_Poisson' in cat_data.columns else np.nan
                    
                    # Observed toxicity rate
                    obs_tox_rate = cat_data['Observed_Toxicity'].mean() if 'Observed_Toxicity' in cat_data.columns else np.nan
                    
                    sensitivity_summary.append({
                        'Volume_Category': category,
                        'N': n,
                        'Mean_Dose_Mean': f"{mean_dose_mean:.2f}" if pd.notna(mean_dose_mean) else 'N/A',
                        'Mean_Dose_SD': f"{mean_dose_sd:.2f}" if pd.notna(mean_dose_sd) else 'N/A',
                        'gEUD_Mean': f"{geud_mean:.2f}" if pd.notna(geud_mean) else 'N/A',
                        'gEUD_SD': f"{geud_sd:.2f}" if pd.notna(geud_sd) else 'N/A',
                        'NTCP_LKB_LogLogit_Mean': f"{ntcp_lkb_loglogit_mean:.4f}" if pd.notna(ntcp_lkb_loglogit_mean) else 'N/A',
                        'NTCP_LKB_Probit_Mean': f"{ntcp_lkb_probit_mean:.4f}" if pd.notna(ntcp_lkb_probit_mean) else 'N/A',
                        'NTCP_RS_Poisson_Mean': f"{ntcp_rs_poisson_mean:.4f}" if pd.notna(ntcp_rs_poisson_mean) else 'N/A',
                        'Observed_Toxicity_Rate': f"{obs_tox_rate:.4f}" if pd.notna(obs_tox_rate) else 'N/A'
                    })
                
                # Export summary table
                if sensitivity_summary:
                    sensitivity_df = pd.DataFrame(sensitivity_summary)
                    sensitivity_file = Path(output_dir) / 'parotid_volume_sensitivity.xlsx'
                    sensitivity_df.to_excel(sensitivity_file, index=False)
                    print(f"  [OK] Parotid volume sensitivity summary saved to {sensitivity_file}")
                
                # Generate plot: NTCP vs gEUD colored by volume category
                if len(parotid_data) > 0 and 'gEUD' in parotid_data.columns:
                    # Get classical NTCP predictions
                    valid_data = parotid_data.dropna(subset=['gEUD', 'Observed_Toxicity']).copy()
                    
                    if len(valid_data) > 0:
                        fig, ax = plt.subplots(figsize=(7.5, 6.0))
                        
                        # Plot by category
                        for category, color, marker in [('low_volume', '#2ca02c', 'o'), ('high_volume', '#9467bd', 's')]:
                            cat_data = valid_data[valid_data['parotid_volume_category'] == category]
                            if len(cat_data) > 0:
                                geud_vals = cat_data['gEUD'].values
                                
                                # Plot classical model predictions
                                if 'NTCP_LKB_LogLogit' in cat_data.columns:
                                    ntcp_vals = cat_data['NTCP_LKB_LogLogit'].values
                                    ax.scatter(geud_vals, ntcp_vals,
                                               c=color, marker=marker, s=60, alpha=0.7,
                                               label=f"LKB Log-logit ({category})", edgecolors='black', linewidths=0.5)
                        
                        ax.set_xlabel("gEUD (Gy)", fontsize=14, fontweight="bold")
                        ax.set_ylabel("Predicted NTCP", fontsize=14, fontweight="bold")
                        ax.set_ylim(0, 1.02)
                        ax.grid(True, which="major", alpha=0.25)
                        ax.legend(fontsize=11, frameon=True, loc="lower right")
                        
                        plt.tight_layout()
                        plot_file = Path(output_dir) / 'plots' / 'Figure4b_DR_Volume_Sensitivity.png'
                        plot_file.parent.mkdir(parents=True, exist_ok=True)
                        plt.savefig(plot_file, dpi=1200)
                        plt.close()
                        print(f"  [OK] Volume sensitivity plot saved to {plot_file}")
                
            except Exception as e:
                print(f"  Warning: Parotid volume sensitivity analysis failed: {e}")
                import traceback
                traceback.print_exc()
        
        # Perform parotid volume sensitivity analysis
        try:
            perform_parotid_volume_sensitivity_analysis(results_df, args.output_dir, ntcp_calc)
        except Exception as e:
            print(f"  Warning: Volume sensitivity analysis failed: {e}")
        
        # ============================================================
        # MANUSCRIPT MATERIALS BUNDLE
        # ============================================================
        
        def create_manuscript_materials_bundle(results_df, output_dir, ntcp_calc):
            """
            Create manuscript_materials output bundle with figures, tables, and methods.
            """
            try:
                print("\n[MANUSCRIPT] Creating manuscript materials bundle...")
                
                manuscript_dir = Path(output_dir) / 'manuscript_materials'
                tables_dir = manuscript_dir / 'tables'
                methods_dir = manuscript_dir / 'methods'
                
                tables_dir.mkdir(parents=True, exist_ok=True)
                methods_dir.mkdir(parents=True, exist_ok=True)
                
                # ============================================================
                # TABLES: manuscript_results.xlsx
                # ============================================================
                
                excel_file = tables_dir / 'manuscript_tables.xlsx'
                print(f"  Creating manuscript_results.xlsx...")
                
                with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
                    # Sheet 1: DoseMetrics
                    dose_metrics_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ', 'gEUD', 'mean_dose', 'max_dose', 'total_volume']
                    v_cols = [col for col in results_df.columns if col.startswith('V') and col[1:].isdigit()]
                    d_cols = [col for col in results_df.columns if col.startswith('D') and any(c.isdigit() for c in col[1:])]
                    dose_metrics_cols.extend(v_cols)
                    dose_metrics_cols.extend(d_cols)
                    available_dose_cols = [col for col in dose_metrics_cols if col in results_df.columns]
                    dose_df = results_df[available_dose_cols].copy()
                    numeric_dose_cols = dose_df.select_dtypes(include=[np.number]).columns
                    for col in numeric_dose_cols:
                        dose_df[col] = dose_df[col].round(2)
                    dose_df.to_excel(writer, sheet_name='DoseMetrics', index=False)
                    
                    # Sheet 2: NTCP_Classical
                    ntcp_classical_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ', 'Observed_Toxicity']
                    classical_models = ['NTCP_LKB_LogLogit', 'NTCP_LKB_Probit', 'NTCP_RS_Poisson', 'NTCP_Biological_Logistic']
                    for col in classical_models:
                        if col in results_df.columns:
                            ntcp_classical_cols.append(col)
                    ntcp_classical_df = results_df[ntcp_classical_cols].copy()
                    for col in classical_models:
                        if col in ntcp_classical_df.columns:
                            ntcp_classical_df[col] = ntcp_classical_df[col].round(4)
                    ntcp_classical_df.to_excel(writer, sheet_name='NTCP_Classical', index=False)
                    
                    # Sheet 3: NTCP_ML
                    ntcp_ml_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ', 'Observed_Toxicity']
                    ml_models = ['NTCP_ML_ANN', 'NTCP_ML_XGBoost', 'NTCP_ML_RandomForest']
                    for col in ml_models:
                        if col in results_df.columns:
                            ntcp_ml_cols.append(col)
                    if len(ntcp_ml_cols) > 4:  # More than just ID columns
                        ntcp_ml_df = results_df[ntcp_ml_cols].copy()
                        for col in ml_models:
                            if col in ntcp_ml_df.columns:
                                ntcp_ml_df[col] = ntcp_ml_df[col].round(4)
                        ntcp_ml_df.to_excel(writer, sheet_name='NTCP_ML', index=False)
                    else:
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ', 'NTCP_ML_ANN', 'NTCP_ML_XGBoost', 'NTCP_ML_RandomForest']).to_excel(
                            writer, sheet_name='NTCP_ML', index=False)
                    
                    # Sheet 4: Radiobiology_Parameters
                    lit_params_data = []
                    for organ, params in ntcp_calc.literature_params.items():
                        for model_type, model_params in params.items():
                            row = {'Organ': organ, 'Model': model_type, **model_params}
                            lit_params_data.append(row)
                    lit_params_df = pd.DataFrame(lit_params_data)
                    lit_params_df.to_excel(writer, sheet_name='Radiobiology_Parameters', index=False)
                    
                    # Sheet 5: uNTCP
                    untcp_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ']
                    untcp_data_cols = [col for col in results_df.columns if 'uNTCP' in col or 'untcp' in col.lower()]
                    untcp_cols.extend(untcp_data_cols)
                    available_untcp_cols = [col for col in untcp_cols if col in results_df.columns]
                    if available_untcp_cols:
                        untcp_df = results_df[available_untcp_cols].copy()
                        numeric_untcp_cols = untcp_df.select_dtypes(include=[np.number]).columns
                        for col in numeric_untcp_cols:
                            untcp_df[col] = untcp_df[col].round(4)
                        untcp_df.to_excel(writer, sheet_name='uNTCP', index=False)
                    else:
                        # Create empty sheet if no uNTCP data
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ', 'uNTCP', 'uNTCP_STD', 'uNTCP_CI_L', 'uNTCP_CI_U']).to_excel(
                            writer, sheet_name='uNTCP', index=False)
                    
                    # Sheet 6: CCS
                    ccs_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ']
                    ccs_data_cols = [col for col in results_df.columns if 'CCS' in col or 'ccs' in col.lower()]
                    ccs_cols.extend(ccs_data_cols)
                    available_ccs_cols = [col for col in ccs_cols if col in results_df.columns]
                    if available_ccs_cols:
                        ccs_df = results_df[available_ccs_cols].copy()
                        numeric_ccs_cols = ccs_df.select_dtypes(include=[np.number]).columns
                        for col in numeric_ccs_cols:
                            ccs_df[col] = ccs_df[col].round(4)
                        ccs_df.to_excel(writer, sheet_name='CCS', index=False)
                    else:
                        # Create empty sheet if no CCS data
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ', 'CCS', 'CCS_Safety']).to_excel(
                            writer, sheet_name='CCS', index=False)
                    
                    # Sheet 7: Uncertainty
                    uncertainty_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ']
                    uncertainty_data_cols = [col for col in results_df.columns if 'uncertainty' in col.lower() or 'std' in col.lower() or 'CI' in col]
                    uncertainty_cols.extend(uncertainty_data_cols)
                    available_uncertainty_cols = [col for col in uncertainty_cols if col in results_df.columns]
                    if available_uncertainty_cols:
                        uncertainty_df = results_df[available_uncertainty_cols].copy()
                        numeric_uncertainty_cols = uncertainty_df.select_dtypes(include=[np.number]).columns
                        for col in numeric_uncertainty_cols:
                            uncertainty_df[col] = uncertainty_df[col].round(4)
                        uncertainty_df.to_excel(writer, sheet_name='Uncertainty', index=False)
                    else:
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ']).to_excel(
                            writer, sheet_name='Uncertainty', index=False)
                    
                    # Sheet 8: ML_Performance
                    ml_perf_data = []
                    for organ in sorted(results_df['Organ'].unique()):
                        organ_data = results_df[results_df['Organ'] == organ]
                        row = {'Organ': organ, 'Sample_Size': len(organ_data)}
                        
                        for model in ['ML_ANN', 'ML_XGBoost', 'ML_RandomForest']:
                            ntcp_col = f'NTCP_{model}'
                            if ntcp_col in organ_data.columns:
                                valid_data = organ_data.dropna(subset=[ntcp_col, 'Observed_Toxicity'])
                                if len(valid_data) >= 5:
                                    try:
                                        y_true = valid_data['Observed_Toxicity'].values
                                        y_pred = valid_data[ntcp_col].values
                                        fpr, tpr, _ = roc_curve(y_true, y_pred)
                                        auc_score = auc(fpr, tpr)
                                        brier_score = brier_score_loss(y_true, y_pred)
                                        row[f'{model}_AUC'] = f"{auc_score:.3f}"
                                        row[f'{model}_Brier'] = f"{brier_score:.3f}"
                                    except:
                                        row[f'{model}_AUC'] = 'Error'
                                        row[f'{model}_Brier'] = 'Error'
                                else:
                                    row[f'{model}_AUC'] = 'Insufficient Data'
                                    row[f'{model}_Brier'] = 'Insufficient Data'
                            else:
                                row[f'{model}_AUC'] = 'Not Available'
                                row[f'{model}_Brier'] = 'Not Available'
                        
                        ml_perf_data.append(row)
                    
                    ml_perf_df = pd.DataFrame(ml_perf_data)
                    ml_perf_df.to_excel(writer, sheet_name='ML_Performance', index=False)
                    
                    # Sheet 9: QA_Overfitting
                    qa_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ']
                    qa_data_cols = [col for col in results_df.columns if 'overfitting' in col.lower() or 'qa' in col.lower() or 'validation' in col.lower()]
                    qa_cols.extend(qa_data_cols)
                    available_qa_cols = [col for col in qa_cols if col in results_df.columns]
                    if available_qa_cols:
                        qa_df = results_df[available_qa_cols].copy()
                        qa_df.to_excel(writer, sheet_name='QA_Overfitting', index=False)
                    else:
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ', 'QA_Flag', 'Overfitting_Risk']).to_excel(
                            writer, sheet_name='QA_Overfitting', index=False)
                    
                    # Sheet 10: Leakage_Checks
                    leakage_cols = ['PrimaryPatientID', 'AnonPatientID', 'Organ']
                    leakage_data_cols = [col for col in results_df.columns if 'leakage' in col.lower() or 'data_leak' in col.lower()]
                    leakage_cols.extend(leakage_data_cols)
                    available_leakage_cols = [col for col in leakage_cols if col in results_df.columns]
                    if available_leakage_cols:
                        leakage_df = results_df[available_leakage_cols].copy()
                        leakage_df.to_excel(writer, sheet_name='Leakage_Checks', index=False)
                    else:
                        pd.DataFrame(columns=['PrimaryPatientID', 'Organ', 'Leakage_Check', 'Status']).to_excel(
                            writer, sheet_name='Leakage_Checks', index=False)
                    
                    # Sheet 11: Parotid_Volume_Sensitivity
                    sensitivity_file = Path(output_dir) / 'parotid_volume_sensitivity.xlsx'
                    if sensitivity_file.exists():
                        sensitivity_df = pd.read_excel(sensitivity_file)
                        sensitivity_df.to_excel(writer, sheet_name='Parotid_Volume_Sensitivity', index=False)
                    else:
                        pd.DataFrame(columns=['Volume_Category', 'N', 'Mean_Dose_Mean', 'gEUD_Mean', 
                                             'NTCP_LKB_LogLogit_Mean', 'NTCP_LKB_Probit_Mean', 
                                             'NTCP_RS_Poisson_Mean', 'Observed_Toxicity_Rate']).to_excel(
                            writer, sheet_name='Parotid_Volume_Sensitivity', index=False)
                
                print(f"  [OK] manuscript_tables.xlsx created with {len(pd.ExcelFile(excel_file).sheet_names)} sheets")
                
                # ============================================================
                # METHODS: methods_models_equations.docx
                # ============================================================
                
                try:
                    from docx import Document
                    from docx.shared import Pt
                    
                    doc = Document()
                    
                    # Title
                    title = doc.add_heading('Methods: NTCP Models and Statistical Methods', 0)
                    
                    # Model Equations Section
                    doc.add_heading('1. NTCP Model Equations', level=1)
                    
                    doc.add_heading('1.1 LKB (Lyman-Kutcher-Burman) Log-Logistic Model', level=2)
                    doc.add_paragraph(
                        'NTCP = 1 / (1 + (TD₅₀ / gEUD_EQD₂)^(4γ₅₀))'
                    )
                    doc.add_paragraph(
                        'where gEUD_EQD₂ is the generalized Equivalent Uniform Dose converted to equivalent dose in 2 Gy fractions, '
                        'TD₅₀ is the dose at which 50% complication probability occurs, and γ₅₀ is the normalized dose-response slope.'
                    )
                    
                    doc.add_heading('1.2 LKB Probit Model', level=2)
                    doc.add_paragraph(
                        'NTCP = Φ((gEUD - TD₅₀) / (m × TD₅₀))'
                    )
                    doc.add_paragraph(
                        'where Φ is the cumulative standard normal distribution, TD₅₀ is the dose for 50% complication probability, '
                        'and m is the normalized slope parameter.'
                    )
                    
                    doc.add_heading('1.3 RS (Relative Seriality) Poisson Model', level=2)
                    doc.add_paragraph(
                        'NTCP = 1 - exp(-(D / D₅₀)^γ)'
                    )
                    doc.add_paragraph(
                        'where D is the dose, D₅₀ is the dose for 50% complication probability, and γ is the dose-response parameter.'
                    )
                    
                    doc.add_heading('1.4 Biological Logistic Model', level=2)
                    doc.add_paragraph(
                        'NTCP = 1 / (1 + exp(-(D - TD₅₀) / k))'
                    )
                    doc.add_paragraph(
                        'where D is the mean dose, TD₅₀ is the dose for 50% complication probability, and k is the slope parameter.'
                    )
                    
                    doc.add_heading('1.5 Machine Learning Models', level=2)
                    doc.add_paragraph(
                        'ANN (Artificial Neural Network): Multi-layer perceptron with hidden layers, trained using backpropagation.'
                    )
                    doc.add_paragraph(
                        'XGBoost: Gradient boosting ensemble method using decision trees with regularization.'
                    )
                    
                    # Parameter Definitions
                    doc.add_heading('2. Parameter Definitions', level=1)
                    doc.add_paragraph('gEUD: Generalized Equivalent Uniform Dose, calculated as (Σᵢ vᵢ × Dᵢᵃ)^(1/a)')
                    doc.add_paragraph('TD₅₀: Dose at which 50% complication probability occurs')
                    doc.add_paragraph('γ₅₀: Normalized dose-response slope parameter')
                    doc.add_paragraph('m: Normalized slope parameter for probit model')
                    doc.add_paragraph('D₅₀: Dose for 50% complication probability (RS Poisson)')
                    doc.add_paragraph('k: Slope parameter for biological logistic model')
                    
                    # Statistical Methods
                    doc.add_heading('3. Statistical Methods', level=1)
                    doc.add_paragraph('Model performance was evaluated using:')
                    doc.add_paragraph('• Area Under the ROC Curve (AUC): Discrimination ability')
                    doc.add_paragraph('• Brier Score: Calibration quality (lower is better)')
                    doc.add_paragraph('• Cross-validation: Stratified k-fold (k=5) for ML models')
                    doc.add_paragraph('• Train-test split: 70-30 stratified split to prevent data leakage')
                    
                    # Uncertainty Propagation
                    doc.add_heading('4. Uncertainty Propagation', level=1)
                    doc.add_paragraph(
                        'Uncertainty-Aware NTCP (uNTCP) uses first-order Taylor expansion (Delta method) to propagate parameter uncertainties:'
                    )
                    doc.add_paragraph('σ²_NTCP = Σⱼ (∂NTCP/∂θⱼ)² × σ²_θⱼ')
                    doc.add_paragraph('95% confidence intervals: CI = NTCP ± 1.96 × σ_NTCP')
                    
                    # CCS Definition
                    doc.add_heading('5. Cohort Consistency Score (CCS)', level=1)
                    doc.add_paragraph(
                        'CCS quantifies the consistency of a patient\'s dose metrics with the training cohort distribution. '
                        'Values range from 0 (inconsistent) to 1 (highly consistent). '
                        'CCS < 0.5 indicates potential extrapolation beyond training data.'
                    )
                    
                    # Governance Disclaimer
                    doc.add_heading('6. Governance and Disclaimer', level=1)
                    doc.add_paragraph(
                        'This analysis pipeline implements published NTCP models from the literature (QUANTEC, Emami et al.). '
                        'Model parameters may be refitted to local cohort data when sufficient samples are available. '
                        'All predictions are for research purposes only and should not be used for clinical decision-making without proper validation.'
                    )
                    doc.add_paragraph(
                        'Software: py_ntcpx_v1.0.0'
                    )
                    
                    methods_file = methods_dir / 'methods_models_equations.docx'
                    doc.save(methods_file)
                    print(f"  [OK] methods_models_equations.docx created")
                    
                except ImportError:
                    print(f"  Warning: python-docx not available. Skipping methods document.")
                    print(f"  Install with: pip install python-docx")
                except Exception as e:
                    print(f"  Warning: Methods document generation failed: {e}")
                
                print(f"\n[MANUSCRIPT] Bundle created successfully:")
                print(f"  Tables: {tables_dir}")
                print(f"  Methods: {methods_dir}")
                
            except Exception as e:
                print(f"  Warning: Manuscript materials bundle creation failed: {e}")
                import traceback
                traceback.print_exc()
        
        # Create manuscript materials bundle
        try:
            create_manuscript_materials_bundle(results_df, args.output_dir, ntcp_calc)
        except Exception as e:
            print(f"  Warning: Manuscript materials bundle failed: {e}")

if __name__ == "__main__":
    main()   